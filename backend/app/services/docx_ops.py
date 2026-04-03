from __future__ import annotations
from typing import List, Optional
from docx.document import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def _get_heading_level(p: Paragraph) -> Optional[int]:
    if not p.style or not p.style.name:
        return None
    name = p.style.name.lower()

    # works for "Heading 1/2/3" and French "Titre 1/2/3"
    if "heading" in name or "titre" in name:
        for n in ("1", "2", "3", "4", "5", "6", "7", "8", "9"):
            if name.strip().endswith(n) or f" {n}" in name:
                return int(n)
        # fallback: "Heading" without number
        return 1
    return None


def _iter_block_items(doc: Document):
    """
    Yield paragraphs and tables in document order.
    """
    parent = doc.element.body
    for child in parent.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def remove_sections_by_heading_titles(doc: Document, titles_to_remove: List[str]) -> None:
    """
    Remove blocks under headings whose text equals any of titles_to_remove,
    until next heading of same or higher level.
    """
    titles_set = {t.strip() for t in titles_to_remove if t and t.strip()}
    if not titles_set:
        return

    blocks = list(_iter_block_items(doc))

    i = 0
    while i < len(blocks):
        b = blocks[i]

        if isinstance(b, Paragraph):
            txt = (b.text or "").strip()
            lvl = _get_heading_level(b)
            if lvl is not None and txt in titles_set:
                # find end boundary
                j = i + 1
                while j < len(blocks):
                    bj = blocks[j]
                    if isinstance(bj, Paragraph):
                        lvl2 = _get_heading_level(bj)
                        if lvl2 is not None and lvl2 <= lvl:
                            break
                    j += 1

                # delete from i to j-1
                for k in range(i, j):
                    elm = blocks[k]._element
                    elm.getparent().remove(elm)

                # rebuild blocks list after deletion
                blocks = list(_iter_block_items(doc))
                # continue from same index i (new content shifted)
                continue

        i += 1


from docx import Document
import re

# accepte: [[INSERER]], [[insérer]], [[insert]], [[insertion schéma 1]] etc.
INSERER_RE = re.compile(r"\[\[\s*(?:ins[eé]rer|insert(?:ion)?)(?:[^\]]*)\]\]", re.IGNORECASE)
SEC_TAG_RE = re.compile(r"\[\[\s*(SEC_[A-Za-z0-9_]+)\s*\]\]")
TEXT_PLACEHOLDER_RE = re.compile(r"\[\[TEXT:(.*?)\]\]")
LOGO_PLACEHOLDER_RE = re.compile(r"\[\[LOGO\]\]")

def _extract_inserer_labels(text: str) -> list[str]:
    labels: list[str] = []
    for match in INSERER_RE.finditer(text or ""):
        raw = match.group(0)
        inner = raw.replace("[[", "").replace("]]", "").strip()
        if inner:
            labels.append(inner)
    return labels

def _normalize_title(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip()).lower()

def docx_has_inserer(path: str) -> bool:
    doc = Document(path)

    # Paragraphes
    for p in doc.paragraphs:
        if INSERER_RE.search(p.text or ""):
            return True

    # Tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if INSERER_RE.search(p.text or ""):
                        return True

    return False

def docx_inserer_heading_titles(path: str) -> set[str]:
    """
    Return normalized heading titles that contain a [[INSERER]] placeholder under them.
    """
    heading_labels, _ = docx_inserer_placeholders(path)
    return set(heading_labels.keys())


def docx_inserer_placeholders(path: str) -> tuple[dict[str, str], list[str]]:
    """
    Return a map of normalized heading titles to their first [[INSERER]] label,
    plus a list of [[INSERER]] labels that are outside of any heading section.
    """
    doc = Document(path)
    last_heading: Optional[str] = None
    current_section_heading: Optional[str] = None
    heading_labels: dict[str, str] = {}
    orphan_labels: list[str] = []
    seen_orphans: set[str] = set()

    def register_labels(labels: list[str]) -> None:
        nonlocal heading_labels, orphan_labels, seen_orphans, current_section_heading
        if current_section_heading:
            key = _normalize_title(current_section_heading)
            if key and key not in heading_labels and labels:
                heading_labels[key] = labels[0]
            return
        for label in labels:
            if label not in seen_orphans:
                seen_orphans.add(label)
                orphan_labels.append(label)

    for item in _iter_block_items(doc):
        if isinstance(item, Paragraph):
            txt = (item.text or "").strip()
            if txt:
                lvl = _get_heading_level(item)
                if lvl is not None:
                    last_heading = txt
                    current_section_heading = None
                    continue
                if SEC_TAG_RE.search(txt):
                    if last_heading:
                        current_section_heading = last_heading
                labels = _extract_inserer_labels(txt)
                if labels:
                    register_labels(labels)
            continue

        for row in item.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = (p.text or "").strip()
                    if not txt:
                        continue
                    if SEC_TAG_RE.search(txt):
                        if last_heading:
                            current_section_heading = last_heading
                    labels = _extract_inserer_labels(txt)
                    if labels:
                        register_labels(labels)

    return heading_labels, orphan_labels


def docx_content_blocks(path: str) -> list[dict[str, str | int | None]]:
    """
    Return ordered content blocks (sections + insert placeholders) based on DOCX order.
    """
    doc = Document(path)
    blocks: list[dict[str, str | int | None]] = []
    last_heading: Optional[str] = None
    current_section_heading: Optional[str] = None
    order_index = 0

    def add_section(sec_key: str) -> None:
        nonlocal order_index
        blocks.append(
            {
                "type": "section",
                "sec_key": sec_key,
                "order_index": order_index,
            }
        )
        order_index += 1

    def add_insert(labels: list[str]) -> None:
        nonlocal order_index
        for label in labels:
            insert_key = (
                _normalize_title(current_section_heading)
                if current_section_heading
                else _normalize_title(label)
            )
            blocks.append(
                {
                    "type": "insert",
                    "order_index": order_index,
                    "insert_label": label,
                    "insert_key": insert_key,
                    "section_key": _normalize_title(current_section_heading) if current_section_heading else None,
                }
            )
            order_index += 1

    for item in _iter_block_items(doc):
        if isinstance(item, Paragraph):
            txt = (item.text or "").strip()
            if not txt:
                continue
            lvl = _get_heading_level(item)
            if lvl is not None:
                last_heading = txt
                current_section_heading = None
                continue
            sec_match = SEC_TAG_RE.search(txt)
            if sec_match:
                if last_heading:
                    current_section_heading = last_heading
                add_section(sec_match.group(1))
            labels = _extract_inserer_labels(txt)
            if labels:
                add_insert(labels)
            continue

        for row in item.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = (p.text or "").strip()
                    if not txt:
                        continue
                    sec_match = SEC_TAG_RE.search(txt)
                    if sec_match:
                        if last_heading:
                            current_section_heading = last_heading
                        add_section(sec_match.group(1))
                    labels = _extract_inserer_labels(txt)
                    if labels:
                        add_insert(labels)

    return blocks


def _scan_general_placeholders(text: str) -> list[dict[str, str | int]]:
    found: list[dict[str, str | int]] = []
    if not text:
        return found

    for match in TEXT_PLACEHOLDER_RE.finditer(text):
        raw = match.group(0)
        description = (match.group(1) or "").strip()
        if not description:
            continue
        found.append(
            {
                "kind": "text",
                "start": match.start(),
                "raw": raw,
                "description": description,
            }
        )

    for match in LOGO_PLACEHOLDER_RE.finditer(text):
        found.append(
            {
                "kind": "logo",
                "start": match.start(),
                "raw": match.group(0),
                "description": "",
            }
        )

    found.sort(key=lambda item: int(item["start"]))
    return found


def docx_general_placeholders(path: str) -> dict[str, object]:
    """
    Extract placeholders visible by USER flow:
      - [[TEXT: <description>]]
      - [[LOGO]]
    """
    doc = Document(path)
    text_placeholders: list[dict[str, str]] = []
    has_logo = False
    text_idx = 1

    def register_text(raw: str, description: str) -> None:
        nonlocal text_idx
        text_placeholders.append(
            {
                "id": f"text_{text_idx}",
                "description": description,
                "raw": raw,
            }
        )
        text_idx += 1

    for item in _iter_block_items(doc):
        if isinstance(item, Paragraph):
            found = _scan_general_placeholders(item.text or "")
            for entry in found:
                if entry["kind"] == "text":
                    register_text(str(entry["raw"]), str(entry["description"]))
                elif entry["kind"] == "logo":
                    has_logo = True
            continue

        for row in item.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    found = _scan_general_placeholders(p.text or "")
                    for entry in found:
                        if entry["kind"] == "text":
                            register_text(str(entry["raw"]), str(entry["description"]))
                        elif entry["kind"] == "logo":
                            has_logo = True

    return {
        "text_placeholders": text_placeholders,
        "has_logo": has_logo,
    }
