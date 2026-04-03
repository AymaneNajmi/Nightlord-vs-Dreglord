from __future__ import annotations

from io import BytesIO
from typing import List, Optional, Tuple

import openpyxl
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def _iter_block_items(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def _delete_element(el):
    el.getparent().remove(el)


def _heading_level(p: Paragraph) -> Optional[int]:
    """
    Detect heading level from style name or style_id.
    Works with 'Heading 1'/'Heading 2' and often 'Titre 1' etc.
    """
    if not p.style:
        return None
    name = (p.style.name or "").lower()
    sid = (p.style.style_id or "").lower()

    for n in range(1, 7):
        if f"heading {n}" in name or f"titre {n}" in name or sid in (f"heading{n}", f"titre{n}"):
            return n
    return None


def remove_sections_by_heading_titles(doc: Document, titles_to_remove: List[str]):
    """
    Remove section blocks starting at a heading paragraph whose text == one of titles_to_remove
    until next heading of same or higher level.
    """
    titles_norm = {t.strip().lower() for t in titles_to_remove if t and t.strip()}
    if not titles_norm:
        return

    blocks = list(_iter_block_items(doc))
    i = 0

    while i < len(blocks):
        b = blocks[i]
        if isinstance(b, Paragraph):
            lvl = _heading_level(b)
            if lvl is not None:
                t = (b.text or "").strip().lower()
                if t in titles_norm:
                    # delete from i until next heading <= lvl
                    j = i
                    while j < len(blocks):
                        bj = blocks[j]
                        if isinstance(bj, Paragraph):
                            lvl2 = _heading_level(bj)
                            if j != i and lvl2 is not None and lvl2 <= lvl:
                                break
                        j += 1

                    for k in range(i, j):
                        _delete_element(blocks[k]._element)

                    blocks = list(_iter_block_items(doc))
                    continue
        i += 1


def replace_excel_placeholders_with_tables(
    doc: Document,
    workbook_bytes: bytes,
    max_rows: int = 60,
    max_cols: int = 14,
):
    """
    Replace [[EXCEL:SheetName]] placeholders by inserting a Word table.
    Keeps text before/after placeholder if present in same paragraph.
    """
    wb = openpyxl.load_workbook(BytesIO(workbook_bytes), data_only=True)
    sheet_map = {ws.title: ws for ws in wb.worksheets}

    def sheet_to_matrix(sheet_name: str):
        ws = sheet_map.get(sheet_name)
        if not ws:
            return None
        rows = []
        for r in ws.iter_rows(values_only=True):
            rows.append([("" if v is None else str(v)) for v in r])

        # trim empty trailing rows
        while rows and all(c == "" for c in rows[-1]):
            rows.pop()

        if not rows:
            return []

        # limit
        rows = rows[:max_rows]
        rows = [row[:max_cols] for row in rows]
        return rows

    # iterate paragraphs snapshot (we will modify doc)
    for p in list(doc.paragraphs):
        txt = p.text or ""
        if "[[EXCEL:" not in txt:
            continue

        start = txt.find("[[EXCEL:")
        end = txt.find("]]", start)
        if end == -1:
            continue

        placeholder = txt[start : end + 2]
        sheet_name = placeholder[len("[[EXCEL:") : -2].strip()

        matrix = sheet_to_matrix(sheet_name)
        if matrix is None:
            # sheet not found -> keep as is
            continue

        before = (txt[:start] or "").strip()
        after = (txt[end + 2 :] or "").strip()

        # replace paragraph text with before only
        p.text = before

        anchor_el = p._element

        # insert table after paragraph
        if matrix:
            tbl = doc.add_table(rows=len(matrix), cols=len(matrix[0]))
            for i, row in enumerate(matrix):
                for j, val in enumerate(row):
                    tbl.cell(i, j).text = val
            tbl_el = tbl._element
            anchor_el.addnext(tbl_el)
            anchor_el = tbl_el

        # insert "after" paragraph if needed
        if after:
            newp = doc.add_paragraph(after)
            anchor_el.addnext(newp._element)
