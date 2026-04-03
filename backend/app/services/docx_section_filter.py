import re
from docx import Document

HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Titre 1", "Titre 2", "Titre 3", "Titre 4"}
SEC_TAG_RE = re.compile(r"\[\[\s*(SEC_[A-Za-z0-9_]+)\s*\]\]")
ASTERISK_TITLE_RE = re.compile(r"\*\s*$")


def _has_trailing_asterisk(title: str) -> bool:
    return bool(ASTERISK_TITLE_RE.search((title or "").strip()))


def _asterisk_titles_without_sections(doc: Document) -> set[str]:
    targets: set[str] = set()
    current_title = None
    current_has_asterisk = False
    current_has_section = False

    for p in doc.paragraphs:
        style_name = getattr(p.style, "name", "") or ""
        is_heading = style_name in HEADING_STYLES
        text = (p.text or "").strip()

        if is_heading:
            if current_title and current_has_asterisk and not current_has_section:
                targets.add(current_title)
            current_title = text
            current_has_asterisk = _has_trailing_asterisk(text)
            current_has_section = False
            continue

        if current_has_asterisk and SEC_TAG_RE.search(text):
            current_has_section = True

    if current_title and current_has_asterisk and not current_has_section:
        targets.add(current_title)

    return targets


def remove_sections_by_titles(docx_in_path: str, docx_out_path: str, titles_to_remove: list[str]) -> None:
    titles_to_remove = [t.strip() for t in titles_to_remove if t and t.strip()]

    doc = Document(docx_in_path)
    auto_titles = _asterisk_titles_without_sections(doc)
    if auto_titles:
        titles_to_remove = titles_to_remove + [t for t in auto_titles if t not in titles_to_remove]

    if not titles_to_remove:
        doc.save(docx_out_path)
        return

    keep = []
    remove_mode = False

    for p in doc.paragraphs:
        style_name = getattr(p.style, "name", "") or ""
        is_heading = style_name in HEADING_STYLES

        if is_heading:
            current_title = (p.text or "").strip()
            remove_mode = current_title in titles_to_remove

        if not remove_mode:
            keep.append(p._p)

    # rebuild document by clearing and re-inserting kept paragraphs XML
    new_doc = Document()
    body = new_doc._body._body
    for p_elm in keep:
        body.append(p_elm)

    new_doc.save(docx_out_path)
