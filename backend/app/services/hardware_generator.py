from __future__ import annotations

import json
import logging
import re
import time
from io import BytesIO
from typing import Any, Dict, List, Tuple
from urllib.parse import quote, urljoin

import requests
from bs4 import BeautifulSoup
from docx import Document
from openai import OpenAI
from pydantic import ValidationError
from pypdf import PdfReader

from app.schemas.hardware import HardwareOutput

logger = logging.getLogger(__name__)

UA = {"User-Agent": "Mozilla/5.0"}

FETCH_TIMEOUT_SECONDS = 35
SEARCH_TIMEOUT_SECONDS = 25
OPENAI_TIMEOUT_SECONDS = 45

CACHE_TTL_SECONDS = 24 * 60 * 60
_cache: Dict[str, Tuple[float, Dict[str, Any]]] = {}

SCHEMA_JSON = r"""
{
  "description_generale": "",
  "specs_techniques": [{ "feature": "", "valeur": "" }],
  "specs_physiques": [{ "spec": "", "valeur": "" }],
  "fonctionnalites": [],
  "aspect_fonctionnel": "",
  "licensing": {
    "network_essentials": "",
    "network_advantage": "",
    "dna_essentials": "",
    "dna_advantage": "",
    "delivery_model": ""
  },
  "uplink_modules": [{ "module_id": "", "description": "" }],
  "power_supplies": [{
    "model": "", "wattage": "", "btu_hr": "", "input_voltage": "", "input_current": "",
    "output_ratings": "", "hold_up_time": "", "input_receptacles": "", "cord_rating": "",
    "dimensions": "", "weight": "", "operating_temp": "", "storage_temp": "", "humidity": "",
    "altitude": "", "led_indicators": ""
  }],
  "stackwise_info": {
    "technology": "",
    "stackpower_supported": "",
    "max_members": "",
    "bandwidth": "",
    "compatibility": "",
    "restrictions": ""
  },
  "performance_scalability": [{ "metric": "", "value": "" }],
  "datasheet_url": "",
  "image_url": ""
}
"""

SYSTEM_PROMPT = (
    "Tu es un ingénieur réseau expert chargé de produire une description matérielle en français basée "
    "STRICTEMENT sur la documentation officielle Cisco.\n\n"
    "RÈGLES :\n"
    "- Tu utilises UNIQUEMENT les informations fournies ci-dessous (extrait officiel Cisco).\n"
    "- ZÉRO invention : si une info n'est pas trouvée dans l'extrait, mets \"\" ou [].\n"
    "- Tu peux reformuler (phrases professionnelles) UNIQUEMENT à partir des infos présentes.\n"
    "- Retourne UNIQUEMENT un JSON valide (pas de markdown, pas de texte).\n\n"
    "Référence matériel : {ref}\n"
    "Datasheet officielle (URL) : {datasheet_url}\n\n"
    "EXTRAIT OFFICIEL (texte de la datasheet/page Cisco) :\n"
    "{content}\n\n"
    "JSON ATTENDU (clés exactes) :\n"
    "{schema_json}"
)


class HardwareGenerationError(RuntimeError):
    def __init__(self, message: str, status_code: int = 400):
        super().__init__(message)
        self.message = message
        self.status_code = status_code


def _cache_key(hardware_ref: str, datasheet_url: str) -> str:
    return f"{hardware_ref.strip().upper()}::{datasheet_url.strip()}"


def get_cached_output(hardware_ref: str, datasheet_url: str) -> Dict[str, Any] | None:
    key = _cache_key(hardware_ref, datasheet_url)
    cached = _cache.get(key)
    if not cached:
        return None
    expires_at, payload = cached
    if time.time() > expires_at:
        _cache.pop(key, None)
        return None
    return payload


def set_cached_output(hardware_ref: str, datasheet_url: str, payload: Dict[str, Any]) -> None:
    key = _cache_key(hardware_ref, datasheet_url)
    _cache[key] = (time.time() + CACHE_TTL_SECONDS, payload)


def cisco_search_urls(query: str, max_results: int = 10) -> List[str]:
    url = f"https://www.cisco.com/c/en/us/search.html?q={quote(query)}"
    try:
        response = requests.get(url, headers=UA, timeout=SEARCH_TIMEOUT_SECONDS)
        response.raise_for_status()
    except requests.RequestException as exc:
        raise HardwareGenerationError(f"Erreur Cisco search: {exc}", status_code=502) from exc

    soup = BeautifulSoup(response.text, "html.parser")

    links: List[str] = []
    for anchor in soup.find_all("a", href=True):
        href = anchor["href"].strip()
        if href.startswith("/"):
            href = urljoin("https://www.cisco.com", href)
        if "cisco.com" not in href:
            continue
        if (
            "/products/" in href
            or "/collateral/" in href
            or "data-sheet" in href
            or "datasheet" in href
        ):
            links.append(href)

    out: List[str] = []
    for item in links:
        if item not in out:
            out.append(item)

    return out[:max_results]


def pick_datasheet_url(urls: List[str]) -> str:
    if not urls:
        return ""

    for url in urls:
        lower = url.lower()
        if lower.endswith(".pdf") and (
            "collateral" in lower or "data-sheet" in lower or "datasheet" in lower
        ):
            return url

    for url in urls:
        lower = url.lower()
        if "collateral" in lower or "data-sheet" in lower or "datasheet" in lower:
            return url

    return urls[0]


def fetch_pdf_text(pdf_url: str, max_pages: int = 12, max_chars: int = 45000) -> str:
    try:
        response = requests.get(pdf_url, headers=UA, timeout=FETCH_TIMEOUT_SECONDS)
        response.raise_for_status()
    except requests.RequestException as exc:
        raise HardwareGenerationError(f"Erreur téléchargement PDF: {exc}", status_code=502) from exc

    reader = PdfReader(BytesIO(response.content))
    text = ""
    for page in reader.pages[:max_pages]:
        text += "\n" + (page.extract_text() or "")
        if len(text) >= max_chars:
            break

    return text[:max_chars]


def fetch_html_text(url: str, max_chars: int = 45000) -> str:
    try:
        response = requests.get(url, headers=UA, timeout=FETCH_TIMEOUT_SECONDS)
        response.raise_for_status()
    except requests.RequestException as exc:
        raise HardwareGenerationError(f"Erreur téléchargement HTML: {exc}", status_code=502) from exc

    soup = BeautifulSoup(response.text, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    text = soup.get_text("\n")
    text = re.sub(r"\n{2,}", "\n", text).strip()
    return text[:max_chars]


def find_official_image_url(html_text: str) -> str:
    match = re.search(
        r"https?://www\.cisco\.com/[^ \n\"']+/c/dam/[^ \n\"']+\.(png|jpg|jpeg|webp)",
        html_text,
        re.I,
    )
    return match.group(0) if match else ""


def url_exists(url: str) -> bool:
    try:
        response = requests.get(url, headers=UA, timeout=15, allow_redirects=True)
        return response.status_code == 200
    except requests.RequestException:
        return False


def guess_c9200_datasheet_urls(hardware_ref: str) -> List[str]:
    _ = hardware_ref.upper().strip()
    candidates = [
        "https://www.cisco.com/c/en/us/products/collateral/switches/catalyst-9200-series-switches/nb-06-cat9200-ser-data-sheet-cte-en.html",
        "https://www.cisco.com/c/en/us/products/collateral/switches/catalyst-9200-series-switches/nb-06-cat9200-ser-data-sheet-cte-en.pdf",
        "https://www.cisco.com/c/en/us/products/switches/catalyst-9200-series-switches/index.html",
        "https://www.cisco.com/c/en/us/products/switches/catalyst-9200l-series-switches/index.html",
    ]
    return candidates


def find_official_cisco_url(hardware_ref: str) -> str:
    candidates = guess_c9200_datasheet_urls(hardware_ref)
    for candidate in candidates:
        if url_exists(candidate):
            return candidate
    return ""


def _extract_json(text: str) -> Dict[str, Any]:
    text = text or ""
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1:
        raise ValueError(f"Aucun JSON trouvé. Début réponse: {text[:300]}")
    return json.loads(text[start : end + 1])


def _clean_empty(data: Any) -> Any:
    if isinstance(data, str) and data.strip().lower() == "empty":
        return ""
    if isinstance(data, list):
        return [_clean_empty(item) for item in data]
    if isinstance(data, dict):
        return {key: _clean_empty(value) for key, value in data.items()}
    return data


def _search_datasheet_url(hardware_ref: str) -> List[str]:
    queries = [
        f"{hardware_ref} datasheet pdf",
        f"{hardware_ref} data sheet",
        f"{hardware_ref} collateral",
        f"{hardware_ref} catalyst",
    ]
    for query in queries:
        try:
            urls = cisco_search_urls(query, max_results=12)
        except HardwareGenerationError:
            continue
        if urls:
            return urls
    return []


def generate_hardware_json(hardware_ref: str) -> HardwareOutput:
    datasheet_url = find_official_cisco_url(hardware_ref)
    if not datasheet_url:
        urls = _search_datasheet_url(hardware_ref)
        datasheet_url = pick_datasheet_url(urls)

    if not datasheet_url:
        raise HardwareGenerationError(
            f"Aucune page Cisco officielle trouvée pour {hardware_ref}. Famille non supportée.",
            status_code=404,
        )

    cached = get_cached_output(hardware_ref, datasheet_url)
    if cached:
        try:
            return HardwareOutput.model_validate(cached)
        except ValidationError:
            _cache.pop(_cache_key(hardware_ref, datasheet_url), None)

    if datasheet_url.lower().endswith(".pdf"):
        content = fetch_pdf_text(datasheet_url)
        image_url = ""
        source_type = "pdf"
    else:
        content = fetch_html_text(datasheet_url)
        image_url = find_official_image_url(content)
        source_type = "html"

    if not content or len(content.strip()) < 200:
        raise HardwareGenerationError(
            f"Contenu Cisco vide ou bloqué. URL={datasheet_url}",
            status_code=422,
        )

    logger.info(
        "hardware_generation hardware_ref=%s datasheet_url=%s source=%s len=%s",
        hardware_ref,
        datasheet_url,
        source_type,
        len(content),
    )

    client = OpenAI(timeout=OPENAI_TIMEOUT_SECONDS)
    prompt = SYSTEM_PROMPT.format(
        ref=hardware_ref,
        datasheet_url=datasheet_url,
        content=content,
        schema_json=SCHEMA_JSON,
    )

    try:
        response = client.responses.create(
            model="gpt-4o",
            input=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": "Produis le JSON maintenant."},
            ],
            temperature=0.1,
        )
    except Exception as exc:  # pragma: no cover - network dependent
        raise HardwareGenerationError(f"Erreur OpenAI: {exc}", status_code=502) from exc

    raw = _extract_json(response.output_text)
    raw = _clean_empty(raw)
    raw["datasheet_url"] = datasheet_url
    if image_url:
        raw["image_url"] = image_url

    try:
        output = HardwareOutput.model_validate(raw)
    except ValidationError as exc:
        raise HardwareGenerationError(f"JSON OpenAI incomplet: {exc}", status_code=502) from exc

    set_cached_output(hardware_ref, datasheet_url, output.model_dump())
    return output


def hardware_output_to_summary_text(hardware_ref: str, data: HardwareOutput | Dict[str, Any]) -> str:
    payload = data.model_dump() if hasattr(data, "model_dump") else dict(data or {})
    lines: list[str] = []

    if hardware_ref:
        lines.append(f"Référence: {hardware_ref}")

    description = (payload.get("description_generale") or "").strip()
    if description:
        lines.append(description)

    features = payload.get("fonctionnalites") or []
    if isinstance(features, list) and features:
        lines.append("Fonctionnalités:")
        for feat in features[:8]:
            value = str(feat or "").strip()
            if value:
                lines.append(f"- {value}")

    perf_rows = payload.get("performance_scalability") or []
    if isinstance(perf_rows, list) and perf_rows:
        lines.append("Performance / Scalabilité:")
        for row in perf_rows[:8]:
            if not isinstance(row, dict):
                continue
            metric = str(row.get("metric") or "").strip()
            value = str(row.get("value") or "").strip()
            if metric and value:
                lines.append(f"- {metric}: {value}")
            elif metric:
                lines.append(f"- {metric}")

    aspect = (payload.get("aspect_fonctionnel") or "").strip()
    if aspect:
        lines.append("Aspect fonctionnel:")
        lines.append(aspect)

    datasheet_url = (payload.get("datasheet_url") or "").strip()
    if datasheet_url:
        lines.append(f"Source: {datasheet_url}")

    return "\n".join(lines).strip() or (f"Hardware Cisco ({hardware_ref})" if hardware_ref else "Hardware Cisco")


def hardware_output_to_summary_html(hardware_ref: str, data: HardwareOutput | Dict[str, Any]) -> str:
    text = hardware_output_to_summary_text(hardware_ref=hardware_ref, data=data)
    paragraphs = [line.strip() for line in text.splitlines() if line.strip()]
    if not paragraphs:
        return ""
    return "".join(f"<p>{line}</p>" for line in paragraphs)


def hardware_output_to_bom_table(data: HardwareOutput | Dict[str, Any]) -> Dict[str, Any]:
    payload = data.model_dump() if hasattr(data, "model_dump") else dict(data or {})
    rows: list[dict[str, str]] = []
    for item in payload.get("uplink_modules") or []:
        if not isinstance(item, dict):
            continue
        rows.append(
            {
                "module_id": str(item.get("module_id") or "").strip(),
                "description": str(item.get("description") or "").strip(),
            }
        )
    return {
        "columns": ["module_id", "description"],
        "rows": rows,
    }


def generate_hardware_content(hardware_ref: str) -> Dict[str, Any]:
    output = generate_hardware_json(hardware_ref)
    output_json = output.model_dump()
    return {
        "output_json": output_json,
        "formatted_summary_text": hardware_output_to_summary_text(hardware_ref, output_json),
        "formatted_summary_html": hardware_output_to_summary_html(hardware_ref, output_json),
        "bom_table": hardware_output_to_bom_table(output_json),
    }


def hardware_to_docx_bytes(hardware_ref: str, data: HardwareOutput) -> bytes:
    if hasattr(data, "model_dump"):
        data = data.model_dump()

    doc = Document()

    title = doc.add_heading("7.1.4 Aspect HARDWARE des équipements", level=1)
    title.runs[0].bold = True

    paragraph = doc.add_paragraph()
    paragraph.add_run(f"Référence : {hardware_ref}").bold = True

    doc.add_page_break()

    doc.add_heading("7.1.4.1 Description générale", level=2)
    doc.add_paragraph(data.get("description_generale", ""))

    doc.add_heading("7.1.4.2 Spécifications techniques", level=2)
    tech = data.get("specs_techniques", [])
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Paramètre"
    table.rows[0].cells[1].text = "Valeur"
    for item in tech:
        row = table.add_row().cells
        row[0].text = item.get("feature", "")
        row[1].text = "" if item.get("valeur", "").lower() == "empty" else item.get("valeur", "")

    doc.add_heading("7.1.4.3 Spécifications physiques", level=2)
    phys = data.get("specs_physiques", [])
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Spécification"
    table.rows[0].cells[1].text = "Valeur"
    for item in phys:
        row = table.add_row().cells
        row[0].text = item.get("spec", "")
        row[1].text = "" if item.get("valeur", "").lower() == "empty" else item.get("valeur", "")

    doc.add_heading("7.1.4.4 Fonctionnalités", level=2)
    for feature in data.get("fonctionnalites", []):
        if feature and feature.lower() != "empty":
            doc.add_paragraph(feature, style="List Bullet")

    doc.add_heading("7.1.4.5 Aspect fonctionnel", level=2)
    doc.add_paragraph(data.get("aspect_fonctionnel", ""))

    doc.add_page_break()
    doc.add_heading("7.1.4.6 Licensing", level=2)
    lic = data.get("licensing", {})
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Licence"
    table.rows[0].cells[1].text = "Description"
    for key, value in lic.items():
        row = table.add_row().cells
        row[0].text = key.replace("_", " ").title()
        row[1].text = "" if str(value).lower() == "empty" else str(value)

    doc.add_heading("7.1.4.7 Modules uplink", level=2)
    uplinks = data.get("uplink_modules", [])
    if uplinks:
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Module"
        table.rows[0].cells[1].text = "Description"
        for uplink in uplinks:
            row = table.add_row().cells
            row[0].text = uplink.get("module_id", "")
            row[1].text = uplink.get("description", "")
    else:
        doc.add_paragraph("Aucun module uplink spécifique référencé dans la datasheet officielle Cisco.")

    doc.add_heading("7.1.4.8 Alimentation (Power Supplies)", level=2)
    psus = data.get("power_supplies", [])
    for psu in psus:
        doc.add_paragraph(f"Modèle : {psu.get('model', '')}", style="List Bullet")
        doc.add_paragraph(f"Puissance : {psu.get('wattage', '')}", style="List Bullet")

    doc.add_heading("7.1.4.9 StackWise / StackPower", level=2)
    stack = data.get("stackwise_info", {})
    for key, value in stack.items():
        if value and value.lower() != "empty":
            doc.add_paragraph(f"{key.replace('_', ' ').title()} : {value}")

    doc.add_heading("7.1.4.10 Performance & Scalabilité", level=2)
    perf = data.get("performance_scalability", [])
    if perf:
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Métrique"
        table.rows[0].cells[1].text = "Valeur"
        for metric in perf:
            row = table.add_row().cells
            row[0].text = metric.get("metric", "")
            row[1].text = metric.get("value", "")

    doc.add_page_break()
    doc.add_heading("7.1.5 Références", level=1)
    doc.add_paragraph(f"Datasheet Cisco : {data.get('datasheet_url', '')}")
    image_url = data.get("image_url", "")
    doc.add_paragraph(f"Image officielle : {'' if image_url.lower() == 'empty' else image_url}")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
