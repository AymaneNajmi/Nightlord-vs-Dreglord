from __future__ import annotations
from dataclasses import dataclass
from typing import List, Tuple, Iterable, Set
import re
from docx import Document

HEADING_STYLE_PREFIXES = (
    "Heading",   # EN: Heading 1
    "Titre",     # FR: Titre 1
    "Überschrift",  # DE sometimes
)

SEC_RE = re.compile(r"\[\[\s*(SEC_[A-Za-z0-9_]+)\s*\]\]")

@dataclass
class HeadingItem:
    level: int
    text: str

def _iter_paragraphs(doc: Document) -> Iterable:
    # body
    for p in doc.paragraphs:
        yield p
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def _is_heading_style(style_name: str) -> Tuple[bool, int]:
    if not style_name:
        return (False, 0)
    s = style_name.strip()

    # Examples: "Heading 1", "Heading 2", "Titre 1", "Titre 2"
    for pref in HEADING_STYLE_PREFIXES:
        if s.startswith(pref):
            # extract digits at end
            digits = "".join(ch for ch in s if ch.isdigit())
            if digits:
                lvl = int(digits)
                return (True, max(1, min(lvl, 6)))
            return (True, 2)  # fallback
    return (False, 0)

def _looks_like_title(p) -> bool:
    # fallback heuristic when template uses bold+big text instead of heading styles
    txt = (p.text or "").strip()
    if len(txt) < 4:
        return False
    if len(txt) > 140:
        return False

    # if fully uppercase or ends with ":" or starts with numbering like 3.1.4
    starts_num = txt[0].isdigit()
    ends_colon = txt.endswith(":")
    upper_ratio = sum(c.isupper() for c in txt) / max(1, sum(c.isalpha() for c in txt))
    mostly_upper = upper_ratio > 0.6 if any(c.isalpha() for c in txt) else False

    # bold ratio
    runs = list(p.runs)
    if runs:
        bold_runs = sum(1 for r in runs if r.bold)
        bold_ratio = bold_runs / len(runs)
    else:
        bold_ratio = 0

    # font size heuristic (points)
    sizes = [r.font.size.pt for r in runs if r.font.size]
    avg_size = sum(sizes)/len(sizes) if sizes else 0

    return (starts_num or ends_colon or mostly_upper) and (bold_ratio >= 0.5 or avg_size >= 12)

def extract_headings(docx_path: str) -> List[HeadingItem]:
    doc = Document(docx_path)
    out: List[HeadingItem] = []
    seen: Set[str] = set()

    for p in _iter_paragraphs(doc):
        txt = (p.text or "").strip()
        if not txt:
            continue

        style_name = getattr(getattr(p, "style", None), "name", "") or ""
        is_head, lvl = _is_heading_style(style_name)

        if is_head:
            key = f"{lvl}|{txt}"
            if key not in seen:
                out.append(HeadingItem(level=lvl, text=txt))
                seen.add(key)
            continue

        # fallback (optional, but helps a lot in real templates)
        if _looks_like_title(p):
            lvl = 3
            key = f"{lvl}|{txt}"
            if key not in seen:
                out.append(HeadingItem(level=lvl, text=txt))
                seen.add(key)

    return out


def _get_heading_level(style_name: str) -> int | None:
    # "Heading 1", "Heading 2", "Heading 3" (EN)
    m = re.match(r"Heading\s+(\d+)", style_name or "")
    if m:
        return int(m.group(1))
    # "Titre 1", "Titre 2" (FR)
    m = re.match(r"Titre\s+(\d+)", style_name or "")
    if m:
        return int(m.group(1))
    return None

def extract_sections_from_docx(path: str):
    """
    Retourne une liste ordonnée:
    [{sec_key, level, title, order_index}, ...]
    Logique:
      - on lit les paragraphes
      - quand on voit un Heading, on garde "current heading"
      - quand on voit [[SEC_x]], on l'associe au dernier heading vu
    """
    doc = Document(path)
    sections = []
    last_heading = None  # (level, title)

    order = 0
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue

        level = _get_heading_level(getattr(p.style, "name", "") if p.style else "")
        if level is not None and text:
            last_heading = (level, text)
            continue

        m = SEC_RE.search(text)
        if m and last_heading:
            sec_key = m.group(1)
            h_level, h_title = last_heading
            sections.append({
                "sec_key": sec_key,
                "level": h_level,
                "title": h_title,
                "order_index": order
            })
            order += 1

    return sections


def extract_asterisk_headings_without_sections(path: str):
    doc = Document(path)
    out = []
    current_title = None
    current_level = None
    current_has_asterisk = False
    current_has_section = False

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue

        level = _get_heading_level(getattr(p.style, "name", "") if p.style else "")
        if level is not None and text:
            if current_title and current_has_asterisk and not current_has_section:
                out.append({"title": current_title, "level": current_level})
            current_title = text
            current_level = level
            current_has_asterisk = text.strip().endswith("*")
            current_has_section = False
            continue

        if current_has_asterisk and SEC_RE.search(text):
            current_has_section = True

    if current_title and current_has_asterisk and not current_has_section:
        out.append({"title": current_title, "level": current_level})

    return out
