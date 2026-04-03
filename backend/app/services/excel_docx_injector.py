import re
from typing import List, Optional

from docx import Document
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook

TOKEN_RE = re.compile(r"\[\[EXCEL:([^\]]+)\]\]")  # [[EXCEL:SheetName]]


def _iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _clear_paragraph(paragraph: Paragraph):
    for run in paragraph.runs:
        run.text = ""


def _sheet_to_matrix(xls_path: str, sheet_name: str, max_rows: int = 80, max_cols: int = 20) -> List[List[str]]:
    wb = load_workbook(xls_path, data_only=True)
    if sheet_name not in wb.sheetnames:
        raise ValueError(f"Feuille Excel introuvable: {sheet_name}")

    ws = wb[sheet_name]

    # range = used range
    max_r = min(ws.max_row or 1, max_rows)
    max_c = min(ws.max_column or 1, max_cols)

    matrix: List[List[str]] = []
    for r in range(1, max_r + 1):
        row_vals: List[str] = []
        empty_row = True
        for c in range(1, max_c + 1):
            v = ws.cell(row=r, column=c).value
            s = "" if v is None else str(v)
            if s.strip():
                empty_row = False
            row_vals.append(s)
        if not empty_row:
            matrix.append(row_vals)

    # trim empty columns at the end
    if not matrix:
        return [["(Feuille vide)"]]

    # remove trailing empty cols
    last_non_empty_col = 0
    for row in matrix:
        for i, val in enumerate(row):
            if str(val).strip():
                last_non_empty_col = max(last_non_empty_col, i)
    last_non_empty_col = max(last_non_empty_col, 0)

    trimmed = [row[: last_non_empty_col + 1] for row in matrix]
    return trimmed


def _insert_table_after_paragraph(doc: Document, paragraph: Paragraph, matrix: List[List[str]]):
    # python-docx: insert by creating table at end then moving it after paragraph
    rows = len(matrix)
    cols = max(len(r) for r in matrix) if matrix else 1

    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"

    for r in range(rows):
        for c in range(cols):
            val = matrix[r][c] if c < len(matrix[r]) else ""
            table.cell(r, c).text = str(val)

    # move table right after paragraph
    p_elm = paragraph._p
    tbl_elm = table._tbl
    p_elm.addnext(tbl_elm)


def inject_excel_tables(docx_in_path: str, docx_out_path: str, excel_path: str) -> None:
    doc = Document(docx_in_path)

    for p in list(_iter_paragraphs(doc)):
        text = p.text or ""
        m = TOKEN_RE.search(text)
        if not m:
            continue

        sheet_name = m.group(1).strip()

        # Remove token from paragraph text (even if token inside runs)
        # easiest: rebuild paragraph plain text without token
        new_text = TOKEN_RE.sub("", text).strip()

        _clear_paragraph(p)
        if new_text:
            p.add_run(new_text)

        matrix = _sheet_to_matrix(excel_path, sheet_name)
        _insert_table_after_paragraph(doc, p, matrix)

    doc.save(docx_out_path)
