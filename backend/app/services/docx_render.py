from __future__ import annotations
from pathlib import Path
from docx import Document

def _iter_paragraphs_and_cells(doc: Document):
    # paragraphs
    for p in doc.paragraphs:
        yield p
    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def _replace_tag_in_paragraph(p, tag: str, text: str):
    if tag not in (p.text or ""):
        return False

    # On remplace tout le paragraphe (POC), en gardant le style global du paragraphe.
    # Si tu veux préserver finement les runs, on peut upgrader après.
    before_style = p.style
    p.text = p.text.replace(tag, text)
    p.style = before_style
    return True

def render_docx_from_sections(template_path: str, sections_text: dict, out_path: str) -> str:
    tpl = Path(template_path)
    if not tpl.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(str(tpl))

    # sections_text = { "SEC_1": "...", ... }
    for sec_key, content in sections_text.items():
        tag = f"[[{sec_key}]]"
        # assure string
        content = (content or "").strip()
        if not content:
            content = "TBD"

        for p in _iter_paragraphs_and_cells(doc):
            _replace_tag_in_paragraph(p, tag, content)

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)
