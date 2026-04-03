from __future__ import annotations

import json
import logging
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Tuple

from docx import Document
from openai import OpenAI
from sqlalchemy.orm import Session

from app.models.ai_template_job import AITemplateJob, AITemplateJobStatus
from app.models.forms import FormOption, FormQuestion, FormSection, FormTemplate

logger = logging.getLogger(__name__)

ALLOWED_QTYPES = {"single_choice", "multi_choice"}
IDENTIFICATION_QTYPE = "text"
SYSTEM_PROMPT = (
    "Tu es un Architecte Réseau & Sécurité senior. "
    "Tu génères un FORMULAIRE (pas un texte LLD). "
    "SOURCE UNIQUE : le contexte EXACT fourni (extrait DOCX). "
    "Interdiction totale d’inventer. "
    "Intent/Example au niveau SECTION uniquement. "
    "Questions uniquement single_choice/multi_choice. "
    "Options uniquement issues du contexte (pas d’options génériques). "
    "Si contexte vide => questions=[]. "
    "Retourne uniquement un JSON conforme au schéma."
)


def _normalize(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").strip().lower())


def _is_option_grounded(normalized_context: str, option_value: str, source_quote: str) -> bool:
    n_value = _normalize(option_value)
    n_quote = _normalize(source_quote)
    if not n_value and not n_quote:
        return False

    # preferred path: explicit citation must exist in context
    if n_quote and n_quote in normalized_context:
        # allow exact match OR containment in either direction to tolerate punctuation/trim differences
        if n_value == n_quote or n_value in n_quote or n_quote in n_value:
            return True

    # fallback path: option value itself is explicitly present in context
    if n_value and n_value in normalized_context:
        return True

    return False


def _is_identification_allowed(context: str, label: str) -> bool:
    patterns = [
        r"\bnom du client\b",
        r"\bidentification\b",
        r"\bclient name\b",
        r"\bnom\b",
    ]
    c = _normalize(context)
    l = _normalize(label)
    return any(re.search(p, c) for p in patterns) and ("nom" in l or "client" in l or "ident" in l)


@dataclass
class OutlineSection:
    section_key: str
    title: str
    level: str
    paragraphs: List[str]


class QualityGateError(RuntimeError):
    def __init__(self, message: str, invalid_options: List[str] | None = None):
        super().__init__(message)
        self.invalid_options = invalid_options or []


def extract_outline_from_docx(docx_path: str) -> List[OutlineSection]:
    doc = Document(docx_path)
    sections: List[OutlineSection] = []
    current: OutlineSection | None = None
    h2_count = 0
    h3_count = 0

    h2_aliases = {"heading 2", "titre 2", "überschrift 2", "encabezado 2"}
    h3_aliases = {"heading 3", "titre 3", "überschrift 3", "encabezado 3"}

    def _heading_level(style_name: str) -> str | None:
        normalized = (style_name or "").strip().lower()
        if normalized in h2_aliases or re.search(r"\b(?:heading|titre|encabezado|überschrift)\s*2\b", normalized):
            return "H2"
        if normalized in h3_aliases or re.search(r"\b(?:heading|titre|encabezado|überschrift)\s*3\b", normalized):
            return "H3"
        return None

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        style_name = (para.style.name or "") if para.style else ""
        level = _heading_level(style_name)
        if level:
            if level == "H2":
                h2_count += 1
                h3_count = 0
                sec_key = f"SEC_{h2_count}"
            else:
                h3_count += 1
                sec_key = f"SEC_{h2_count}_{h3_count}"
            current = OutlineSection(section_key=sec_key, title=text or sec_key, level=level, paragraphs=[])
            sections.append(current)
            continue
        if style_name == "Heading 1":
            continue
        if current and text:
            current.paragraphs.append(text)

    logger.info("ai_form_builder_rich extracted headings: total=%s h2=%s h3=%s", len(sections), h2_count, h3_count)
    return sections


def build_section_context(section: OutlineSection, max_chars: int = 20000) -> str:
    body = "\n\n".join(section.paragraphs).strip()
    context = f"Titre: {section.title}\n\nTexte:\n{body}" if body else ""
    if len(context) > max_chars:
        context = context[:max_chars]
    logger.info("ai_form_builder_rich context_length section=%s length=%s", section.section_key, len(context))
    return context


def _schema() -> Dict[str, Any]:
    return {
        "type": "object",
        "additionalProperties": False,
        "required": ["section_key", "title", "level", "intent", "example", "questions"],
        "properties": {
            "section_key": {"type": "string"},
            "title": {"type": "string", "minLength": 1},
            "level": {"type": "string", "enum": ["H2", "H3"]},
            "intent": {"type": "string", "minLength": 3},
            "example": {"type": "string", "minLength": 3},
            "questions": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "required": ["order", "label", "qtype", "is_required", "options"],
                    "properties": {
                        "order": {"type": "integer", "minimum": 0},
                        "label": {"type": "string", "minLength": 3},
                        "qtype": {"type": "string", "enum": ["single_choice", "multi_choice", "text"]},
                        "is_required": {"type": "boolean"},
                        "options": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "additionalProperties": False,
                                "required": ["order", "value", "source_quote"],
                                "properties": {
                                    "order": {"type": "integer", "minimum": 0},
                                    "value": {"type": "string", "minLength": 1},
                                    "source_quote": {"type": "string", "minLength": 1},
                                },
                            },
                        },
                    },
                },
            },
        },
    }


def _extract_candidate_options(context: str, max_items: int = 60) -> List[str]:
    candidates: List[str] = []
    seen = set()
    raw_lines = [line.strip() for line in (context or "").splitlines() if line.strip()]
    for line in raw_lines:
        parts = re.split(r"[,;|/•\-]", line)
        for part in parts:
            token = re.sub(r"\s+", " ", part).strip(" .:\t")
            if len(token) < 2 or len(token) > 80:
                continue
            key = _normalize(token)
            if key and key not in seen:
                seen.add(key)
                candidates.append(token)
                if len(candidates) >= max_items:
                    return candidates
    return candidates


def _user_prompt(
    section: OutlineSection,
    context: str,
    continuation: bool = False,
    invalid_options: List[str] | None = None,
) -> str:
    candidate_options = _extract_candidate_options(context)
    candidate_options_text = "\n".join(f"- {item}" for item in candidate_options[:50]) or "- (aucune)"
    invalid_options = invalid_options or []
    invalid_options_text = "\n".join(f"- {item}" for item in invalid_options) if invalid_options else ""
    prompt = (
        "CONTEXTE (SOURCE UNIQUE)\n"
        f"Titre: {section.title}\n"
        f"Texte:\n{context}\n\n"
        "LISTE CANDIDATE D'OPTIONS EXTRAITES DU CONTEXTE\n"
        f"{candidate_options_text}\n\n"
        "TACHE\n"
        "Générer :\n"
        "- intent: 2 à 5 phrases (cadrage)\n"
        "- example: 1 phrase\n"
        "- questions: riches, structurantes, sans texte libre\n"
        "- qtype: single_choice ou multi_choice uniquement\n"
        "- options: 1 à 8 options, toutes présentes dans le texte (pas d’invention)\n"
        "- ordre 0..n\n\n"
        "IMPORTANT\n"
        "- Ne pas créer de questions text sauf si le texte mentionne explicitement un champ d’identification (ex: Nom du client).\n"
        "- Si la matière est limitée : réduire le nombre de questions plutôt que d’inventer.\n"
        "- Chaque option DOIT être un extrait exact du contexte (copie mot pour mot, pas de reformulation).\n"
        "- Pour chaque option, renseigner source_quote avec la citation exacte dans le contexte.\n"
        "- ONLY JSON, NO TEXT."
    )
    if invalid_options:
        prompt += (
            "\n\nCORRECTION OBLIGATOIRE\n"
            "Les options suivantes ont été rejetées car absentes du contexte:\n"
            f"{invalid_options_text}\n"
            "Regénère en remplaçant ces options uniquement par des extraits exacts du contexte."
        )
    if continuation:
        prompt += "\nContinue JSON from last valid token."
    return prompt


def call_openai_generate_section(
    section: OutlineSection,
    context: str,
    invalid_options: List[str] | None = None,
) -> Dict[str, Any]:
    client = OpenAI()
    model = os.getenv("OPENAI_MODEL_FORM", "gpt-4.1")
    attempts: List[Tuple[float, int, bool]] = [
        (0.1, 1700, False),
        (0.0, 2300, False),
        (0.0, 2600, True),
    ]
    last_error = "OpenAI generation failed"

    for idx, (temperature, max_tokens, continuation) in enumerate(attempts, start=1):
        try:
            response = client.responses.create(
                model=model,
                temperature=temperature,
                max_output_tokens=max_tokens,
                input=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {
                        "role": "user",
                        "content": _user_prompt(
                            section,
                            context,
                            continuation=continuation,
                            invalid_options=invalid_options,
                        ),
                    },
                ],
                text={
                    "format": {
                        "type": "json_schema",
                        "name": "rich_section",
                        "strict": True,
                        "schema": _schema(),
                    }
                },
            )
            payload_text = getattr(response, "output_text", None)
            if not payload_text:
                raise RuntimeError("Empty output_text")
            parsed = json.loads(payload_text)
            logger.info("ai_form_builder_rich openai_success section=%s attempt=%s", section.section_key, idx)
            return parsed
        except Exception as exc:  # noqa: BLE001
            last_error = str(exc)
            logger.exception(
                "ai_form_builder_rich openai_error section=%s attempt=%s error=%s",
                section.section_key,
                idx,
                exc,
            )
    raise RuntimeError(last_error)


def validate_section_output(
    section_data: Dict[str, Any],
    context: str,
    context_empty: bool = False,
    drop_invalid_questions: bool = False,
) -> Dict[str, Any]:
    intent = (section_data.get("intent") or "").strip()
    example = (section_data.get("example") or "").strip()
    questions = section_data.get("questions") or []

    if not intent or not example:
        raise RuntimeError("intent/example cannot be empty")
    if context_empty and questions:
        raise RuntimeError("questions must be empty for empty context")

    seen_labels = set()
    non_matching_options = 0
    total_options = 0
    normalized_context = _normalize(context)
    invalid_options: List[str] = []

    filtered_questions: List[Dict[str, Any]] = []

    for question in questions:
        label = (question.get("label") or "").strip()
        qtype = question.get("qtype")
        options = question.get("options") or []

        normalized_label = _normalize(label)
        if normalized_label in seen_labels:
            raise QualityGateError(f"duplicate label: {label}")
        seen_labels.add(normalized_label)

        if qtype not in ALLOWED_QTYPES:
            if not (qtype == IDENTIFICATION_QTYPE and _is_identification_allowed(context, label)):
                raise QualityGateError(f"unsupported qtype: {qtype}")

        if qtype in ALLOWED_QTYPES and len(options) < 1:
            if drop_invalid_questions:
                logger.warning(
                    "ai_form_builder_rich dropping_question_not_enough_options label=%s options=%s",
                    label,
                    len(options),
                )
                continue
            raise QualityGateError(f"not enough options for question: {label}")

        for option in options:
            total_options += 1
            raw_value = str(option.get("value") or "")
            source_quote = str(option.get("source_quote") or "")
            if not _is_option_grounded(normalized_context, raw_value, source_quote):
                invalid_options.append(raw_value)
                non_matching_options += 1

        filtered_questions.append(question)

    mismatch_ratio = (non_matching_options / total_options) if total_options else 0.0
    if mismatch_ratio > 0.30:
        raise QualityGateError(
            f"too many options not grounded in context ({non_matching_options}/{total_options})",
            invalid_options=invalid_options,
        )

    if drop_invalid_questions:
        section_data["questions"] = filtered_questions

    logger.info(
        "ai_form_builder_rich quality_gate section=%s questions=%s mismatching_options=%s",
        section_data.get("section_key"),
        len(questions),
        non_matching_options,
    )
    return section_data


def persist_to_db(
    db: Session,
    job: AITemplateJob,
    section_outputs: List[Dict[str, Any]],
) -> FormTemplate:
    form = db.query(FormTemplate).filter(FormTemplate.id == job.form_template_id).first() if job.form_template_id else None
    if not form:
        if not job.techno_id or not job.template_doc_id:
            raise RuntimeError("job is missing techno_id/template_doc_id")
        form = FormTemplate(
            techno_id=job.techno_id,
            doc_id=job.template_doc_id,
            name=f"{job.techno_name} (rich)",
            version=1,
            created_by="ai_form_builder_rich",
        )
        db.add(form)
        db.flush()

    db.query(FormSection).filter(FormSection.form_id == form.id).delete()
    db.flush()

    for sec_idx, section_data in enumerate(section_outputs):
        section = FormSection(
            form_id=form.id,
            sec_key=section_data["section_key"],
            heading_level=2 if section_data["level"] == "H2" else 3,
            heading_title=section_data["title"],
            order_index=sec_idx,
            section_intent=section_data["intent"],
            section_example=section_data["example"],
            status=section_data.get("status", "OK"),
            error_message=section_data.get("error_message"),
        )
        db.add(section)
        db.flush()

        for q in section_data.get("questions", []):
            question = FormQuestion(
                section_id=section.id,
                order_index=q["order"],
                label=q["label"],
                qtype=q["qtype"],
                is_required=bool(q.get("is_required", False)),
            )
            db.add(question)
            db.flush()
            for opt in q.get("options", []):
                db.add(
                    FormOption(
                        question_id=question.id,
                        order_index=opt["order"],
                        label=opt["value"],
                        value=opt["value"],
                    )
                )

    db.commit()
    db.refresh(form)
    return form


def run_generate_rich_job(db: Session, job: AITemplateJob) -> Dict[str, Any]:
    if not job.source_files:
        raise RuntimeError("job has no source files")
    source_docx = next((f.get("stored_path") for f in job.source_files if str(f.get("stored_path", "")).lower().endswith(".docx")), None)
    if not source_docx:
        raise RuntimeError("no DOCX found in source files")
    if not Path(source_docx).exists():
        raise RuntimeError("source DOCX not found on disk")

    outline = extract_outline_from_docx(source_docx)
    if not outline:
        raise RuntimeError("no H2/H3 sections found in DOCX")

    section_outputs: List[Dict[str, Any]] = []
    failed_sections: List[Dict[str, str]] = []
    job.status = AITemplateJobStatus.RUNNING
    job.error_message = None
    db.commit()

    for section in outline:
        context = build_section_context(section, max_chars=20000)
        if not context.strip():
            out = {
                "section_key": section.section_key,
                "title": section.title,
                "level": section.level,
                "intent": "Section sans contenu détaillé dans le document source.",
                "example": "Aucun exemple disponible.",
                "questions": [],
            }
            section_outputs.append(validate_section_output(out, context, context_empty=True))
            continue

        try:
            ai_output = call_openai_generate_section(section, context)
            ai_output.setdefault("section_key", section.section_key)
            ai_output.setdefault("title", section.title)
            ai_output.setdefault("level", section.level)
            try:
                valid = validate_section_output(ai_output, context)
            except QualityGateError as quality_exc:
                if not quality_exc.invalid_options:
                    raise
                logger.warning(
                    "ai_form_builder_rich quality_retry section=%s invalid_options=%s",
                    section.section_key,
                    quality_exc.invalid_options,
                )
                ai_output_retry = call_openai_generate_section(
                    section,
                    context,
                    invalid_options=quality_exc.invalid_options[:20],
                )
                ai_output_retry.setdefault("section_key", section.section_key)
                ai_output_retry.setdefault("title", section.title)
                ai_output_retry.setdefault("level", section.level)
                valid = validate_section_output(
                    ai_output_retry,
                    context,
                    drop_invalid_questions=True,
                )
            section_outputs.append(valid)
        except Exception as exc:  # noqa: BLE001
            logger.exception(
                "ai_form_builder_rich section_failed_continue section=%s error=%s",
                section.section_key,
                exc,
            )
            failed_sections.append({"section_key": section.section_key, "error": str(exc)})
            section_outputs.append(
                {
                    "section_key": section.section_key,
                    "title": section.title,
                    "level": section.level,
                    "intent": "Section en erreur pendant la génération.",
                    "example": "Aucun exemple disponible (erreur).",
                    "questions": [],
                    "status": "ERROR",
                    "error_message": str(exc),
                }
            )

    form = persist_to_db(db, job, section_outputs)
    job.form_template_id = form.id
    job.output_payload = {"sections": section_outputs, "failed_sections": failed_sections}

    if failed_sections:
        first = failed_sections[0]
        job.status = AITemplateJobStatus.FAILED
        job.error_message = (
            f"{len(failed_sections)} section(s) en erreur. "
            f"Première: {first['section_key']} -> {first['error']}"
        )
    else:
        job.status = AITemplateJobStatus.DONE
        job.error_message = None

    db.commit()

    return {
        "status": job.status.value if hasattr(job.status, "value") else str(job.status),
        "form_template_id": form.id,
        "failed_sections": failed_sections,
    }
