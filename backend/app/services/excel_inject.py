from __future__ import annotations
from typing import Optional
from docx import Document
from openpyxl import load_workbook


def _iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def inject_excel_tables(docx_in_path: str, docx_out_path: str, excel_path: str) -> None:
    doc = Document(docx_in_path)
    wb = load_workbook(excel_path, data_only=True)

    # find placeholders like [[EXCEL:SheetName]]
    for p in _iter_all_paragraphs(doc):
        text = p.text or ""
        if "[[EXCEL:" not in text:
            continue

        # naive parse: take first placeholder in paragraph
        start = text.find("[[EXCEL:")
        end = text.find("]]", start)
        if end == -1:
            continue

        placeholder = text[start : end + 2]
        sheet_name = placeholder.replace("[[EXCEL:", "").replace("]]", "").strip()

        if sheet_name not in wb.sheetnames:
            # replace placeholder to make it visible in doc
            p.text = text.replace(placeholder, f"[Worksheet '{sheet_name}' not found]")
            continue

        ws = wb[sheet_name]

        # remove placeholder paragraph text
        p.text = text.replace(placeholder, "").strip()

        # create table under paragraph (simple)
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        # add Word table at end, then move it later is complex → simplest: add after by adding it near end
        # Many projects accept it inserted at end of doc; if you need exact position, we can enhance it.
        table = doc.add_table(rows=len(rows), cols=len(rows[0]) if rows[0] else 1)
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                table.cell(r_idx, c_idx).text = "" if val is None else str(val)

    doc.save(docx_out_path)
