from __future__ import annotations

import base64
import json
import logging
import re
import unicodedata
from io import BytesIO
from html.parser import HTMLParser
from pathlib import Path
from typing import Dict, List, Optional, Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Inches, RGBColor
from openpyxl import load_workbook


EXCEL_TAG_RE = re.compile(r"\[\[\s*EXC?EL\s*:\s*([^\]]+?)\s*\]\]", re.IGNORECASE)
SEC_TAG_RE = re.compile(r"\[\[\s*(SEC_[A-Za-z0-9_]+)\s*\]\]")
INSERER_TAG_RE = re.compile(
    r"\[\[\s*(?:ins[eé]rer|insert(?:ion)?)(?:[^\]]*)\]\]",
    re.IGNORECASE,
)  # [[INSERER]], [[insérer ...]], [[insert ...]], [[insertion ...]]
TEXT_PLACEHOLDER_RE = re.compile(r"\[\[TEXT:(.*?)\]\]")
LOGO_PLACEHOLDER_RE = re.compile(r"\[\[LOGO\]\]")

def _extract_inserer_label(text: str) -> Optional[str]:
    match = INSERER_TAG_RE.search(text or "")
    if not match:
        return None
    raw = match.group(0)
    inner = raw.replace("[[", "").replace("]]", "").strip()
    return inner or None
ASTERISK_TITLE_RE = re.compile(r"\*\s*$")

NUM_PREFIX_RE = re.compile(r"^\s*\d+(?:\.\d+)*\s+")
WARNING_PREFIX_RE = re.compile(r"^[\s\u2022\-\*]*AVERTISSEMENT\s*:\s*r[ée]ponse\s*manquante\.\s*", re.IGNORECASE)

logger = logging.getLogger(__name__)


# ----------------------------
# Low-level iteration helpers
# ----------------------------
def _iter_block_items(doc: Document):
    parent_elm = doc.element.body
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def _iter_story_block_items(container):
    """Yield Paragraph/Table in document order for a docx story container (body/header/footer)."""
    element = getattr(container, "element", None)
    if element is not None and hasattr(element, "body"):
        parent_elm = element.body
    else:
        parent_elm = getattr(container, "_element", None)
    if parent_elm is None:
        return
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, container)
        elif child.tag.endswith("}tbl"):
            yield Table(child, container)


def _iter_all_story_containers(doc: Document):
    """
    Iterate over body + headers/footers so placeholder replacement also works
    outside detected heading sections.
    """
    yield doc
    for section in doc.sections:
        header = getattr(section, "header", None)
        footer = getattr(section, "footer", None)
        if header is not None:
            yield header
        if footer is not None:
            yield footer

def _is_heading(par: Paragraph) -> bool:
    style = (par.style.name or "") if par.style else ""
    return style.startswith("Heading")

def _heading_level(par: Paragraph) -> int:
    style = (par.style.name or "") if par.style else ""
    m = re.search(r"Heading\s+(\d+)", style)
    return int(m.group(1)) if m else 0

def _delete_paragraph(paragraph: Paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def _delete_table(tbl: Table):
    t = tbl._element
    t.getparent().remove(t)
    tbl._tbl = tbl._element = None

def _normalize_title(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip()).lower()

def _has_trailing_asterisk(title: str) -> bool:
    return bool(ASTERISK_TITLE_RE.search((title or "").strip()))


# ----------------------------
# 1) Remove sections by titles
# ----------------------------
def remove_sections_by_titles(doc: Document, titles_to_remove: List[str]):
    if not titles_to_remove:
        return

    wanted = {_normalize_title(t) for t in titles_to_remove if t and t.strip()}
    if not wanted:
        return

    items = list(_iter_block_items(doc))
    i = 0
    while i < len(items):
        it = items[i]
        if isinstance(it, Paragraph) and _is_heading(it):
            level = _heading_level(it)
            title_clean = _normalize_title(it.text)

            match = any(w in title_clean for w in wanted)
            if match:
                _delete_paragraph(it)

                j = i + 1
                while j < len(items):
                    nxt = items[j]
                    if isinstance(nxt, Paragraph) and _is_heading(nxt):
                        nxt_level = _heading_level(nxt)
                        if nxt_level and nxt_level <= level:
                            break
                    if isinstance(nxt, Paragraph):
                        _delete_paragraph(nxt)
                    else:
                        _delete_table(nxt)
                    j += 1

                items = list(_iter_block_items(doc))
                i = 0
                continue

        i += 1


# ----------------------------
# 1bis) Remove headings ending with "*" without [[SEC_x]]
# ----------------------------
def remove_asterisk_headings_without_sections(doc: Document):
    items = list(_iter_block_items(doc))
    i = 0

    while i < len(items):
        it = items[i]
        if isinstance(it, Paragraph) and _is_heading(it):
            level = _heading_level(it)
            title = (it.text or "").strip()

            if _has_trailing_asterisk(title):
                j = i + 1
                has_sec_tag = False
                while j < len(items):
                    nxt = items[j]
                    if isinstance(nxt, Paragraph) and _is_heading(nxt):
                        nxt_level = _heading_level(nxt)
                        if nxt_level and nxt_level <= level:
                            break
                    if isinstance(nxt, Paragraph):
                        if SEC_TAG_RE.search(nxt.text or ""):
                            has_sec_tag = True
                            break
                    else:
                        for row in nxt.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    if SEC_TAG_RE.search(p.text or ""):
                                        has_sec_tag = True
                                        break
                                if has_sec_tag:
                                    break
                            if has_sec_tag:
                                break
                    if has_sec_tag:
                        break
                    j += 1

                if not has_sec_tag:
                    for k in range(i, j):
                        block = items[k]
                        if isinstance(block, Paragraph):
                            _delete_paragraph(block)
                        else:
                            _delete_table(block)
                    items = list(_iter_block_items(doc))
                    i = 0
                    continue

        i += 1


# ----------------------------
# 2) Replace [[SEC_x]] tags
# ----------------------------
def _clear_paragraph_runs(paragraph: Paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def _add_multiline_text(paragraph: Paragraph, value: str) -> None:
    lines = value.splitlines()
    if not lines:
        paragraph.add_run("")
        return
    run = paragraph.add_run(lines[0])
    for line in lines[1:]:
        run.add_break()
        run.add_text(line)


def _replace_text_placeholders_in_paragraph(paragraph: Paragraph, values_by_id: Dict[str, str], counters: Dict[str, int]) -> None:
    text = paragraph.text or ""
    if not TEXT_PLACEHOLDER_RE.search(text):
        return

    chunks: list[tuple[str, str]] = []
    last = 0
    for match in TEXT_PLACEHOLDER_RE.finditer(text):
        before = text[last:match.start()]
        if before:
            chunks.append(("text", before))

        counters["text"] += 1
        placeholder_id = f"text_{counters['text']}"
        value = str(values_by_id.get(placeholder_id) or "")
        chunks.append(("multiline", value))
        last = match.end()

    tail = text[last:]
    if tail:
        chunks.append(("text", tail))

    _clear_paragraph_runs(paragraph)
    for kind, val in chunks:
        if kind == "text":
            paragraph.add_run(val)
        else:
            _add_multiline_text(paragraph, val)


def _replace_logo_in_paragraph(paragraph: Paragraph, logo_path: Optional[str]) -> None:
    text = paragraph.text or ""
    if not LOGO_PLACEHOLDER_RE.search(text):
        return

    if not logo_path:
        cleaned = LOGO_PLACEHOLDER_RE.sub("", text)
        _clear_paragraph_runs(paragraph)
        if cleaned:
            paragraph.add_run(cleaned)
        return

    parts = LOGO_PLACEHOLDER_RE.split(text, maxsplit=1)
    before = parts[0] if parts else ""
    after = parts[1] if len(parts) > 1 else ""

    _clear_paragraph_runs(paragraph)
    if before:
        paragraph.add_run(before)
    run = paragraph.add_run()
    run.add_picture(str(logo_path), width=Inches(1.5))
    if after:
        paragraph.add_run(after)


def inject_general_placeholders(doc: Document, text_placeholder_values: Optional[Dict[str, str]], logo_path: Optional[str]) -> None:
    values_by_id = {str(k): str(v or "") for k, v in (text_placeholder_values or {}).items()}
    counters = {"text": 0}

    for p in doc.paragraphs:
        _replace_text_placeholders_in_paragraph(p, values_by_id, counters)
        _replace_logo_in_paragraph(p, logo_path)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_text_placeholders_in_paragraph(p, values_by_id, counters)
                    _replace_logo_in_paragraph(p, logo_path)

def replace_sec_tags(
    doc: Document,
    sections_text: Dict[str, str],
):
    if not sections_text:
        return

    for p in doc.paragraphs:
        m = SEC_TAG_RE.search(p.text or "")
        if not m:
            continue

        sec_key = m.group(1)
        new_text = (sections_text.get(sec_key) or "").strip()
        new_text = WARNING_PREFIX_RE.sub("", new_text).strip()
        if not new_text:
            new_text = "TBD"
        is_tbd = new_text == "TBD"

        only_tag = _normalize_title(SEC_TAG_RE.sub("", p.text)) == ""
        if only_tag:
            _clear_paragraph_runs(p)
            if new_text:
                run = p.add_run(new_text)
                if is_tbd:
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            continue

        full_text = p.text or ""
        parts = SEC_TAG_RE.split(full_text, maxsplit=1)
        before_text = parts[0] if parts else ""
        after_text = parts[2] if len(parts) > 2 else ""
        _clear_paragraph_runs(p)
        if before_text:
            p.add_run(before_text)
        if new_text:
            run = p.add_run(new_text)
            if is_tbd:
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        if after_text:
            p.add_run(after_text)


# ----------------------------
# 3) Inject Excel tables
# ----------------------------
def _clean_cell(v):
    if v is None:
        return ""
    s = str(v)
    s = s.replace("\r", " ").replace("\n", " ")
    s = " ".join(s.split())
    return s

def _insert_table_after(paragraph: Paragraph, rows: List[List[str]]):
    doc = paragraph.part.document

    ncols = max(len(r) for r in rows) if rows else 1
    tbl = doc.add_table(rows=0, cols=ncols)
    tbl.style = "Table Grid"
    tbl.autofit = True

    for row in rows:
        tr = tbl.add_row().cells
        for i in range(ncols):
            tr[i].text = _clean_cell(row[i] if i < len(row) else "")

    if ncols == 4:
        widths = [Inches(1.4), Inches(1.5), Inches(1.4), Inches(3.2)]
        for row in tbl.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w

    paragraph._p.addnext(tbl._tbl)
    return tbl

def inject_excel_tables(doc: Document, worksheet_path: Optional[str]):
    if not worksheet_path:
        return

    xls_path = Path(worksheet_path)
    if not xls_path.exists():
        return

    wb = load_workbook(filename=str(xls_path), data_only=True)

    for p in list(doc.paragraphs):
        m = EXCEL_TAG_RE.search(p.text or "")
        if not m:
            continue

        sheet_name = m.group(1).strip()
        if sheet_name not in wb.sheetnames:
            low = sheet_name.lower()
            match = next((s for s in wb.sheetnames if s.lower() == low), None)
            if not match:
                p.text = EXCEL_TAG_RE.sub("", p.text)
                continue
            sheet_name = match

        ws = wb[sheet_name]

        values = []
        for row in ws.iter_rows(values_only=True):
            if row is None:
                continue
            if any(c is not None and str(c).strip() != "" for c in row):
                values.append([("" if c is None else c) for c in row])

        if not values:
            p.text = EXCEL_TAG_RE.sub("", p.text)
            continue

        p.text = EXCEL_TAG_RE.sub("", p.text).strip()
        _insert_table_after(p, values)


# ----------------------------
# 4) Renumber headings
# ----------------------------
def renumber_headings(doc: Document):
    counters = [0, 0, 0, 0, 0, 0]

    for p in doc.paragraphs:
        if not _is_heading(p):
            continue
        level = _heading_level(p)
        if not (1 <= level <= 6):
            continue

        counters[level - 1] += 1
        for i in range(level, 6):
            counters[i] = 0

        nums = [str(counters[i]) for i in range(level) if counters[i] > 0]
        prefix = ".".join(nums)

        title = NUM_PREFIX_RE.sub("", p.text).strip()
        p.text = f"{prefix} {title}"


# ----------------------------
# 5) Inject TinyMCE HTML into [[INSERER]]
# ----------------------------
def _insert_paragraph_after(anchor: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    return Paragraph(new_p, anchor._parent)

def _insert_picture_after(anchor: Paragraph, image_bytes: bytes, width_in_inches: float = 6.0):
    p = _insert_paragraph_after(anchor)
    run = p.add_run()
    run.add_picture(BytesIO(image_bytes), width=Inches(width_in_inches))
    return p

class _TinyHtmlParser(HTMLParser):
    """
    Convert a small subset of HTML to blocks:
      - paragraphs (p/div)
      - line breaks
      - bold/italic/underline
      - unordered/ordered lists
      - img with base64 data URL
    """
    def __init__(self):
        super().__init__()
        self.blocks: List[Dict[str, Any]] = []
        self._cur_runs: List[Dict[str, Any]] = []
        self._in_p = False

        self._style_stack: List[Dict[str, bool]] = []
        self._list_stack: List[str] = []  # "ul" or "ol"
        self._in_li = False

    def _cur_style(self) -> Dict[str, bool]:
        style = {"bold": False, "italic": False, "underline": False}
        for s in self._style_stack:
            for k in style:
                style[k] = style[k] or bool(s.get(k))
        return style

    def _start_paragraph(self):
        if self._in_p and self._cur_runs:
            self._end_paragraph()
        self._in_p = True
        self._cur_runs = []

    def _end_paragraph(self):
        txt = "".join(r["text"] for r in self._cur_runs).strip()
        # even if empty, keep if it contains image blocks separately
        self.blocks.append({"type": "p", "runs": self._cur_runs[:]})
        self._cur_runs = []
        self._in_p = False

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        attrs = dict(attrs or [])

        if tag in ("p", "div"):
            self._start_paragraph()
            return

        if tag == "br":
            if not self._in_p:
                self._start_paragraph()
            self._cur_runs.append({"text": "\n", **self._cur_style()})
            return

        if tag in ("strong", "b"):
            self._style_stack.append({"bold": True})
            return
        if tag in ("em", "i"):
            self._style_stack.append({"italic": True})
            return
        if tag == "u":
            self._style_stack.append({"underline": True})
            return

        if tag in ("ul", "ol"):
            self._list_stack.append(tag)
            return

        if tag == "li":
            self._in_li = True
            self._start_paragraph()
            return

        if tag == "img":
            src = (attrs.get("src") or "").strip()
            if src.startswith("data:image/") and "base64," in src:
                b64 = src.split("base64,", 1)[1]
                try:
                    img_bytes = base64.b64decode(b64)
                    self.blocks.append({"type": "img", "bytes": img_bytes})
                except Exception:
                    pass
            return

    def handle_endtag(self, tag):
        tag = tag.lower()

        if tag in ("strong", "b", "em", "i", "u"):
            if self._style_stack:
                self._style_stack.pop()
            return

        if tag == "li":
            # close list item paragraph
            if self._in_p:
                # mark paragraph as list paragraph
                list_type = self._list_stack[-1] if self._list_stack else "ul"
                self.blocks.append({"type": "li", "list": list_type, "runs": self._cur_runs[:]})
                self._cur_runs = []
                self._in_p = False
            self._in_li = False
            return

        if tag in ("ul", "ol"):
            if self._list_stack:
                self._list_stack.pop()
            return

        if tag in ("p", "div"):
            if self._in_p:
                self._end_paragraph()
            return

    def handle_data(self, data):
        txt = (data or "")
        if not txt:
            return
        if not self._in_p:
            self._start_paragraph()
        st = self._cur_style()
        self._cur_runs.append({"text": txt, **st})

def _apply_runs_to_paragraph(p: Paragraph, runs: List[Dict[str, Any]]):
    # python-docx doesn’t keep multiple styles if you just do p.text
    for r in runs:
        t = r.get("text", "")
        if not t:
            continue
        run = p.add_run(t)
        run.bold = bool(r.get("bold"))
        run.italic = bool(r.get("italic"))
        run.underline = bool(r.get("underline"))

def _get_heading_text(par: Paragraph) -> Optional[str]:
    style = (par.style.name or "") if par.style else ""
    if style.startswith("Heading") or style.startswith("Titre"):
        return (par.text or "").strip()
    return None

def _inject_html_into_paragraph(paragraph: Paragraph, html: str):
    only_tag = _normalize_title(INSERER_TAG_RE.sub("", paragraph.text or "")) == ""
    html = (html or "").strip()

    if not html:
        if only_tag:
            _delete_paragraph(paragraph)
        else:
            paragraph.text = INSERER_TAG_RE.sub("", paragraph.text)
        return

    parser = _TinyHtmlParser()
    parser.feed(html)
    blocks = parser.blocks

    if only_tag:
        anchor = paragraph
        anchor.text = ""
        insert_after = anchor

        for b in blocks:
            if b["type"] == "img":
                insert_after = _insert_picture_after(insert_after, b["bytes"], width_in_inches=6.0)
                continue

            if b["type"] == "li":
                np = _insert_paragraph_after(insert_after)
                try:
                    np.style = "List Bullet" if b["list"] == "ul" else "List Number"
                except Exception:
                    prefix = "• " if b["list"] == "ul" else "1) "
                    np.add_run(prefix)
                _apply_runs_to_paragraph(np, b.get("runs", []))
                insert_after = np
                continue

            np = _insert_paragraph_after(insert_after)
            _apply_runs_to_paragraph(np, b.get("runs", []))
            insert_after = np

        _delete_paragraph(anchor)
        return

    paragraph.text = INSERER_TAG_RE.sub("", paragraph.text).strip()
    insert_after = paragraph

    for b in blocks:
        if b["type"] == "img":
            insert_after = _insert_picture_after(insert_after, b["bytes"], width_in_inches=6.0)
            continue

        if b["type"] == "li":
            np = _insert_paragraph_after(insert_after)
            try:
                np.style = "List Bullet" if b["list"] == "ul" else "List Number"
            except Exception:
                prefix = "• " if b["list"] == "ul" else "1) "
                np.add_run(prefix)
            _apply_runs_to_paragraph(np, b.get("runs", []))
            insert_after = np
            continue

        np = _insert_paragraph_after(insert_after)
        _apply_runs_to_paragraph(np, b.get("runs", []))
        insert_after = np

def _build_insert_map_from_payload(data: Any) -> tuple[Dict[str, str], int]:
    """Return normalized {placeholder -> html} map + count of NULL section_id entries."""
    if data is None:
        return {}, 0

    if isinstance(data, dict) and data.get("__type") == "map" and isinstance(data.get("items"), dict):
        items = {str(k): str(v or "") for k, v in data.get("items", {}).items() if v is not None}
        return items, 0

    entries: list[dict[str, Any]] = []
    if isinstance(data, list):
        entries = [item for item in data if isinstance(item, dict)]
    elif isinstance(data, dict):
        maybe_entries = data.get("items") if isinstance(data.get("items"), list) else data.get("insertions")
        if isinstance(maybe_entries, list):
            entries = [item for item in maybe_entries if isinstance(item, dict)]

    if not entries:
        return {}, 0

    result: Dict[str, str] = {}
    null_section_count = 0
    for entry in entries:
        key = (
            entry.get("placeholder")
            or entry.get("insert_key")
            or entry.get("label")
            or entry.get("key")
            or ""
        )
        value = entry.get("html")
        if value is None:
            value = entry.get("insert_html")
        if value is None:
            value = entry.get("content")

        if entry.get("section_id") is None:
            null_section_count += 1

        key_str = str(key or "").strip()
        if not key_str:
            continue
        result[key_str] = str(value or "")

    return result, null_section_count


def _parse_insert_html_payload(insert_html: Optional[str]) -> tuple[Optional[str], Optional[Dict[str, str]], int]:
    if not insert_html:
        return None, None, 0

    raw = (insert_html or "").strip()
    if raw.startswith("{") or raw.startswith("["):
        try:
            data = json.loads(raw)
        except json.JSONDecodeError:
            return raw, None, 0

        insert_map, null_section_count = _build_insert_map_from_payload(data)
        if insert_map:
            return None, insert_map, null_section_count

    return raw, None, 0


def inject_insert_html_map(doc: Document, insert_map: Dict[str, str], null_section_insertions: int = 0):
    if not insert_map:
        logger.debug("insertions_map empty: nothing to apply")
        return

    def _normalize_insert_lookup(value: Optional[str]) -> str:
        normalized = _normalize_title(value or "")
        if not normalized:
            return ""
        return "".join(
            c for c in unicodedata.normalize("NFKD", normalized)
            if not unicodedata.combining(c)
        )

    normalized_map = {
        _normalize_insert_lookup(str(key or "")): str(value or "")
        for key, value in insert_map.items()
        if str(key or "").strip()
    }

    def resolve_insert_html(current_heading: Optional[str], label: Optional[str]) -> str:
        normalized_label = _normalize_insert_lookup(label or "")
        if normalized_label:
            direct = normalized_map.get(normalized_label, "")
            if direct:
                return direct
            for prefix in ("insérer", "inserer", "insert", "insertion"):
                if not normalized_label.startswith(prefix):
                    continue
                suffix = normalized_label[len(prefix):].strip()
                if suffix:
                    plain = normalized_map.get(suffix, "")
                    if plain:
                        return plain

        if current_heading:
            return normalized_map.get(_normalize_insert_lookup(current_heading), "")
        return ""

    placeholders_found = 0
    placeholders_replaced = 0

    logger.debug(
        "insertions_map loaded: total_insertions=%s null_section_insertions=%s",
        len(insert_map),
        null_section_insertions,
    )

    for container in _iter_all_story_containers(doc):
        current_heading = None
        for item in list(_iter_story_block_items(container)):
            if isinstance(item, Paragraph):
                heading_text = _get_heading_text(item)
                if heading_text:
                    current_heading = _normalize_title(heading_text)

                if INSERER_TAG_RE.search(item.text or ""):
                    placeholders_found += 1
                    label = _extract_inserer_label(item.text or "")
                    html = resolve_insert_html(current_heading, label)
                    if (html or "").strip():
                        placeholders_replaced += 1
                    _inject_html_into_paragraph(item, html)
                continue

            for row in item.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if not INSERER_TAG_RE.search(p.text or ""):
                            continue
                        placeholders_found += 1
                        label = _extract_inserer_label(p.text or "")
                        html = resolve_insert_html(current_heading, label)
                        if (html or "").strip():
                            placeholders_replaced += 1
                        _inject_html_into_paragraph(p, html)

    logger.debug(
        "insertions_apply done: placeholders_found=%s placeholders_replaced=%s placeholders_without_match=%s",
        placeholders_found,
        placeholders_replaced,
        max(placeholders_found - placeholders_replaced, 0),
    )

def inject_insert_html(doc: Document, insert_html: Optional[str]):
    """
    Replace every [[INSERER]] / [[insérer]] in the DOCX by the HTML content.
    If insert_html is empty => remove the placeholder only.
    """
    html, insert_map, null_section_insertions = _parse_insert_html_payload(insert_html)
    if insert_map:
        inject_insert_html_map(doc, insert_map, null_section_insertions)
        return

    html = (html or "").strip()

    for container in _iter_all_story_containers(doc):
        for item in list(_iter_story_block_items(container)):
            if isinstance(item, Paragraph):
                if INSERER_TAG_RE.search(item.text or ""):
                    _inject_html_into_paragraph(item, html)
                continue

            for row in item.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if INSERER_TAG_RE.search(p.text or ""):
                            _inject_html_into_paragraph(p, html)


def mark_update_fields(doc: Document) -> None:
    settings = getattr(doc._part, "settings", None)
    if settings is None:
        return
    settings = getattr(settings, "element", None)
    if settings is None:
        return
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


# ----------------------------
# Main pipeline
# ----------------------------
def apply_doc_pipeline(
    template_path: str,
    out_path: str,
    sections_text: Dict[str, str],
    titles_to_remove: List[str],
    worksheet_path: Optional[str] = None,
    insert_html: Optional[str] = None,   # ✅ NEW
    text_placeholder_values: Optional[Dict[str, str]] = None,
    logo_path: Optional[str] = None,
):
    doc = Document(template_path)

    # 1) remove sections
    remove_sections_by_titles(doc, titles_to_remove)

    # 1bis) remove headings ending with "*" without [[SEC_x]]
    remove_asterisk_headings_without_sections(doc)

    # 1bis) renumber headings after removals
    renumber_headings(doc)

    # 2) replace [[SEC_x]]
    replace_sec_tags(doc, sections_text)

    # 3) inject excel
    inject_excel_tables(doc, worksheet_path)

    # ✅ 4) replace [[INSERER]] by TinyMCE content (at exact location)
    inject_insert_html(doc, insert_html)

    # 5) replace [[TEXT: ...]] and [[LOGO]]
    inject_general_placeholders(doc, text_placeholder_values, logo_path)

    mark_update_fields(doc)

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
