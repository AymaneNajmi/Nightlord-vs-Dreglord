import hashlib
import json
import logging
import os
import re
import time
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple

from docx import Document
from openai import OpenAI
from sqlalchemy.orm import Session

from app.models.ai_template_job import AITemplateJob, AITemplateJobStatus
from app.models.forms import FormOption, FormQuestion, FormSection, FormTemplate
from app.models.techno import DocType, Techno
from app.models.template_doc import TemplateDoc
from app.schemas.ai_template_builder import (
    AITemplateOutput,
    FormQuestionSpec,
    FormSectionSpec,
    SectionQuestionsOutput,
    StyleGuide,
    TemplateOutlineItem,
)
from app.services.docx_headings import extract_headings
from app.services.docx_pipeline import renumber_headings
from app.services.question_quality import (
    enforce_question_quality,
    is_editorial_section,
)
from app.services.llm_provider import LLMProvider, call_llm_json, get_provider
from app.services.section_context import build_section_context, extract_sections_from_docx_by_headings

BASE_DIR = Path(__file__).resolve().parent.parent
AI_TEMPLATES_DIR = BASE_DIR / "storage" / "ai_templates"
AI_TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

MAX_FILE_SIZE = 25 * 1024 * 1024
MAX_PROMPT_CHARS = 20000
OUTLINE_CONTEXT_CHARS = int(os.getenv("AI_TEMPLATE_OUTLINE_CONTEXT_CHARS", "9000"))
QUESTION_CONTEXT_CHARS = int(os.getenv("AI_TEMPLATE_QUESTION_CONTEXT_CHARS", "7000"))
MAX_HEADING_COUNT = 120
SNIPPET_CHUNKS = 3
SNIPPET_SEPARATOR = "\n--- EXTRACT ---\n"
MAX_CHUNKS = int(os.getenv("AI_TEMPLATE_MAX_CHUNKS", "2"))
MAX_JOB_SECONDS = int(os.getenv("AI_TEMPLATE_MAX_SECONDS", "420"))
OUTLINE_MODEL = os.getenv("OPENAI_MODEL_OUTLINE", "gpt-4.1-mini")
FORM_MODEL = os.getenv("OPENAI_MODEL_FORM", "gpt-4.1")

logger = logging.getLogger(__name__)

SEC_PLACEHOLDER_RE = re.compile(r"^\[\[\s*SEC_[A-Za-z0-9_]+\s*\]\]$", re.IGNORECASE)
EXCEL_PLACEHOLDER_RE = re.compile(r"^\[\[\s*Excel\s*:\s*[^\]]+\]\]$", re.IGNORECASE)
INSERT_PLACEHOLDER_RE = re.compile(r"^\[\[\s*ins[ée]rer\s+[^\]]+\]\]$", re.IGNORECASE)
ANY_PLACEHOLDER_RE = re.compile(r"\[\[[^\]]+\]\]")
MEANINGFUL_SECTION_MIN_CHARS = int(os.getenv("AI_TEMPLATE_MEANINGFUL_SECTION_MIN_CHARS", "150"))
MEANINGFUL_SECTION_MIN_WORDS = int(os.getenv("AI_TEMPLATE_MEANINGFUL_SECTION_MIN_WORDS", "12"))
MIN_SECTION_QUESTIONS = int(os.getenv("AI_TEMPLATE_MIN_SECTION_QUESTIONS", "2"))
MAX_SECTION_QUESTIONS = int(os.getenv("AI_TEMPLATE_MAX_SECTION_QUESTIONS", "5"))


def _safe_filename(name: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9_.-]", "_", name.strip())
    return cleaned[:180] or "document"


def _hash_file(path: Path) -> str:
    sha = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(8192), b""):
            sha.update(chunk)
    return sha.hexdigest()


def store_uploads(job_id: int, files: Iterable[Tuple[str, bytes]]) -> List[Dict[str, Any]]:
    job_dir = AI_TEMPLATES_DIR / f"job_{job_id}"
    job_dir.mkdir(parents=True, exist_ok=True)
    metadata: List[Dict[str, Any]] = []
    for filename, content in files:
        safe = _safe_filename(filename)
        path = job_dir / safe
        path.write_bytes(content)
        metadata.append(
            {
                "filename": safe,
                "size": len(content),
                "sha256": _hash_file(path),
                "stored_path": str(path),
            }
        )
    return metadata


def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    parts: List[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_tables_from_docx(path: str) -> List[str]:
    doc = Document(path)
    tables: List[str] = []
    for table in doc.tables:
        rows: List[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            tables.append("\n".join(rows))
    return tables


def extract_headings_from_docx(path: str) -> List[str]:
    headings = extract_headings(path)
    titles: List[str] = []
    for item in headings:
        if isinstance(item, dict):
            title = item.get("title") or item.get("text")
            level = item.get("level")
        else:
            title = getattr(item, "text", None)
            level = getattr(item, "level", None)
        if title and level and int(level) <= 4:
            normalized_title = str(title).strip()
            if not _is_meaningful_heading(normalized_title):
                continue
            titles.append(normalized_title)
    return titles


def _normalize_heading_key(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", (text or "").strip())
    cleaned = cleaned.rstrip(":").strip()
    return cleaned.lower()


def _title_case_score(text: str) -> float:
    words = [w for w in re.split(r"\s+", (text or "").strip()) if w]
    if not words:
        return 0.0
    titled = 0
    for w in words:
        core = w.strip(" _-.,:;()[]{}")
        if core and core[:1].isalpha() and core[:1].isupper():
            titled += 1
    return titled / len(words)


def _canonical_title_key(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", (text or ""))
    normalized = "".join(ch for ch in normalized if not unicodedata.combining(ch))
    normalized = re.sub(r"[^a-z0-9]+", " ", normalized.casefold()).strip()
    return normalized


def _clean_title(title: str) -> str:
    cleaned = re.sub(r"\s+", " ", (title or "").strip())
    if not cleaned:
        return ""

    slug_match = re.match(r"^([a-z0-9-]{3,})\s+(.+)$", cleaned)
    if slug_match and "-" in slug_match.group(1):
        cleaned = slug_match.group(2).strip()

    duplicate_match = re.match(r"^([A-Za-zÀ-ÖØ-öø-ÿ0-9_-]+)\s+(.+)$", cleaned)
    if duplicate_match:
        first = duplicate_match.group(1).strip(" _-")
        second = duplicate_match.group(2).strip()
        if first and second:
            second_first = second.split()[0].strip(" _-:;")
            if first.casefold() == second_first.casefold():
                cleaned = second

    repeated_phrase = re.match(r"^(.+?)\s+(.+)$", cleaned)
    if repeated_phrase:
        left = repeated_phrase.group(1).strip()
        right = repeated_phrase.group(2).strip()
        left_key = _canonical_title_key(left)
        right_key = _canonical_title_key(right)
        if left_key and left_key == right_key:
            cleaned = right if _title_case_score(right) >= _title_case_score(left) else left

    words = [w for w in cleaned.split() if w]
    if len(words) >= 2 and len(words) % 2 == 0:
        mid = len(words) // 2
        left = " ".join(words[:mid]).strip()
        right = " ".join(words[mid:]).strip()
        left_key = _canonical_title_key(left)
        right_key = _canonical_title_key(right)
        if left_key and left_key == right_key:
            cleaned = right if _title_case_score(right) >= _title_case_score(left) else left

    return cleaned.rstrip(":").strip()


def _strip_parent_namespace_prefix(title: str, parent_title: str) -> str:
    candidate = (title or "").strip()
    parent = (parent_title or "").strip()
    if not candidate or not parent:
        return candidate

    prefix_pattern = re.compile(
        rf"^\s*{re.escape(parent)}(?:\s*[-–—:]\s*|\.(?:\d+|[A-Za-z0-9_-]+)\s+)",
        re.IGNORECASE,
    )
    stripped = prefix_pattern.sub("", candidate).strip()
    if stripped != candidate:
        return stripped

    # Fallback tolerant au bruit de casse/espaces: compare sur une forme canonique.
    canonical_candidate = _canonical_title_key(candidate)
    canonical_parent = _canonical_title_key(parent)
    if not canonical_candidate or not canonical_parent:
        return candidate

    parts = canonical_candidate.split()
    parent_parts = canonical_parent.split()
    if len(parts) > len(parent_parts) and parts[: len(parent_parts)] == parent_parts:
        remainder = re.sub(r"^\s*.+?(?:\.(?:\d+|[A-Za-z0-9_-]+)\s+|\s+)", "", candidate).strip()
        return remainder or candidate

    return candidate


def _clean_outline_titles(outline: List[TemplateOutlineItem]) -> List[TemplateOutlineItem]:
    parent_titles: Dict[int, str] = {}
    for item in outline:
        current_title = _clean_title(item.title)
        level = max(1, int(item.level or 1))
        parent_title = parent_titles.get(level - 1, "") if level > 1 else ""
        if parent_title:
            current_title = _strip_parent_namespace_prefix(current_title, parent_title)
            current_title = _clean_title(current_title)
        item.title = current_title
        parent_titles[level] = item.title or ""
        for deeper in [lv for lv in list(parent_titles.keys()) if lv > level]:
            parent_titles.pop(deeper, None)
    return outline


def _align_outline_titles_with_source(
    outline: List[TemplateOutlineItem],
    heading_title_map: Dict[str, str] | None = None,
) -> List[TemplateOutlineItem]:
    if not heading_title_map:
        return outline
    for item in outline:
        key = _normalize_heading_key(item.title or "")
        if key and key in heading_title_map:
            item.title = heading_title_map[key]
    return outline


def _is_sec_placeholder_heading(text: str) -> bool:
    normalized = (text or "").strip()
    if not normalized:
        return False
    lowered = normalized.lower()
    compact = re.sub(r"\s+", "", lowered)
    if re.fullmatch(r"sec_\d+(?:_\d+)?", compact):
        return True
    if "[[sec_" in lowered:
        return True
    if re.fullmatch(r"\[\[sec_\d+(?:_\d+)?\]\]", compact):
        return True
    return False


def _is_hardware_code_heading(text: str) -> bool:
    normalized = (text or "").strip().rstrip("=:")
    if not normalized:
        return False
    # Ex: C9115AXI-E, R-ISE-VMC-K9, CP-8800-A-KEM
    compact = re.sub(r"\s+", "", normalized)
    if " " not in normalized and re.fullmatch(r"[A-Za-z0-9_./-]{5,}", compact):
        has_digit = any(ch.isdigit() for ch in compact)
        has_sep = any(ch in "-_/" for ch in compact)
        if has_digit and has_sep:
            return True
    return False


def _is_meaningful_heading(text: str) -> bool:
    candidate = (text or "").strip()
    if not candidate:
        return False
    if _is_sec_placeholder_heading(candidate):
        return False
    if _is_hardware_code_heading(candidate):
        return False
    return True


def extract_heading_levels_from_docs(paths: Iterable[str]) -> Dict[str, int]:
    level_map: Dict[str, int] = {}
    for path in paths:
        if not path.lower().endswith(".docx"):
            continue
        for item in extract_headings(path):
            title = getattr(item, "text", None)
            level = getattr(item, "level", None)
            if title and level:
                _, stripped = _strip_number_prefix(title)
                key = _normalize_heading_key(stripped or title)
                if key and key not in level_map and int(level) <= 4:
                    level_map[key] = int(level)
    return level_map


def extract_heading_numbers_from_docs(paths: Iterable[str]) -> Dict[str, str]:
    number_map: Dict[str, str] = {}
    for path in paths:
        if not path.lower().endswith(".docx"):
            continue
        for item in extract_headings(path):
            title = getattr(item, "text", None)
            level = getattr(item, "level", None)
            if title and level and int(level) <= 4:
                number, stripped = _strip_number_prefix(title)
                key = _normalize_heading_key(stripped or title)
                if key and number and key not in number_map:
                    number_map[key] = number
    return number_map




def extract_heading_titles_from_docs(paths: Iterable[str]) -> Dict[str, str]:
    title_map: Dict[str, str] = {}
    for path in paths:
        if not path.lower().endswith(".docx"):
            continue
        for item in extract_headings(path):
            title = getattr(item, "text", None)
            level = getattr(item, "level", None)
            if title and level and int(level) <= 4:
                _, stripped = _strip_number_prefix(title)
                normalized_title = (stripped or title).strip()
                key = _normalize_heading_key(normalized_title)
                if key and key not in title_map:
                    title_map[key] = normalized_title
    return title_map

def extract_headings_from_docs(paths: Iterable[str]) -> List[str]:
    titles: List[str] = []
    for path in paths:
        if path.lower().endswith(".docx"):
            titles.extend(extract_headings_from_docx(path))
    seen = set()
    deduped = []
    for title in titles:
        key = title.strip().lower()
        if key and key not in seen:
            seen.add(key)
            deduped.append(title.strip())
    return deduped


def extract_headings_from_text(text: str) -> List[str]:
    titles: List[str] = []
    for line in (text or "").splitlines():
        raw = line.strip()
        if not raw:
            continue
        if len(raw) > 160:
            continue
        if not _is_meaningful_heading(raw):
            continue
        if "[[SEC_" in raw.upper():
            continue
        if re.search(r"\b\d{1,3}(?:\.\d{1,3}){3}\b", raw):
            continue
        digit_ratio = sum(ch.isdigit() for ch in raw) / max(1, len(raw))
        if digit_ratio > 0.4:
            continue
        if re.match(r"^\d+(\.\d+)*\s+.+", raw):
            titles.append(raw)
        elif raw.endswith(":") and len(raw) > 4:
            titles.append(raw.rstrip(":").strip())
    seen = set()
    deduped = []
    for title in titles:
        key = title.strip().lower()
        if key and key not in seen:
            seen.add(key)
            deduped.append(title.strip())
    return deduped


def _filter_heading_candidates(headings: List[str]) -> List[str]:
    filtered: List[str] = []
    for heading in headings:
        text = heading.strip()
        if not text:
            continue
        if not _is_meaningful_heading(text):
            continue
        if "[[SEC_" in text.upper():
            continue
        if re.search(r"\b\d{1,3}(?:\.\d{1,3}){3}\b", text):
            continue
        digit_ratio = sum(ch.isdigit() for ch in text) / max(1, len(text))
        if digit_ratio > 0.4:
            continue
        filtered.append(text)
    return filtered


def _build_outline_from_headings(
    headings: List[str],
    heading_levels: Dict[str, int],
    heading_numbers: Dict[str, str],
    style_map: Dict[int, str],
) -> List[TemplateOutlineItem]:
    outline: List[TemplateOutlineItem] = []
    for idx, heading in enumerate(headings, start=1):
        number, title = _strip_number_prefix(heading)
        key = _normalize_heading_key(title or heading)
        level = heading_levels.get(key) or _derive_level_from_title(heading)
        number = heading_numbers.get(key) or number
        outline.append(
            TemplateOutlineItem(
                id=f"SEC_{idx}_1",
                level=level,
                number=number or "",
                title=title or heading,
                style_name=_style_name_for_level(level, style_map),
                content="",
                markers=[],
            )
        )
    return outline


def extract_text_from_pdf(path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    parts: List[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def extract_text_from_docs(paths: Iterable[str]) -> str:
    chunks: List[str] = []
    for path in paths:
        if path.lower().endswith(".docx"):
            chunks.append(extract_text_from_docx(path))
        elif path.lower().endswith(".pdf"):
            chunks.append(extract_text_from_pdf(path))
        else:
            raise ValueError(f"Format non supporté: {path}")
    return "\n\n".join(chunks)


def extract_tables_from_docs(paths: Iterable[str]) -> List[str]:
    tables: List[str] = []
    for path in paths:
        if path.lower().endswith(".docx"):
            tables.extend(extract_tables_from_docx(path))
    return tables


def redact_confidential(text: str, names: Iterable[str]) -> Tuple[str, Dict[str, Any]]:
    redacted = text or ""
    report_items: List[Dict[str, Any]] = []
    total = 0
    for idx, raw in enumerate(names, start=1):
        name = (raw or "").strip()
        if not name:
            continue
        pattern = re.compile(re.escape(name), re.IGNORECASE)
        replacement = "CLIENT"
        redacted, count = pattern.subn(replacement, redacted)
        total += count
        report_items.append(
            {
                "name": name,
                "replacement": replacement,
                "occurrences": count,
            }
        )
    return redacted, {
        "generated_at": datetime.utcnow().isoformat(),
        "total_redactions": total,
        "items": report_items,
    }


def replace_client_names(text: str, client_names: List[str], replacement: str = "CLIENT") -> str:
    if not text:
        return text
    updated = text
    for raw in client_names:
        name = (raw or "").strip()
        if not name:
            continue
        pattern = re.compile(re.escape(name), re.IGNORECASE)
        updated = pattern.sub(replacement, updated)
    return updated


def deep_replace(obj: Any, client_names: List[str]) -> Any:
    if isinstance(obj, str):
        return replace_client_names(obj, client_names)
    if isinstance(obj, list):
        return [deep_replace(item, client_names) for item in obj]
    if isinstance(obj, dict):
        return {key: deep_replace(value, client_names) for key, value in obj.items()}
    return obj


def sanitize_network_text(text: str) -> Tuple[str, Dict[str, Any]]:
    sanitized = text or ""
    report_items: List[Dict[str, Any]] = []

    patterns: List[Tuple[str, str, str]] = [
        ("cidr", r"\b(?:\d{1,3}\.){3}\d{1,3}/(?:[0-9]|[12][0-9]|3[0-2])\b", "TBD_CIDR"),
        (
            "ipv4",
            r"\b(?:(?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)\.){3}(?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)\b",
            "TBD_IP",
        ),
        (
            "ipv6",
            r"\b(?:[0-9a-fA-F]{1,4}:){2,7}[0-9a-fA-F]{1,4}\b|\b::1\b|\b::\b",
            "TBD_IPV6",
        ),
        ("email", r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", "TBD_EMAIL"),
        ("url", r"\b(?:https?|ftp)://[^\s)\]>]+", "TBD_URL"),
        ("mac", r"\b(?:[0-9A-Fa-f]{2}[:-]){5}[0-9A-Fa-f]{2}\b", "TBD_MAC"),
        (
            "secret_assignment",
            r"(?im)(\b(?:token|secret|password|passwd|psk|api[_-]?key|private[_-]?key|key)\b\s*[:=]\s*)([^\s,;]+)",
            r"\1TBD_SECRET",
        ),
    ]

    for name, pattern, replacement in patterns:
        sanitized, count = re.subn(pattern, replacement, sanitized)
        report_items.append(
            {
                "type": name,
                "replacement": replacement,
                "occurrences": count,
            }
        )

    config_line_pattern = re.compile(
        r"(?im)^\s*(?:access-list|crypto|username|set\s+|set$|ikev2|ipsec|tunnel-group|"
        r"snmp-server|radius-server|tacacs-server|enable\s+secret|line\s+vty|"
        r"community\s+|pre-shared-key|psk\s+|key\s+)(?:\b|\s|[:=]).*$"
    )
    sanitized, config_count = config_line_pattern.subn("[REDACTED_CONFIG_LINE]", sanitized)
    report_items.append(
        {
            "type": "config_line",
            "replacement": "[REDACTED_CONFIG_LINE]",
            "occurrences": config_count,
        }
    )

    total = sum(item["occurrences"] for item in report_items)
    return sanitized, {
        "generated_at": datetime.utcnow().isoformat(),
        "total_redactions": total,
        "items": report_items,
    }


def _extract_json(text: str) -> Dict[str, Any]:
    text = text or ""
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1:
        raise ValueError(f"JSON introuvable. Début réponse: {text[:200]}")
    return json.loads(text[start : end + 1])


def _schema_json(schema: Dict[str, Any]) -> str:
    return json.dumps(schema, ensure_ascii=False, indent=2)


def _sanitize_section_questions_payload(data: Dict[str, Any]) -> Dict[str, Any]:
    questions = data.get("questions")
    if not isinstance(questions, list):
        return data

    for question in questions:
        if not isinstance(question, dict):
            continue
        choices = question.get("choices")
        if not isinstance(choices, list):
            continue
        normalized_choices: List[str] = []
        seen = set()
        for choice in choices:
            if not isinstance(choice, str):
                continue
            cleaned = re.sub(r"\s+", " ", choice).strip()
            if not cleaned:
                continue
            key = cleaned.lower()
            if key in seen:
                continue
            seen.add(key)
            normalized_choices.append(cleaned)
            if len(normalized_choices) >= 10:
                break
        question["choices"] = normalized_choices

    return data


SYSTEM_PROMPT_OUTLINE = """
IDENTITÉ
Tu es un architecte réseau et sécurité senior chargé de produire un modèle de dossier d'ingénierie.

OBJECTIF
Générer:
- Une structure documentaire très détaillée (H1/H2/H3/H4).
- Des marqueurs [[SEC_*]] pour les sections à renseigner.

RÈGLES STRICTES
- Ne jamais inventer des valeurs spécifiques client (IP réelles, noms, modèles exacts non cités).
- Tu peux enrichir avec des bonnes pratiques génériques (LLD-ready).
- Si une information manque, créer une section [[SEC_*]] et des questions pour la collecter.
- Les sections éditoriales restent textuelles (objectif, périmètre, contexte) sans commandes/config.
- Les sections techniques doivent contenir des placeholders [[SEC_X_Y]].
- Les tableaux doivent être placés avec [[Excel: NOM_TABLEAU]].
- Les figures/diagrammes doivent être placés avec [[insérer NOM_ELEMENT]].

SORTIE
Retourner uniquement un JSON strictement valide conforme au schéma fourni.
""".strip()


SYSTEM_PROMPT_QUESTIONS = """
Tu génères des questions de formulaire pour une section unique.

RÈGLES STRICTES
- Types autorisés: single_choice, multi_choice uniquement.
- Types interdits: boolean, table, text, number, date, validation, default, placeholder, show_if.
- Générer entre 0 et 8 questions maximum.
- Si le contexte est vide, ambigu ou insuffisant, retourner questions=[].
- Chaque question doit être explicitement justifiée par une information présente dans le contexte fourni.
- Si une information n'est pas présente dans le contexte, ne pas poser la question correspondante.
- Mieux vaut 0 question qu'une question générique non justifiée par le contexte.
- Chaque question doit contenir 2 à 10 options dans choices, extraites/reformulées à partir du contexte.
- Produire aussi un champ purpose (2 à 3 phrases) décrivant l'objectif de la section et ce que le texte final doit couvrir, sans commandes ni configuration brute.

SORTIE
Retourner uniquement un JSON strictement valide conforme au schéma fourni.
""".strip()


def _looks_like_refusal(text: str) -> bool:
    if not text:
        return False
    snippet = text.lower()
    refusal_terms = [
        "i'm sorry",
        "i cannot",
        "i can't assist",
        "cannot assist",
        "can't assist",
        "je ne peux pas",
        "je suis désolé",
        "je suis desole",
        "désolé",
        "desole",
        "refus",
        "policy",
    ]
    return any(term in snippet for term in refusal_terms)


def _looks_like_timeout(exc: Exception) -> bool:
    message = str(exc).lower()
    return "timeout" in message or "timed out" in message


def _strip_number_prefix(title: str) -> Tuple[str, str]:
    if not title:
        return "", ""
    cleaned = title.strip()
    match = re.match(
        r"^(?P<number>(?:[IVXLCDM]+|[A-Z]|\d+)(?:\.(?:[IVXLCDM]+|[A-Z]|\d+))*)\s+(?P<title>.+)$",
        cleaned,
    )
    if match:
        return match.group("number"), match.group("title").strip()
    return "", cleaned


def _detect_numbering_pattern(titles: List[str]) -> str:
    counts = {"decimal": 0, "roman": 0, "alphanumeric": 0}
    for title in titles:
        number, _ = _strip_number_prefix(title)
        if not number:
            continue
        if re.match(r"^\d+(\.\d+)*$", number):
            counts["decimal"] += 1
        elif re.match(r"^[IVXLCDM]+(\.[IVXLCDM]+)*$", number, re.IGNORECASE):
            counts["roman"] += 1
        elif re.match(r"^[A-Z](\.[A-Z]|\.\d+)*$", number):
            counts["alphanumeric"] += 1
    if not any(counts.values()):
        return "decimal"
    return max(counts, key=counts.get)


def _detect_title_case(titles: List[str]) -> str:
    if not titles:
        return "sentence"
    upper = 0
    title_case = 0
    sentence = 0
    for raw in titles:
        _, title = _strip_number_prefix(raw)
        letters = [c for c in title if c.isalpha()]
        if not letters:
            continue
        upper_ratio = sum(c.isupper() for c in letters) / len(letters)
        if upper_ratio > 0.7:
            upper += 1
            continue
        words = [w for w in re.split(r"\s+", title) if w]
        if words and all(word[0].isupper() for word in words if word[0].isalpha()):
            title_case += 1
        else:
            sentence += 1
    if upper >= title_case and upper >= sentence:
        return "UPPER"
    if title_case >= sentence:
        return "title"
    return "sentence"


def _extract_style_map_from_docx(path: str) -> Dict[int, str]:
    doc = Document(path)
    style_map: Dict[int, str] = {}
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style_name = getattr(getattr(para, "style", None), "name", "") or ""
        level_match = re.match(r"(Heading|Titre)\s+(\d+)", style_name)
        if level_match:
            level = int(level_match.group(2))
            if level not in style_map:
                style_map[level] = style_name
    return style_map


def extract_style_guide_from_docs(paths: Iterable[str], headings: List[str]) -> StyleGuide:
    style_map: Dict[int, str] = {}
    for path in paths:
        if path.lower().endswith(".docx"):
            for level, style_name in _extract_style_map_from_docx(path).items():
                if level not in style_map:
                    style_map[level] = style_name
    numbering_pattern = _detect_numbering_pattern(headings)
    title_case = _detect_title_case(headings)
    levels = sorted({max(1, min(4, _derive_level_from_title(h))) for h in headings}) or [1, 2, 3, 4]
    detected_examples = headings[:8]
    return StyleGuide(
        numbering_pattern=numbering_pattern,
        title_case=title_case,
        indentation=levels,
        detected_examples=detected_examples,
    )


def _compress_text(text: str, limit: int) -> str:
    if len(text) <= limit:
        return text
    chunk_size = max(1, limit // SNIPPET_CHUNKS)
    parts = []
    parts.append(text[:chunk_size])
    if SNIPPET_CHUNKS >= 3:
        mid_start = max(0, len(text) // 2 - chunk_size // 2)
        parts.append(text[mid_start : mid_start + chunk_size])
    parts.append(text[-chunk_size:])
    return SNIPPET_SEPARATOR.join(parts)


def _split_text(text: str, chunk_size: int) -> List[str]:
    if not text:
        return [""]
    chunks: List[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + chunk_size)
        chunks.append(text[start:end])
        start = end
    return chunks[:MAX_CHUNKS]




def build_outline_prompt(
    redacted_text: str,
    template_type: str | None,
    headings: List[str],
    style_guide: StyleGuide,
    chunk_index: int = 1,
    chunk_total: int = 1,
) -> str:
    template_label = template_type or "LLD/HLD"
    clipped_headings = headings[:MAX_HEADING_COUNT]
    headings_text = "\n".join(f"- {title}" for title in clipped_headings) or "- (aucun titre extrait)"
    clipped_text = _compress_text(redacted_text, OUTLINE_CONTEXT_CHARS)
    style_guide_text = json.dumps(style_guide.model_dump(), ensure_ascii=False, indent=2)
    return f"""
CONTEXTE
Tu dois générer un template documentaire {template_label} et un formulaire basé sur des documents d'ingénierie.
Cette requête traite un EXTRACT {chunk_index}/{chunk_total} du document total.

STYLE GUIDE (MUST FOLLOW)
{style_guide_text}

ATTENDUS STRUCTURE
- Hiérarchie détaillée H1/H2/H3/H4.
- Respecter STRICTEMENT la numérotation, la casse et le style des titres du guide ci-dessus.
- Pour chaque titre, fournir: level, number, title, style_name (ex: Heading 1/2 ou style custom détecté).
- Sections typiques: objectifs, périmètre, architecture existante/cible, BOM, connectivité,
  plan d'adressage, VLAN, routage, matrice de flux, HA, supervision, tests.
- Marqueurs obligatoires:
  - [[SEC_X_Y]] pour sections à remplir.
  - [[Excel: NOM_TABLEAU]] pour tableaux (BOM, VLAN, routage, flux, connectivité, etc.).
  - [[insérer NOM_ELEMENT]] pour figures (topologie logique/physique, schéma HA, etc.).
- Tous les titres fournis dans la liste "TITRES EXTRAITS" doivent apparaître dans l'outline.
- Pour les grands titres (H1/H2), conserver exactement le libellé source (casse incluse), sans reformulation.
- Ne préfixe jamais les titres par un slug, un id, ou un texte technique. Le champ title doit être exactement le titre lisible du document.
- Interdit de préfixer les sous-titres par le titre parent (pas de namespace type "smc server.1").
- Le contenu doit être meilleur que la source: plus complet, plus clair, plus structuré, LLD-ready.
- Si l'information manque, créer une section "À renseigner" avec [[SEC_*]] et questions associées.

ATTENDU FORMULAIRE
- Le champ form.sections doit être vide (les questions seront générées dans un second passage).

SCHÉMA JSON
{_schema_json(AITemplateOutput.model_json_schema())}

IMPORTANT: sortie JSON stricte uniquement, aucun texte ou markdown hors JSON.

TITRES EXTRAITS (à inclure):
{headings_text}

CONTENU SOURCÉ (déjà expurgé):
{clipped_text}
""".strip()


def build_questions_prompt(
    sec_id: str,
    section_title: str,
    context_pack: str,
    style_guide: StyleGuide,
    doc_type: str | None,
    techno: str | None,
    extra_instructions: str | None = None,
) -> str:
    style_guide_text = json.dumps(style_guide.model_dump(), ensure_ascii=False, indent=2)
    doc_type_label = doc_type or "ingénierie"
    techno_label = techno or "N/A"
    schema_payload = SectionQuestionsOutput.model_json_schema()
    clipped_context = _compress_text(context_pack, QUESTION_CONTEXT_CHARS)

    lines = [
        "SECTION CIBLE",
        f"- sec_id: {sec_id}",
        f"- titre: {section_title}",
        f"- doc_type: {doc_type_label}",
        f"- techno: {techno_label}",
        "",
        "STYLE GUIDE",
        style_guide_text,
        "",
        "CONTEXTE DE LA SECTION (PARAGRAPHES + TABLEAUX)",
        clipped_context,
        "",
        "EXIGENCES FORMULAIRE",
        "- Types autorisés: single_choice, multi_choice uniquement.",
        "- Générer entre 2 et 5 questions (inclus).",
        "- Si contexte exploitable: minimum 2 questions obligatoires.",
        "- Chaque question doit être basée sur le contexte, jamais sur le titre seul.",
        "- choices: 2 à 10 options par question.",
        "- Pas de validation/default/show_if/placeholder.",
        '- Interdit absolu de réutiliser le bloc générique domaines/schéma/VLAN/HA/BOM.',
    ]

    if extra_instructions:
        lines.extend(["", extra_instructions.strip()])

    lines.extend(
        [
            "",
            "SCHÉMA JSON",
            _schema_json(schema_payload),
            "",
            "Réponds UNIQUEMENT avec un JSON strict conforme au schéma.",
        ]
    )

    return "\n".join(lines).strip()


def _normalize_outline_ids(outline: List[TemplateOutlineItem]) -> List[TemplateOutlineItem]:
    for idx, item in enumerate(outline, start=1):
        if not str(item.id or "").startswith("SEC_"):
            item.id = f"SEC_{idx}_1"
    return outline


def _outline_is_valid(outline: List[TemplateOutlineItem]) -> bool:
    if not outline:
        return False
    return any((item.title or "").strip() for item in outline)


def _derive_level_from_title(title: str) -> int:
    number, _ = _strip_number_prefix(title)
    match = re.match(r"^(\d+|[IVXLCDM]+|[A-Z])(?:\.(\d+|[IVXLCDM]+|[A-Z]))*", number.strip())
    if not match:
        return 2
    parts = number.split(".") if number else []
    return max(1, min(4, len(parts) if parts else 1))


def _style_name_for_level(level: int, style_map: Dict[int, str] | None = None) -> str:
    if style_map and style_map.get(level):
        return style_map[level]
    return f"Heading {level}"


def _augment_outline_with_headings(
    outline: List[TemplateOutlineItem],
    headings: List[str],
    style_map: Dict[int, str] | None = None,
    heading_levels: Dict[str, int] | None = None,
) -> List[TemplateOutlineItem]:
    existing = {_normalize_heading_key(item.title) for item in outline if item.title}
    next_index = len(outline) + 1
    for heading in headings:
        _, stripped = _strip_number_prefix(heading)
        key = _normalize_heading_key(stripped or heading)
        if not key or key in existing:
            continue
        level = heading_levels.get(key) if heading_levels else None
        level = level or _derive_level_from_title(heading)
        number, title = _strip_number_prefix(heading)
        outline.append(
            TemplateOutlineItem(
                id=f"SEC_{next_index}_1",
                level=level,
                number=number,
                title=title or heading,
                style_name=_style_name_for_level(level, style_map),
                content="",
                markers=[],
            )
        )
        existing.add(key)
        next_index += 1
    return outline


def _merge_outputs(outputs: List[AITemplateOutput]) -> AITemplateOutput:
    merged_outline: List[TemplateOutlineItem] = []
    title_seen: set[str] = set()
    placeholders: List[dict] = []
    excel_tables: List[dict] = []
    insert_items: List[dict] = []
    form_sections: dict[str, dict] = {}
    style_guide: StyleGuide | None = None

    for output in outputs:
        if style_guide is None:
            style_guide = output.style_guide
        for item in output.template_outline:
            key = _normalize_heading_key(item.title or "")
            if key and key not in title_seen:
                merged_outline.append(item)
                title_seen.add(key)
        for placeholder in output.placeholders:
            placeholders.append(placeholder.model_dump() if hasattr(placeholder, "model_dump") else placeholder)
        for table in output.excel_tables:
            excel_tables.append(table.model_dump() if hasattr(table, "model_dump") else table)
        for insert_item in output.insert_items:
            insert_items.append(insert_item.model_dump() if hasattr(insert_item, "model_dump") else insert_item)
        for section in output.form.sections:
            sec_id = section.sec_id
            if sec_id not in form_sections:
                form_sections[sec_id] = {
                    "sec_id": sec_id,
                    "purpose": section.purpose,
                    "questions": [],
                }
            elif section.purpose and not form_sections[sec_id].get("purpose"):
                form_sections[sec_id]["purpose"] = section.purpose
            for question in section.questions:
                form_sections[sec_id]["questions"].append(
                    question.model_dump() if hasattr(question, "model_dump") else question
                )

    return AITemplateOutput(
        style_guide=style_guide or outputs[0].style_guide,
        template_outline=merged_outline,
        placeholders=placeholders,
        excel_tables=excel_tables,
        insert_items=insert_items,
        form={"sections": list(form_sections.values())},
    )


def _is_placeholder_line(line: str) -> bool:
    candidate = (line or "").strip()
    if not candidate:
        return False
    return bool(
        SEC_PLACEHOLDER_RE.match(candidate)
        or EXCEL_PLACEHOLDER_RE.match(candidate)
        or INSERT_PLACEHOLDER_RE.match(candidate)
    )


def _extract_descriptive_text(content: str) -> str:
    lines = [ln.strip() for ln in (content or "").splitlines() if ln.strip()]
    non_placeholder_lines = [ln for ln in lines if not _is_placeholder_line(ln)]
    if non_placeholder_lines:
        return "\n".join(non_placeholder_lines).strip()
    without_placeholders = ANY_PLACEHOLDER_RE.sub("", content or "")
    cleaned = re.sub(r"\s+", " ", without_placeholders).strip()
    return cleaned


def _extract_effective_section_text(content: str) -> str:
    lines = [ln.strip() for ln in (content or "").splitlines()]
    kept_lines: List[str] = []
    for line in lines:
        if not line:
            continue
        if _is_placeholder_line(line):
            continue
        without_placeholders = ANY_PLACEHOLDER_RE.sub("", line)
        normalized = re.sub(r"\s+", " ", without_placeholders).strip()
        if normalized:
            kept_lines.append(normalized)
    return "\n".join(kept_lines).strip()


def _has_meaningful_section_text(content: str) -> bool:
    effective_text = _extract_effective_section_text(content)
    if len(effective_text.strip()) < MEANINGFUL_SECTION_MIN_CHARS:
        return False
    words = [w for w in re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ0-9]{2,}", effective_text) if not w.isdigit()]
    return len(words) >= MEANINGFUL_SECTION_MIN_WORDS


def normalize_template_placeholders(
    outline: List[TemplateOutlineItem],
) -> Tuple[List[TemplateOutlineItem], List[Dict[str, Any]]]:
    excel_keywords = [
        "bom", "nomenclature", "vlan", "routage", "routing", "adressage",
        "matrice de flux", "flux", "connectivité", "connectivite", "tableau", "table",
    ]
    insert_keywords = ["topologie", "architecture", "schéma", "schema", "figure", "diagramme", "photo"]
    strip_logs: List[Dict[str, Any]] = []

    for item in outline:
        sec_id = item.id or ""
        old_content = item.content or ""
        old_len = len(old_content)

        marker_pool = list(item.markers or [])
        marker_pool.extend(ANY_PLACEHOLDER_RE.findall(old_content))
        marker_pool = [m.strip() for m in marker_pool if (m or "").strip()]

        descriptive_text = _extract_descriptive_text(old_content)
        item.descriptive_text = descriptive_text or None

        sec_marker = f"[[{sec_id}]]"
        excel_marker = next((m for m in marker_pool if EXCEL_PLACEHOLDER_RE.match(m)), None)
        insert_marker = next((m for m in marker_pool if INSERT_PLACEHOLDER_RE.match(m)), None)

        lower_title = (item.title or "").lower()
        if not excel_marker and any(k in lower_title for k in excel_keywords):
            table_name = (item.title or "Tableau").replace(":", "").strip() or "Tableau"
            excel_marker = f"[[Excel: {table_name}]]"
        if not insert_marker and any(k in lower_title for k in insert_keywords):
            element_name = (item.title or "Élément").replace(":", "").strip() or "Élément"
            insert_marker = f"[[insérer {element_name}]]"

        strict_lines = [sec_marker]
        if excel_marker:
            strict_lines.append(excel_marker)
        if insert_marker:
            strict_lines.append(insert_marker)
        new_content = "\n".join(strict_lines)

        item.markers = strict_lines.copy()
        item.content = new_content

        if old_content.strip() != new_content.strip() or descriptive_text:
            strip_logs.append(
                {
                    "sec_id": sec_id,
                    "old_len": old_len,
                    "new_content": new_content,
                }
            )

    return outline, strip_logs



def _merge_section_purpose(base_purpose: str | None, descriptive_text: str | None) -> str | None:
    base = (base_purpose or "").strip()
    desc = (descriptive_text or "").strip()
    if desc and base:
        if desc in base:
            return base
        return f"{desc}\n\n{base}"
    return desc or base or None


def _ensure_outline_style_fields(
    outline: List[TemplateOutlineItem],
    style_map: Dict[int, str] | None = None,
    heading_levels: Dict[str, int] | None = None,
    heading_numbers: Dict[str, str] | None = None,
) -> List[TemplateOutlineItem]:
    for item in outline:
        key = _normalize_heading_key(item.title or "")
        if heading_levels and key in heading_levels:
            item.level = heading_levels[key]
        level = max(1, min(4, item.level))
        if not item.style_name:
            item.style_name = _style_name_for_level(level, style_map)
        if heading_numbers and key in heading_numbers:
            item.number = heading_numbers[key]
        if not item.number:
            number, title = _strip_number_prefix(item.title or "")
            if number:
                item.number = number
                item.title = title or item.title
        if item.number is None:
            item.number = ""
    return outline


def _fallback_output_from_headings(headings: List[str], style_guide: StyleGuide) -> AITemplateOutput:
    default_business_headings = [
        "1 Objectifs",
        "2 Contexte",
        "3 Périmètre",
        "4 Existant",
        "5 Architecture cible",
        "6 Plan d'adressage",
        "7 VLAN",
        "8 Routage",
        "9 Sécurité",
        "10 Supervision",
        "11 Migration",
        "12 Tests & validation",
    ]
    cleaned_headings = [h.strip() for h in headings if _is_meaningful_heading(h)]
    if cleaned_headings:
        noisy_count = sum(1 for h in headings if _is_hardware_code_heading(h) or _is_sec_placeholder_heading(h))
        noise_ratio = noisy_count / max(1, len(headings))
        base_headings = cleaned_headings if noise_ratio <= 0.40 else default_business_headings
    else:
        base_headings = default_business_headings
    outline: List[TemplateOutlineItem] = []
    placeholders = []
    form_sections = []
    for idx, heading in enumerate(base_headings, start=1):
        level = _derive_level_from_title(heading)
        number, title = _strip_number_prefix(heading)
        sec_id = f"SEC_{idx}_1"
        outline.append(
            TemplateOutlineItem(
                id=sec_id,
                level=level,
                number=number,
                title=title or heading,
                style_name=_style_name_for_level(level),
                content=f"[[{sec_id}]]",
                markers=[f"[[{sec_id}]]"],
            )
        )
        placeholders.append(
            {
                "sec_id": sec_id,
                "description": heading,
                "recommended_question_types": ["single_choice"],
                "intent": f"Informations requises pour {heading}",
                "example": None,
                "validation_rules": None,
                "placeholder": sec_id,
            }
        )
        form_sections.append(
            {
                "sec_id": sec_id,
                "purpose": None,
                "questions": [],
            }
        )
    return AITemplateOutput(
        style_guide=style_guide,
        template_outline=outline,
        placeholders=placeholders,
        excel_tables=[],
        insert_items=[],
        form={"sections": form_sections},
    )


def _enrich_form_sections(
    output: AITemplateOutput,
) -> AITemplateOutput:
    return output


def _requested_section_count(prompt: str) -> int:
    return len(set(re.findall(r"SEC_[A-Za-z0-9_]+", prompt or "")))


def _effective_system_prompt(system_prompt: str) -> str:
    if os.getenv("OPENAI_MINIMAL_SYSTEM_PROMPT", "0").lower() in {"1", "true", "yes", "on"}:
        return "Tu es un assistant. Réponds uniquement avec un JSON strict conforme au schéma fourni."
    return system_prompt


def _log_openai_error(exc: Exception, fallback_raw: str = "") -> None:
    raw_candidate = fallback_raw
    for attr in ("response", "body", "message", "last_response"):
        value = getattr(exc, attr, None)
        if value:
            raw_candidate = str(value)
            break
    logger.exception(
        "[OPENAI_ERROR] type=%s error=%s raw=%s",
        type(exc).__name__,
        str(exc),
        (raw_candidate or "")[:1000],
    )


def _call_openai_json_schema(
    prompt: str,
    schema_payload: Dict[str, Any],
    model: str,
    system_prompt: str,
    max_output_tokens: int,
) -> str:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY missing in environment")
    client = OpenAI(api_key=api_key, timeout=120)
    effective_system_prompt = _effective_system_prompt(system_prompt)

    request_payload = {
        "model": model,
        "max_output_tokens": max_output_tokens,
        "response_format": {
            "type": "json_schema",
            "json_schema": {
                "name": "ai_template_output",
                "schema": schema_payload,
                "strict": True,
            },
        },
        "input": [
            {"role": "system", "content": effective_system_prompt},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.1,
    }
    logger.info(
        "[OPENAI_REQUEST] model=%s payload_size=%s requested_sections=%s minimal_system_prompt=%s",
        model,
        len(json.dumps(request_payload, ensure_ascii=False)),
        _requested_section_count(prompt),
        effective_system_prompt != system_prompt,
    )

    try:
        response = client.responses.create(**request_payload)
        output_text = response.output_text or ""
        logger.info("[OPENAI_RAW_RESPONSE] %s", output_text[:500])
        return output_text
    except TypeError:
        fallback_payload = {
            "model": model,
            "max_output_tokens": max_output_tokens,
            "input": [
                {
                    "role": "system",
                    "content": (
                        f"{effective_system_prompt}\n\n"
                        "Tu dois répondre UNIQUEMENT avec un JSON strict conforme au schéma fourni."
                    ),
                },
                {"role": "user", "content": f"SCHÉMA JSON:\n{json.dumps(schema_payload, ensure_ascii=False)}"},
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.1,
        }
        logger.info(
            "[OPENAI_REQUEST] model=%s payload_size=%s requested_sections=%s minimal_system_prompt=%s mode=fallback_no_response_format",
            model,
            len(json.dumps(fallback_payload, ensure_ascii=False)),
            _requested_section_count(prompt),
            effective_system_prompt != system_prompt,
        )
        try:
            response = client.responses.create(**fallback_payload)
            output_text = response.output_text or ""
            logger.info("[OPENAI_RAW_RESPONSE] %s", output_text[:500])
            return output_text
        except Exception as exc:
            _log_openai_error(exc)
            raise
    except Exception as exc:
        _log_openai_error(exc)
        raise


def call_llm_and_parse_json_schema(prompt: str, provider: LLMProvider) -> AITemplateOutput:
    try:
        try:
            raw = call_llm_json(
                provider=provider,
                prompt=prompt,
                system_prompt=SYSTEM_PROMPT_OUTLINE,
                json_schema=AITemplateOutput.model_json_schema(),
                model_key="outline_model",
                max_output_tokens=10000,
            )
        except Exception as api_exc:
            if _looks_like_timeout(api_exc):
                raise RuntimeError(f"{provider.value} request timed out.") from api_exc
            raise RuntimeError(str(api_exc)) from api_exc
        if _looks_like_refusal(raw):
            logger.warning("LLM_REFUSAL_DETECTED stage=outline_initial provider=%s snippet=%s", provider.value, raw[:200])
            raise RuntimeError(f"{provider.value} refusal: {raw[:200]}")
        data = _extract_json(raw)
        return AITemplateOutput.model_validate(data)
    except Exception as exc:
        repair_prompt = (
            "JSON ONLY. Respecte strictement le schéma. Aucun texte hors JSON.\n\n" + prompt
        )
        try:
            repaired = call_llm_json(
                provider=provider,
                prompt=repair_prompt,
                system_prompt=SYSTEM_PROMPT_OUTLINE,
                json_schema=AITemplateOutput.model_json_schema(),
                model_key="outline_model",
                max_output_tokens=10000,
            )
        except Exception as api_exc:
            if _looks_like_timeout(api_exc):
                raise RuntimeError(f"{provider.value} request timed out.") from api_exc
            raise RuntimeError(str(api_exc)) from api_exc
        if _looks_like_refusal(repaired):
            logger.warning("LLM_REFUSAL_DETECTED stage=outline_repair provider=%s snippet=%s", provider.value, repaired[:200])
            raise RuntimeError(f"{provider.value} refusal: {repaired[:200]}")
        try:
            data = _extract_json(repaired)
            return AITemplateOutput.model_validate(data)
        except Exception as repair_exc:
            raise RuntimeError(
                "JSON invalide après une relance obligatoire."
            ) from (repair_exc or exc)


def build_docx_from_outline(
    outline: List[TemplateOutlineItem],
    output_path: Path,
    template_path: str | None = None,
    template_has_numbering: bool = False,
) -> str:
    if template_path and Path(template_path).exists():
        doc = Document(template_path)
        body = doc.element.body
        for element in list(body):
            body.remove(element)
    else:
        doc = Document()
    has_numbers = any((item.number or "").strip() for item in outline)
    for item in outline:
        level = max(1, min(4, item.level))
        style_name = item.style_name or _style_name_for_level(level)
        title = item.title or ""
        title = re.sub(r"^(\d+(?:\.\d+)*)([A-Za-zÀ-ÿ])", r"\1 \2", title)
        number = (item.number or "").strip()
        if not template_has_numbering and number and not title.lstrip().startswith(number):
            title = f"{number} {title}".strip()
        paragraph = doc.add_paragraph(title)
        if style_name in doc.styles:
            paragraph.style = style_name
        else:
            heading_style = _style_name_for_level(level)
            if heading_style in doc.styles:
                paragraph.style = heading_style
        if item.content:
            doc.add_paragraph(item.content)
    if not template_has_numbering and not has_numbers:
        renumber_headings(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return str(output_path)


def _append_docx(target: Document, source: Document) -> None:
    for element in source.element.body:
        target.element.body.append(element)


def _merge_cover_template(cover_path: Path, template_path: Path) -> None:
    cover_doc = Document(str(cover_path))
    template_doc = Document(str(template_path))
    cover_doc.add_page_break()
    _append_docx(cover_doc, template_doc)
    cover_doc.save(str(template_path))


def _help_text(question: Dict[str, Any]) -> str:
    parts = []
    if question.get("help_text"):
        parts.append(str(question["help_text"]))
    if question.get("intent"):
        parts.append(f"Intent: {question['intent']}")
    if question.get("example"):
        parts.append(f"Exemple: {question['example']}")
    validation = question.get("validation") or {}
    if validation:
        parts.append("Validation: " + json.dumps(validation, ensure_ascii=False))
    if question.get("default") is not None:
        parts.append(f"Valeur par défaut: {question['default']}")
    return "\n".join(parts) if parts else None


def _form_name_for_techno(techno_name: str) -> str:
    base = techno_name.strip() or "Formulaire"
    return f"Formulaire {base}"


def _normalize_question_type(qtype: str) -> str:
    if not qtype:
        return "single_choice"
    lowered = qtype.strip().lower()
    if lowered in {"select", "single_select", "single"}:
        return "single_choice"
    if lowered in {"multi_select", "multi"}:
        return "multi_choice"
    return lowered


def _snake_case(value: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9]+", "_", value.strip().lower())
    cleaned = re.sub(r"_+", "_", cleaned).strip("_")
    return cleaned or "question"


def _normalize_question_payload(questions: List[FormQuestionSpec], sec_id: str) -> List[FormQuestionSpec]:
    for question in questions:
        if not question.key and question.label:
            question.key = _snake_case(question.label)
    return questions


def _detect_fallback_topic(title: str) -> str:
    normalized = (title or "").lower()
    if any(token in normalized for token in ["wifi", "wlan", "ssid", "capwap", "ewc", "access point"]):
        return "wifi"
    if any(token in normalized for token in ["wan", "mpls", "vpls", "internet", "liaison"]):
        return "wan"
    if any(token in normalized for token in ["ngfw", "pare-feu", "firewall", "nac", "ise", "sécurité", "security", "ips"]):
        return "security"
    if any(token in normalized for token in ["adressage", "ip", "routage", "ospf", "bgp", "gateway", "passerelle"]):
        return "routing"
    if any(token in normalized for token in ["vlan", "spanning", "ether-channel", "couche 2", "lan", "campus", "stack"]):
        return "switching"
    if any(token in normalized for token in ["bom", "équipement", "hardware", "rack", "câblage", "hostname", "nomenclature"]):
        return "inventory"
    return "generic"


def _extract_context_choices(context_pack: str, max_items: int = 10) -> List[str]:
    raw = context_pack or ""
    tokens = re.split(r"[\n,;|:/()\[\]\-]+", raw)
    stopwords = {
        "section",
        "contexte",
        "contenu",
        "tableaux",
        "pertinents",
        "paragraphe",
        "paragraphes",
        "document",
        "niveau",
        "titre",
        "phase",
        "site",
        "sites",
        "projet",
        "requis",
    }
    choices: List[str] = []
    seen = set()
    for token in tokens:
        t = re.sub(r"\s+", " ", token).strip(" .\t\r\n")
        if len(t) < 3 or len(t) > 40:
            continue
        if re.fullmatch(r"[0-9.]+", t):
            continue
        key = t.lower()
        if key in stopwords or key.startswith("sec_"):
            continue
        if key not in seen:
            seen.add(key)
            choices.append(t)
            if len(choices) >= max_items:
                break
    return choices


def _contextual_fallback_questions(sec_id: str, title_clean: str, context_pack: str, is_editorial: bool) -> List[FormQuestionSpec]:
    choices = _extract_context_choices(context_pack, max_items=12)
    if len(choices) < 4:
        return []

    short_choices = choices[:8]
    second_choices = choices[2:10] if len(choices) > 6 else choices[:6]

    questions: List[FormQuestionSpec] = [
        FormQuestionSpec(
            key=_snake_case(f"elements_section_{title_clean}"),
            label=f"Quels éléments sont explicitement cités dans la section \"{title_clean}\" ?",
            type="multi_choice",
            choices=short_choices,
            required=True,
            help_text=None,
            intent="Forcer une collecte alignée sur les termes réels du paragraphe.",
            example=short_choices[:2],
            validation={"kind": "enum", "rules": {"min_items": 1, "max_items": min(8, len(short_choices))}},
            default=None,
            placeholder=sec_id,
            show_if=None,
        ),
        FormQuestionSpec(
            key=_snake_case(f"priorite_section_{title_clean}"),
            label="Quel élément est prioritaire pour cette section ?",
            type="single_choice",
            choices=short_choices,
            required=True,
            help_text=None,
            intent="Identifier la priorité technique principale selon le contenu source.",
            example=short_choices[0],
            validation={"kind": "enum", "rules": {"allowed": short_choices}},
            default=None,
            placeholder=sec_id,
            show_if=None,
        ),
        FormQuestionSpec(
            key=_snake_case(f"sous_elements_section_{title_clean}"),
            label="Quels éléments complémentaires doivent être pris en compte ?",
            type="multi_choice",
            choices=second_choices,
            required=True,
            help_text=None,
            intent="Capturer les sous-éléments réellement présents dans le contexte de section.",
            example=second_choices[:2],
            validation={"kind": "enum", "rules": {"min_items": 1, "max_items": min(8, len(second_choices))}},
            default=None,
            placeholder=sec_id,
            show_if=None,
        ),
    ]

    if not is_editorial:
        questions.append(
            FormQuestionSpec(
                key=_snake_case(f"ha_section_{title_clean}"),
                label="Une redondance / haute disponibilité est-elle mentionnée ou requise ?",
                type="boolean",
                required=True,
                choices=None,
                help_text=None,
                intent="Valider l'exigence de résilience si elle apparaît dans le contexte.",
                example=True,
                validation={"kind": "enum", "rules": {"allowed": [True, False]}},
                default=None,
                placeholder=sec_id,
                show_if=None,
            )
        )

    return questions


def _fallback_questions_for_section(sec_id: str, title: str, context_pack: str | None = None) -> List[FormQuestionSpec]:
    effective_text = _extract_effective_section_text(context_pack or "")
    title_clean = re.sub(r"\s+", " ", (title or "section").strip())
    base_choices = [
        "Éléments déjà définis",
        "Éléments partiellement définis",
        "Éléments à préciser",
        "Éléments non mentionnés",
    ]
    second_choices = [
        "Contraintes techniques",
        "Exigences de sécurité",
        "Impacts d'intégration",
        "Dépendances externes",
        "Critères de validation",
    ]
    snippet = effective_text[:220]

    return [
        FormQuestionSpec(
            key=_snake_case(f"elements_cles_{sec_id}"),
            label=f"Quels sont les éléments clés explicitement décrits dans la section '{title_clean}' ?",
            type="multi_choice",
            choices=base_choices,
            required=True,
            help_text=(f"Extrait utile: {snippet}" if snippet else None),
        ),
        FormQuestionSpec(
            key=_snake_case(f"exigences_{sec_id}"),
            label=f"Quelles exigences opérationnelles ou techniques doivent être retenues pour '{title_clean}' ?",
            type="multi_choice",
            choices=second_choices,
            required=True,
            help_text=(f"Basé sur le contenu de section ({len(effective_text)} caractères utiles)." if effective_text else None),
        ),
    ]


def generate_questions_for_section(
    sec_id: str,
    title: str,
    context_pack: str,
    style_guide: StyleGuide,
    doc_type: str | None,
    techno: str | None,
    extra_instructions: str | None = None,
    provider: LLMProvider = LLMProvider.OPENAI,
) -> SectionQuestionsOutput:
    effective_text = _extract_effective_section_text(context_pack or "")
    if not _has_meaningful_section_text(context_pack or ""):
        return SectionQuestionsOutput(
            sec_id=sec_id,
            section_title=title,
            purpose="But: contexte insuffisant/non exploitable, aucune question générée.",
            questions=[],
        )

    prompt = build_questions_prompt(
        sec_id=sec_id,
        section_title=title,
        context_pack=context_pack,
        style_guide=style_guide,
        doc_type=doc_type,
        techno=techno,
        extra_instructions=extra_instructions,
    )
    schema_payload = SectionQuestionsOutput.model_json_schema()
    last_exc: Exception | None = None

    for attempt in range(3):
        retry_triggered = attempt > 0
        try:
            raw = call_llm_json(
                provider=provider,
                prompt=prompt,
                system_prompt=SYSTEM_PROMPT_QUESTIONS,
                json_schema=schema_payload,
                model_key="form_model",
                max_output_tokens=8000,
            )
        except Exception as api_exc:
            if _looks_like_timeout(api_exc):
                last_exc = RuntimeError(f"{provider.value} request timed out.")
            else:
                last_exc = RuntimeError(str(api_exc))
            continue

        if _looks_like_refusal(raw):
            logger.warning("LLM_REFUSAL_DETECTED stage=questions_initial provider=%s sec_id=%s snippet=%s", provider.value, sec_id, raw[:200])
            last_exc = RuntimeError(f"{provider.value} refusal: {raw[:200]}")
            continue

        try:
            data = _sanitize_section_questions_payload(_extract_json(raw))
            output = SectionQuestionsOutput.model_validate(data)
        except Exception:
            repair_prompt = "JSON ONLY. Respecte strictement le schéma. Aucun texte hors JSON.\n\n" + prompt
            repaired = call_llm_json(
                provider=provider,
                prompt=repair_prompt,
                system_prompt=SYSTEM_PROMPT_QUESTIONS,
                json_schema=schema_payload,
                model_key="form_model",
                max_output_tokens=8000,
            )
            if _looks_like_refusal(repaired):
                logger.warning("LLM_REFUSAL_DETECTED stage=questions_repair provider=%s sec_id=%s snippet=%s", provider.value, sec_id, repaired[:200])
                last_exc = RuntimeError(f"{provider.value} refusal: {repaired[:200]}")
                continue
            data = _sanitize_section_questions_payload(_extract_json(repaired))
            output = SectionQuestionsOutput.model_validate(data)

        output.questions = _normalize_question_payload(output.questions, sec_id)[:MAX_SECTION_QUESTIONS]
        question_types = [((q.type or "unknown").strip().lower() or "unknown") for q in output.questions]
        logger.info(
            "[AI_OUTPUT] Section %r (%s) -> %s questions (%s)",
            title,
            sec_id,
            len(output.questions),
            ", ".join(question_types),
        )

        if len(output.questions) < MIN_SECTION_QUESTIONS:
            last_exc = RuntimeError(f"Questions insuffisantes: {len(output.questions)} < {MIN_SECTION_QUESTIONS}")
            if attempt == 0:
                prompt = build_questions_prompt(
                    sec_id=sec_id,
                    section_title=title,
                    context_pack=context_pack,
                    style_guide=style_guide,
                    doc_type=doc_type,
                    techno=techno,
                    extra_instructions=(
                        (extra_instructions or "")
                        + f"\nFALLBACK STRICT: Tu dois générer au moins {MIN_SECTION_QUESTIONS} questions et au plus {MAX_SECTION_QUESTIONS}."
                        + "\nInterdiction de retourner un tableau questions vide quand le contexte contient du texte exploitable."
                    ).strip(),
                )
                continue
            break

        try:
            enforce_question_quality(
                output.questions,
                is_editorial_section(title),
                effective_text,
            )
            logger.info(
                "[QUALITY_GATE] Section %r (%s) -> %s questions validées (retry=%s)",
                title,
                sec_id,
                len(output.questions),
                retry_triggered,
            )
            return output
        except Exception as quality_exc:
            last_exc = quality_exc
            prompt = build_questions_prompt(
                sec_id=sec_id,
                section_title=title,
                context_pack=context_pack,
                style_guide=style_guide,
                doc_type=doc_type,
                techno=techno,
                extra_instructions=(
                    (extra_instructions or "")
                    + "\nRÉVISION QUALITÉ OBLIGATOIRE: "
                    + str(quality_exc)
                    + "\n- respecter strictement les types autorisés (single_choice/multi_choice)."
                    + f"\n- {MIN_SECTION_QUESTIONS}..{MAX_SECTION_QUESTIONS} questions obligatoires."
                    + "\n- choices obligatoires (2..10)."
                ).strip(),
            )

    logger.warning("Question generation failed for sec_id=%s title=%r: %s", sec_id, title, last_exc)
    fallback_questions = _fallback_questions_for_section(sec_id=sec_id, title=title, context_pack=effective_text)
    return SectionQuestionsOutput(
        sec_id=sec_id,
        section_title=title,
        purpose="But: fallback déterministe activé suite à un échec IA.",
        questions=fallback_questions[:MAX_SECTION_QUESTIONS],
    )


def _question_stats(questions: List[FormQuestionSpec]) -> Dict[str, Any]:
    type_dist: Dict[str, int] = {}
    labels: List[str] = []
    for question in questions:
        qtype = (question.type or "").strip().lower() or "unknown"
        type_dist[qtype] = type_dist.get(qtype, 0) + 1
        if question.label:
            labels.append(question.label)
    return {"q_count": len(questions), "top_labels": labels[:3], "type_dist": type_dist}


def _question_signature(questions: List[FormQuestionSpec], limit: int = 6) -> Tuple[str, ...]:
    labels: List[str] = []
    for question in questions[:limit]:
        label = re.sub(r"\s+", " ", (question.label or "").strip().lower())
        if label:
            labels.append(label)
    return tuple(labels)


def _is_overly_repeated_signature(
    signature: Tuple[str, ...],
    seen_signatures: Dict[Tuple[str, ...], int],
) -> bool:
    if not signature:
        return False
    # Tolère une répétition occasionnelle, mais bloque les sets identiques en chaîne.
    return seen_signatures.get(signature, 0) >= 2


def _canonical_text(value: str) -> str:
    text = (value or "").strip().lower()
    text = text.replace("’", "'")
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    text = re.sub(r"[^a-z0-9\s]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _looks_like_legacy_generic_set(questions: List[FormQuestionSpec]) -> bool:
    labels = {
        _canonical_text(q.label or "")
        for q in questions
        if (q.label or "").strip()
    }
    if not labels:
        return False

    legacy_technical_fragments = [
        "quels domaines techniques sont dans le perimetre",
        "quel schema d architecture cible est attendu",
        "quels sites sont concernes",
        "le plan vlan detaille est il requis",
        "renseigner les vlan attendus id nom usage zone",
        "une architecture haute disponibilite est elle requise",
        "si ha quel mode de redondance failover doit etre applique",
        "quels controles de securite doivent etre integres",
        "une bom inventaire materiel est elle attendue",
        "renseigner les equipements attendus type modele quantite site",
    ]
    legacy_editorial_fragments = [
        "quel est le role principal du document pour cette section",
        "quelles technologies sont couvertes par cette section",
        "quel site ou entite est concernee",
        "le projet est il phase deploiement en etapes",
        "definir les phases et ce qui change a chaque phase",
    ]

    def _hit_count(fragments: List[str]) -> int:
        hits = 0
        for fragment in fragments:
            if any(fragment in label for label in labels):
                hits += 1
        return hits

    technical_hits = _hit_count(legacy_technical_fragments)
    editorial_hits = _hit_count(legacy_editorial_fragments)
    return technical_hits >= 4 or editorial_hits >= 3


def create_form_from_output(
    db: Session,
    techno: Techno,
    template_doc: TemplateDoc,
    output: AITemplateOutput,
    created_by: str | None,
    client_names: List[str],
) -> FormTemplate:
    create_started_at = time.perf_counter()
    flush_count = 0
    name = _form_name_for_techno(techno.name)
    last = (
        db.query(FormTemplate)
        .filter(FormTemplate.techno_id == techno.id, FormTemplate.name == name)
        .order_by(FormTemplate.version.desc())
        .first()
    )
    if last:
        db.query(FormTemplate).filter(
            FormTemplate.techno_id == techno.id, FormTemplate.name == name
        ).update({FormTemplate.is_active: False})
        form = FormTemplate(
            techno_id=techno.id,
            doc_id=template_doc.id,
            name=name,
            version=last.version + 1,
            is_active=True,
            parent_id=last.parent_id or last.id,
            created_by=created_by,
        )
    else:
        form = FormTemplate(
            techno_id=techno.id,
            doc_id=template_doc.id,
            name=name,
            version=1,
            is_active=True,
            created_by=created_by,
        )
    db.add(form)
    db.flush()
    flush_count += 1
    if not form.parent_id:
        form.parent_id = form.id
        db.flush()
        flush_count += 1

    outline_map = {item.id: item for item in output.template_outline}
    placeholder_map = {item.sec_id: item for item in output.placeholders}
    logger.info(
        "form_create_start form_id=%s section_count=%s",
        form.id,
        len(output.form.sections),
    )

    section_rows: List[FormSection] = []
    section_specs: List[FormSectionSpec] = []
    section_started_at: Dict[str, float] = {}

    for idx, section in enumerate(output.form.sections):
        section_started_at[section.sec_id] = time.perf_counter()
        outline_item = outline_map.get(section.sec_id)
        heading_level = outline_item.level if outline_item else 2
        heading_title = f"{outline_item.number} {outline_item.title}".strip() if outline_item else section.sec_id
        heading_title = replace_client_names(heading_title, client_names)
        placeholder_info = placeholder_map.get(section.sec_id)
        purpose_text = replace_client_names(getattr(section, "purpose", None), client_names)

        form_section = FormSection(
            form_id=form.id,
            sec_key=section.sec_id,
            heading_level=heading_level,
            heading_title=heading_title,
            order_index=idx,
            section_intent=getattr(placeholder_info, "intent", None) if placeholder_info else None,
            section_example=getattr(placeholder_info, "example", None) if placeholder_info else None,
            purpose_text=purpose_text,
        )
        section_rows.append(form_section)
        section_specs.append(section)

    if section_rows:
        db.add_all(section_rows)
        db.flush()
        flush_count += 1

    all_questions: List[FormQuestion] = []
    questions_by_sec_id: Dict[str, List[FormQuestion]] = {}
    options_count_by_sec_id: Dict[str, int] = {}

    for form_section, section in zip(section_rows, section_specs):
        first_question = section.questions[0] if section.questions else None
        logger.info(
            "form_section_input form_id=%s sec_key=%s heading_title=%s q_count=%s first_label=%s first_type=%s first_choices_count=%s first_validation_keys=%s",
            form.id,
            section.sec_id,
            form_section.heading_title,
            len(section.questions),
            getattr(first_question, "label", None),
            getattr(first_question, "type", None),
            len(getattr(first_question, "choices", []) or []),
            sorted(list((getattr(first_question, "validation", {}) or {}).keys())),
        )
        questions_for_section: List[FormQuestion] = []
        option_count = 0
        for q_idx, question in enumerate(section.questions):
            question_data = question.model_dump()
            help_text = question.help_text or _help_text(question_data)
            normalized_qtype = _normalize_question_type(question.type)
            questions_for_section.append(
                FormQuestion(
                    section_id=form_section.id,
                    label=replace_client_names(question.label, client_names),
                    qtype=normalized_qtype,
                    is_required=bool(question.required),
                    help_text=help_text,
                    order_index=q_idx,
                    placeholder_key=getattr(question, "placeholder", None),
                    question_key=question.key,
                    show_if_json=getattr(question, "show_if", None),
                )
            )
            option_count += len(question.choices or [])

        questions_by_sec_id[section.sec_id] = questions_for_section
        options_count_by_sec_id[section.sec_id] = option_count
        all_questions.extend(questions_for_section)

    questions_insert_started_at = time.perf_counter()
    if all_questions:
        db.add_all(all_questions)
        db.flush()
        flush_count += 1
    questions_insert_elapsed_ms = (time.perf_counter() - questions_insert_started_at) * 1000

    options_insert_started_at = time.perf_counter()
    all_options: List[FormOption] = []
    for section in section_specs:
        q_rows = questions_by_sec_id.get(section.sec_id, [])
        for form_question, question in zip(q_rows, section.questions):
            for o_idx, opt in enumerate(question.choices or []):
                sanitized_opt = replace_client_names(opt, client_names)
                all_options.append(
                    FormOption(
                        question_id=form_question.id,
                        label=sanitized_opt,
                        value=sanitized_opt,
                        order_index=o_idx,
                    )
                )
    if all_options:
        db.bulk_save_objects(all_options)
    options_insert_elapsed_ms = (time.perf_counter() - options_insert_started_at) * 1000

    total_questions_inserted = len(all_questions)
    total_options_inserted = len(all_options)

    for form_section, section in zip(section_rows, section_specs):
        sec_questions = len(questions_by_sec_id.get(section.sec_id, []))
        sec_options = options_count_by_sec_id.get(section.sec_id, 0)
        logger.info(
            "form_section_profile sec_id=%s section_id=%s q_count=%s opt_count=%s time_ms=%.2f",
            section.sec_id,
            form_section.id,
            sec_questions,
            sec_options,
            (time.perf_counter() - section_started_at[section.sec_id]) * 1000,
        )

    logger.info(
        "form_create_end form_id=%s section_count=%s total_questions=%s total_options=%s flush_count=%s commit_count=%s questions_insert_ms=%.2f options_insert_ms=%.2f elapsed_ms=%.2f",
        form.id,
        len(output.form.sections),
        total_questions_inserted,
        total_options_inserted,
        flush_count,
        0,
        questions_insert_elapsed_ms,
        options_insert_elapsed_ms,
        (time.perf_counter() - create_started_at) * 1000,
    )

    return form


def run_ai_template_job(db: Session, job_id: int, client_names: List[str]) -> None:
    job_started_at = time.perf_counter()
    job = db.query(AITemplateJob).filter(AITemplateJob.id == job_id).first()
    if not job:
        raise ValueError("Job introuvable")
    if job.status == AITemplateJobStatus.RUNNING:
        return
    job.status = AITemplateJobStatus.RUNNING
    provider = get_provider(job.llm_provider)
    job.progress = 5
    job.logs = (job.logs or "") + f"\nDémarrage du job. Provider LLM: {provider.value}."
    db.commit()

    try:
        files_meta = job.source_files or []
        paths = [item["stored_path"] for item in files_meta if item.get("stored_path")]
        extracted_text = extract_text_from_docs(paths)
        doc_tables = extract_tables_from_docs(paths)
        docx_sections: List[dict] = []
        for path in paths:
            if path.lower().endswith(".docx"):
                docx_sections.extend(extract_sections_from_docx_by_headings(path))
        headings = extract_headings_from_docs(paths)
        if len(headings) < 6:
            headings = extract_headings_from_text(extracted_text)
        headings = _filter_heading_candidates(headings)
        headings = [replace_client_names(h, client_names) for h in headings]
        if len(headings) < 6:
            headings = [
                "1 Objectif",
                "2 Contexte",
                "3 Périmètre",
                "4 Existant",
                "5 Cible",
                "6 VLAN",
                "7 Routage",
                "8 Sécurité",
                "9 Supervision",
                "10 Migration",
                "11 Tests",
            ]
        job.logs = (job.logs or "") + f"\nHeadings détectés: {len(headings)}."
        heading_levels = extract_heading_levels_from_docs(paths)
        heading_numbers = extract_heading_numbers_from_docs(paths)
        heading_titles = extract_heading_titles_from_docs(paths)
        seen = set()
        deduped_headings = []
        for title in headings:
            key = title.strip().lower()
            if key and key not in seen:
                seen.add(key)
                deduped_headings.append(title.strip())
        headings = deduped_headings
        heading_sample = headings[:10]
        job.logs = (job.logs or "") + f"\nHEADINGS_SOURCE_SAMPLE={heading_sample}"
        logger.info("HEADINGS_SOURCE_SAMPLE=%s", heading_sample)
        style_guide = extract_style_guide_from_docs(paths, headings)
        style_map: Dict[int, str] = {}
        for path in paths:
            if path.lower().endswith(".docx"):
                for level, style_name in _extract_style_map_from_docx(path).items():
                    if level not in style_map:
                        style_map[level] = style_name
        doc_tables = [replace_client_names(table, client_names) for table in doc_tables]
        docx_sections = deep_replace(docx_sections, client_names)
        redacted_text, report = redact_confidential(extracted_text, client_names)
        sanitized_text, sanitization_report = sanitize_network_text(redacted_text)
        sanitized_text = replace_client_names(sanitized_text, client_names)
        job.redaction_report = {
            "client_redaction": report,
            "network_sanitization": sanitization_report,
        }
        job.progress = 25
        job.logs = (job.logs or "") + "\nTexte extrait, expurgé et sanitizé."
        db.commit()

        if len(headings) > MAX_HEADING_COUNT:
            job.logs = (job.logs or "") + f"\nTitres tronqués: {len(headings)} -> {MAX_HEADING_COUNT}."
        chunks = _split_text(sanitized_text, MAX_PROMPT_CHARS)
        if len(chunks) > 1:
            job.logs = (job.logs or "") + f"\nTexte expurgé découpé en {len(chunks)} extraits."
        outputs: List[AITemplateOutput] = []
        fallback_output: AITemplateOutput | None = None
        started_at = time.monotonic()
        for idx, chunk in enumerate(chunks, start=1):
            outline_started_at = time.perf_counter()
            if time.monotonic() - started_at > MAX_JOB_SECONDS:
                fallback_output = _fallback_output_from_headings(headings, style_guide)
                job.logs = (job.logs or "") + "\nFallback déclenché (timeout IA)."
                logger.info(
                    "outline_openai_timing chunk=%s/%s elapsed_ms=%.2f chunk_chars=%s status=timeout_fallback",
                    idx,
                    len(chunks),
                    (time.perf_counter() - outline_started_at) * 1000,
                    len(chunk),
                )
                break
            prompt = build_outline_prompt(
                chunk,
                job.template_type,
                headings,
                style_guide,
                chunk_index=idx,
                chunk_total=len(chunks),
            )
            try:
                outputs.append(call_llm_and_parse_json_schema(prompt, provider=provider))
                logger.info(
                    "outline_openai_timing chunk=%s/%s elapsed_ms=%.2f chunk_chars=%s status=ok",
                    idx,
                    len(chunks),
                    (time.perf_counter() - outline_started_at) * 1000,
                    len(chunk),
                )
            except RuntimeError as exc:
                fallback_output = _fallback_output_from_headings(headings, style_guide)
                logger.info(
                    "outline_openai_timing chunk=%s/%s elapsed_ms=%.2f chunk_chars=%s status=openai_fallback",
                    idx,
                    len(chunks),
                    (time.perf_counter() - outline_started_at) * 1000,
                    len(chunk),
                )
                snippet = chunk[:200].replace("\n", " ")
                job.logs = (
                    (job.logs or "")
                    + f"\nFallback déclenché (OpenAI): {exc}"
                    + f"\nRefusal chunk={idx}/{len(chunks)} len={len(chunk)} snippet={snippet}"
                )
                break
        if fallback_output:
            output = fallback_output
        else:
            output = _merge_outputs(outputs)

        output.template_outline = _normalize_outline_ids(output.template_outline)
        output.template_outline = _clean_outline_titles(output.template_outline)
        output.template_outline = _align_outline_titles_with_source(output.template_outline, heading_titles)
        cleaned_outline_sample = [item.title for item in output.template_outline[:10]]
        job.logs = (job.logs or "") + f"\nOUTLINE_TITLES_AFTER_CLEAN_SAMPLE={cleaned_outline_sample}"
        logger.info("OUTLINE_TITLES_AFTER_CLEAN_SAMPLE=%s", cleaned_outline_sample)
        job.logs = (job.logs or "") + f"\nOutline après merge/normalize: {len(output.template_outline)}."
        output.template_outline = _augment_outline_with_headings(
            output.template_outline,
            headings,
            style_map,
            heading_levels,
        )
        output.template_outline = _align_outline_titles_with_source(output.template_outline, heading_titles)
        output.template_outline = _ensure_outline_style_fields(
            output.template_outline,
            style_map,
            heading_levels,
            heading_numbers,
        )
        if headings and not _outline_is_valid(output.template_outline):
            output.template_outline = _build_outline_from_headings(
                headings,
                heading_levels,
                heading_numbers,
                style_map,
            )
        output.template_outline, stripped_logs = normalize_template_placeholders(output.template_outline)
        for stripped in stripped_logs:
            log_line = (
                f"\nTEMPLATE_CONTENT_STRIPPED sec_id={stripped['sec_id']} "
                f"old_len={stripped['old_len']} new_content={stripped['new_content']!r}"
            )
            job.logs = (job.logs or "") + log_line
            logger.info(log_line.strip())

        for outline_item in output.template_outline:
            outline_item.title = replace_client_names(outline_item.title, client_names)
            outline_item.content = replace_client_names(outline_item.content, client_names)
        output.form.sections = []
        seen_question_signatures: Dict[Tuple[str, ...], int] = {}
        all_generated_labels: list[str] = []
        section_context_lengths: Dict[str, int] = {}
        doc_type_value = job.doc_type or DocType.INGENIERIE
        if hasattr(doc_type_value, "value"):
            doc_type_value = doc_type_value.value
        for item in output.template_outline:
            sec_id = item.id or ""
            if not sec_id.startswith("SEC_"):
                continue
            context_pack = build_section_context(
                title=item.title,
                sec_id=sec_id,
                full_text=sanitized_text,
                doc_tables=doc_tables,
                max_chars=int(os.getenv("AI_TEMPLATE_QUESTION_CONTEXT_CHARS", "3500")),
                docx_sections=docx_sections,
                doc_type=str(doc_type_value),
                techno=job.techno_name,
                allow_empty_context=True,
            )
            context_pack = replace_client_names(context_pack, client_names)
            effective_text = _extract_effective_section_text(context_pack or "")
            effective_len = len(effective_text.strip())
            has_meaningful_text = _has_meaningful_section_text(context_pack or "")
            ctx_len = len((context_pack or "").strip())
            section_context_lengths[sec_id] = ctx_len
            low_context = not has_meaningful_text
            low_context_notice = None
            if low_context:
                job.logs = (job.logs or "") + (
                    f"\nSECTION_SKIP_QUESTIONS sec_id={sec_id} title={item.title!r} ctx_len={ctx_len} effective_len={effective_len}"
                )
                logger.info(
                    "SECTION_SKIP_QUESTIONS sec_id=%s title=%r ctx_len=%s effective_len=%s",
                    sec_id,
                    item.title,
                    ctx_len,
                    effective_len,
                )

            dedup_block = ""
            if all_generated_labels:
                recent = all_generated_labels[-40:]
                dedup_block = "QUESTIONS DEJA POSEES (INTERDIT DE REPETER OU REFORMULER):\n" + "\n".join(
                    f"✗ {x}" for x in recent
                )

            def _combined_instructions(extra_notice: str | None = None) -> str | None:
                combined = "\n".join(
                    filter(
                        None,
                        [low_context_notice, dedup_block, extra_notice],
                    )
                )
                return combined or None

            used_retry = False
            used_fallback = False
            if not has_meaningful_text:
                questions = []
                purpose = "But: contexte vide/non exploitable, aucune question générée."
                questions_started_at = time.perf_counter()
                logger.info(
                    "section_questions_timing sec_id=%s title=%r elapsed_ms=%.2f retry=%s fallback=%s q_count=%s skipped=%s",
                    sec_id,
                    item.title,
                    (time.perf_counter() - questions_started_at) * 1000,
                    False,
                    False,
                    0,
                    True,
                )
            else:
                try:
                    questions_started_at = time.perf_counter()
                    try:
                        questions_output = generate_questions_for_section(
                            sec_id=sec_id,
                            title=item.title,
                            context_pack=context_pack,
                            style_guide=style_guide,
                            doc_type=doc_type_value,
                            techno=job.techno_name,
                            extra_instructions=_combined_instructions(),
                            provider=provider,
                        )
                    except RuntimeError as exc:
                        used_retry = True
                        retry_notice = (
                            "RÉVISION OBLIGATOIRE: respecter uniquement single_choice/multi_choice, "
                            "choices 2..10 et imposer 2..5 questions."
                        )
                        retry_instructions = _combined_instructions(retry_notice)
                        questions_output = generate_questions_for_section(
                            sec_id=sec_id,
                            title=item.title,
                            context_pack=context_pack,
                            style_guide=style_guide,
                            doc_type=doc_type_value,
                            techno=job.techno_name,
                            extra_instructions=retry_instructions,
                            provider=provider,
                        )
                        job.logs = (job.logs or "") + f"\nRetry questions {sec_id}: {exc}"
                        logger.info("[QUALITY_GATE] External retry triggered for section %r (%s): %s", item.title, sec_id, exc)
                    logger.info(
                        "section_questions_timing sec_id=%s title=%r elapsed_ms=%.2f retry=%s fallback=%s q_count=%s",
                        sec_id,
                        item.title,
                        (time.perf_counter() - questions_started_at) * 1000,
                        used_retry,
                        False,
                        len(questions_output.questions),
                    )
                    questions = questions_output.questions
                    purpose = replace_client_names(questions_output.purpose, client_names)
                    for question in questions:
                        question.label = replace_client_names(question.label, client_names)
                        question.choices = [replace_client_names(choice, client_names) for choice in (question.choices or [])]
                except Exception as exc:
                    if low_context:
                        job.logs = (job.logs or "") + (
                            f"\nSECTION_LOW_CONTEXT_FALLBACK sec_id={sec_id} title={item.title!r}"
                        )
                        logger.warning(
                            "SECTION_LOW_CONTEXT_FALLBACK sec_id=%s title=%r error=%s",
                            sec_id,
                            item.title,
                            exc,
                        )
                    used_fallback = True
                    questions = []
                    purpose = "But: aucune question générée (IA en échec)."
                    job.logs = (job.logs or "") + f"\nQuestions fallback {sec_id}: {exc}"
                    logger.warning("[QUALITY_GATE] Fallback triggered for section %r (%s): %s", item.title, sec_id, exc)
                    logger.info(
                        "section_questions_timing sec_id=%s title=%r elapsed_ms=%.2f retry=%s fallback=%s q_count=%s",
                        sec_id,
                        item.title,
                        (time.perf_counter() - questions_started_at) * 1000,
                        used_retry,
                        True,
                        0,
                    )

            signature = _question_signature(questions)
            looks_legacy_generic = _looks_like_legacy_generic_set(questions)
            if not used_fallback and (looks_legacy_generic or _is_overly_repeated_signature(signature, seen_question_signatures)):
                used_retry = True
                dedupe_notice = "RÉVISION DÉDUPLICATION OBLIGATOIRE: ce set de questions est trop similaire à des sections précédentes."
                if looks_legacy_generic:
                    dedupe_notice = (
                        "RÉVISION OBLIGATOIRE: ce set correspond au template générique hérité (domaines/schéma/VLAN/HA/BOM). "
                        "Interdit de réutiliser ce bloc. Génère des questions spécifiques au titre exact et au contenu."
                    )
                retry_instructions = _combined_instructions(dedupe_notice)
                try:
                    questions_output = generate_questions_for_section(
                        sec_id=sec_id,
                        title=item.title,
                        context_pack=context_pack,
                        style_guide=style_guide,
                        doc_type=doc_type_value,
                        techno=job.techno_name,
                        extra_instructions=retry_instructions,
                        provider=provider,
                    )
                    questions = questions_output.questions
                    purpose = replace_client_names(questions_output.purpose, client_names)
                    for question in questions:
                        question.label = replace_client_names(question.label, client_names)
                        question.choices = [replace_client_names(choice, client_names) for choice in (question.choices or [])]
                    signature = _question_signature(questions)
                    job.logs = (job.logs or "") + f"\nRetry dedupe questions {sec_id}: signature répétée"
                except Exception as exc:
                    used_fallback = True
                    questions = []
                    purpose = "But: aucune question générée (IA en échec) — fallback questions désactivé."
                    signature = _question_signature(questions)
                    job.logs = (job.logs or "") + f"\nQuestions fallback dedupe {sec_id}: {exc}"

            seen_norm = {_canonical_text(lbl) for lbl in all_generated_labels[-200:] if (lbl or "").strip()}
            filtered_questions: List[FormQuestionSpec] = []
            for question in questions:
                normalized_label = _canonical_text(question.label or "")
                if normalized_label and normalized_label in seen_norm:
                    continue
                if normalized_label:
                    seen_norm.add(normalized_label)
                filtered_questions.append(question)
            questions = filtered_questions
            signature = _question_signature(questions)
            seen_question_signatures[signature] = seen_question_signatures.get(signature, 0) + 1
            question_types = [((q.type or "unknown").strip().lower() or "unknown") for q in questions]
            logger.info(
                "[AI_OUTPUT] Section %r (%s) -> %s questions (%s)",
                item.title,
                sec_id,
                len(questions),
                ", ".join(question_types),
            )
            logger.info(
                "[QUALITY_GATE] Section %r (%s) -> questions_final_count=%s retry=%s fallback=%s",
                item.title,
                sec_id,
                len(questions),
                used_retry,
                used_fallback,
            )
            purpose = _merge_section_purpose(purpose, getattr(item, "descriptive_text", None))
            output.form.sections.append(
                FormSectionSpec(
                    sec_id=sec_id,
                    purpose=purpose,
                    questions=questions,
                )
            )
            for question in questions:
                if question.label:
                    all_generated_labels.append(question.label)
            stats = _question_stats(questions)
            ctx_preview = (context_pack or "")[:200].replace("\n", " ")
            editorial_flag = is_editorial_section(item.title)
            job.logs = (
                (job.logs or "")
                + (
                    f"\nSECTION_DEBUG sec_id={sec_id} title={item.title!r} editorial={editorial_flag} "
                    f"ctx_len={len(context_pack)} ctx_preview={ctx_preview!r} "
                    f"q_count={stats['q_count']} top3_labels={stats['top_labels']} type_dist={stats['type_dist']}"
                )
            )
            logger.info(
                "section_generation_debug sec_id=%s title=%s editorial=%s ctx_len=%s q_count=%s top3_labels=%s type_dist=%s",
                sec_id,
                item.title,
                editorial_flag,
                len(context_pack),
                stats["q_count"],
                stats["top_labels"],
                stats["type_dist"],
            )
        job.logs = (job.logs or "") + f"\nForm sections générées: {len(output.form.sections)}."
        job.output_payload = output.model_dump()
        job.progress = 55
        job.logs = (job.logs or "") + "\nStructure IA générée."
        db.commit()

        techno = db.query(Techno).filter(Techno.name == job.techno_name).first()
        if not techno:
            techno = Techno(
                name=job.techno_name,
                description=None,
                doc_type=job.doc_type or DocType.INGENIERIE,
                created_by=job.created_by,
            )
            db.add(techno)
            db.commit()
            db.refresh(techno)

        safe_name = _safe_filename(job.techno_name)
        output_dir = AI_TEMPLATES_DIR / f"techno_{techno.id}" / f"job_{job.id}"
        output_path = output_dir / f"{safe_name}_template_ai.docx"
        template_source_path = next((p for p in paths if p.lower().endswith(".docx")), None)
        has_number_prefixes = any(_strip_number_prefix(h)[0] for h in headings)
        template_numbering = not has_number_prefixes
        docx_started_at = time.perf_counter()
        build_docx_from_outline(
            output.template_outline,
            output_path,
            template_source_path,
            template_numbering,
        )
        logger.info(
            "build_docx_from_outline_timing elapsed_ms=%.2f output_path=%s",
            (time.perf_counter() - docx_started_at) * 1000,
            output_path,
        )
        if job.cover_template_id:
            cover_template = (
                db.query(TemplateDoc).filter(TemplateDoc.id == job.cover_template_id).first()
            )
            if cover_template and cover_template.stored_path:
                _merge_cover_template(Path(cover_template.stored_path), output_path)
        template_doc = TemplateDoc(
            techno_id=techno.id,
            filename=output_path.name,
            stored_path=str(output_path),
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            is_generated=True,
            outline_json=output.model_dump(),
            template_text="\n".join([item.content for item in output.template_outline]),
        )
        db.add(template_doc)
        db.commit()
        db.refresh(template_doc)
        job.template_doc_id = template_doc.id
        job.progress = 70
        job.logs = (job.logs or "") + "\nTemplateDoc généré."
        db.commit()

        try:
            for section in output.form.sections:
                outline_item = next((it for it in output.template_outline if it.id == section.sec_id), None)
                section_title = outline_item.title if outline_item else section.sec_id
                logger.info(
                    "pre_form_save_section sec_id=%s title=%r ctx_len=%s q_count=%s",
                    section.sec_id,
                    section_title,
                    section_context_lengths.get(section.sec_id, 0),
                    len(section.questions),
                )
            create_form_started_at = time.perf_counter()
            form = create_form_from_output(
                db=db,
                techno=techno,
                template_doc=template_doc,
                output=output,
                created_by=str(job.created_by) if job.created_by else None,
                client_names=client_names,
            )
            logger.info(
                "create_form_from_output_timing elapsed_ms=%.2f form_id=%s",
                (time.perf_counter() - create_form_started_at) * 1000,
                getattr(form, "id", None),
            )
        except Exception:
            logger.exception("create_form_from_output_failed job_id=%s techno_id=%s", job.id, techno.id)
            raise
        db.commit()
        db.refresh(form)
        inserted_sections = db.query(FormSection).filter(FormSection.form_id == form.id).count()
        job.logs = (job.logs or "") + (
            f"\nForm sections insérées: {inserted_sections} (générées: {len(output.form.sections)})."
        )
        job.form_template_id = form.id
        job.progress = 85
        job.logs = (job.logs or "") + "\nFormTemplate généré."

        techno.template_doc_id = template_doc.id
        techno.form_template_id = form.id
        if job.cover_template_id:
            techno.cover_template_id = job.cover_template_id
        db.commit()

        job.techno_id = techno.id
        job.status = AITemplateJobStatus.DONE
        job.progress = 100
        job.logs = (job.logs or "") + "\nJob terminé."
        logger.info(
            "run_ai_template_job_timing job_id=%s status=%s elapsed_ms=%.2f",
            job.id,
            job.status,
            (time.perf_counter() - job_started_at) * 1000,
        )
        db.commit()
    except Exception as exc:
        job.status = AITemplateJobStatus.FAILED
        job.error_message = str(exc)
        job.logs = (job.logs or "") + f"\nErreur: {exc}"
        logger.info(
            "run_ai_template_job_timing job_id=%s status=%s elapsed_ms=%.2f",
            job.id,
            job.status,
            (time.perf_counter() - job_started_at) * 1000,
        )
        db.commit()
        return
