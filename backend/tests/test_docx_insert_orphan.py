import json
import os
import re
import sys
import tempfile
import types
import unittest
from pathlib import Path

from docx import Document

ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
sys.path.insert(0, ROOT_DIR)

if "openpyxl" not in sys.modules:
    fake_openpyxl = types.ModuleType("openpyxl")
    fake_openpyxl.load_workbook = lambda *args, **kwargs: None
    sys.modules["openpyxl"] = fake_openpyxl

from app.services.docx_pipeline import apply_doc_pipeline


class DocxInsertOrphanTests(unittest.TestCase):
    def _render_with_insert_html(self, build_template, insert_html):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp = Path(tmpdir)
            template_path = tmp / "template.docx"
            out_path = tmp / "out.docx"

            doc = build_template()
            doc.save(template_path)

            apply_doc_pipeline(
                template_path=str(template_path),
                out_path=str(out_path),
                sections_text={},
                titles_to_remove=[],
                worksheet_path=None,
                insert_html=insert_html,
            )

            return Document(str(out_path))

    def _render(self, build_template, insert_map):
        return self._render_with_insert_html(
            build_template,
            json.dumps({"__type": "map", "items": insert_map}),
        )

    def test_orphan_insert_paragraph_before_any_heading_is_replaced(self) -> None:
        def build_template():
            d = Document()
            d.add_paragraph("[[insérer Schéma réseau]]")
            d.add_heading("1. Contexte", level=1)
            d.add_paragraph("Texte de section")
            return d

        out_doc = self._render(
            build_template,
            {"insérer schéma réseau": "<p>Contenu orphelin inséré</p>"},
        )

        paragraph_texts = [p.text for p in out_doc.paragraphs if (p.text or "").strip()]
        full_text = "\n".join(paragraph_texts)

        self.assertIn("Contenu orphelin inséré", full_text)
        self.assertIsNone(re.search(r"\[\[\s*ins[eé]rer", full_text, flags=re.IGNORECASE))

    def test_orphan_insert_uses_placeholder_key_even_when_not_in_section(self) -> None:
        def build_template():
            d = Document()
            d.add_paragraph("[[insérer schéma hors section]]")
            d.add_heading("1. Contexte", level=1)
            d.add_paragraph("[[SEC_CTX]]")
            return d

        out_doc = self._render(
            build_template,
            {
                "schéma hors section": "<p>Remplacement sans section_id</p>",
                "1. contexte": "<p>Contenu section</p>",
            },
        )

        paragraph_texts = [p.text for p in out_doc.paragraphs if (p.text or "").strip()]
        full_text = "\n".join(paragraph_texts)

        self.assertIn("Remplacement sans section_id", full_text)
        self.assertNotIn("Contenu section", full_text)
        self.assertIsNone(re.search(r"\[\[\s*ins[eé]rer", full_text, flags=re.IGNORECASE))

    def test_orphan_insert_accepts_non_accented_key_variant(self) -> None:
        def build_template():
            d = Document()
            d.add_paragraph("[[insérer schéma wan]]")
            return d

        out_doc = self._render(
            build_template,
            {"schema wan": "<p>Contenu clé sans accent</p>"},
        )

        paragraph_texts = [p.text for p in out_doc.paragraphs if (p.text or "").strip()]
        full_text = "\n".join(paragraph_texts)

        self.assertIn("Contenu clé sans accent", full_text)
        self.assertIsNone(re.search(r"\[\[\s*ins[eé]rer", full_text, flags=re.IGNORECASE))

    def test_orphan_insert_accepts_insert_keyword_variant_before_any_heading(self) -> None:
        def build_template():
            d = Document()
            d.add_paragraph("[[insert network diagram]]")
            d.add_heading("1. Contexte", level=1)
            return d

        out_doc = self._render(
            build_template,
            {"network diagram": "<p>Insert keyword replacement</p>"},
        )

        paragraph_texts = [p.text for p in out_doc.paragraphs if (p.text or "").strip()]
        full_text = "\n".join(paragraph_texts)

        self.assertIn("Insert keyword replacement", full_text)
        self.assertNotRegex(full_text, r"\[\[\s*(?:ins[eé]rer|insert(?:ion)?)")

    def test_orphan_insert_with_null_section_id_payload_is_replaced(self) -> None:
        def build_template():
            d = Document()
            d.add_paragraph("[[insérer schéma hors section]]")
            d.add_heading("1. Contexte", level=1)
            d.add_paragraph("[[SEC_CTX]]")
            return d

        payload = json.dumps(
            {
                "insertions": [
                    {
                        "section_id": None,
                        "placeholder": "schéma hors section",
                        "html": "<p>Contenu insertion section_id null</p>",
                    }
                ]
            }
        )

        out_doc = self._render_with_insert_html(build_template, payload)

        paragraph_texts = [p.text for p in out_doc.paragraphs if (p.text or "").strip()]
        full_text = "\n".join(paragraph_texts)

        self.assertIn("Contenu insertion section_id null", full_text)
        self.assertIsNone(re.search(r"\[\[\s*ins[eé]rer", full_text, flags=re.IGNORECASE))

    def test_orphan_insert_inside_table_before_any_heading_is_replaced(self) -> None:
        def build_template():
            d = Document()
            t = d.add_table(rows=1, cols=1)
            t.cell(0, 0).text = "[[insérer tableau orphelin]]"
            d.add_heading("1. Contexte", level=1)
            d.add_paragraph("Suite")
            return d

        out_doc = self._render(
            build_template,
            {"insérer tableau orphelin": "<p>Valeur cellule injectée</p>"},
        )

        cell_texts = []
        for table in out_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_texts.append("\n".join(p.text for p in cell.paragraphs))

        full_cells = "\n".join(cell_texts)
        self.assertIn("Valeur cellule injectée", full_cells)
        self.assertNotRegex(full_cells, r"\[\[\s*ins[eé]rer")


if __name__ == "__main__":
    unittest.main()
