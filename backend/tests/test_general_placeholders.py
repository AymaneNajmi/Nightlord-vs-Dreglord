from pathlib import Path

from docx import Document

from app.services.docx_ops import docx_general_placeholders
from app.services.docx_pipeline import apply_doc_pipeline


PNG_1X1 = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\x0cIDATx\x9cc``\x00\x00\x00\x02\x00\x01"
    b"\x0b\xe7\x02\x9b\x00\x00\x00\x00IEND\xaeB`\x82"
)


def test_docx_general_placeholders_extracts_ordered_text_and_logo(tmp_path: Path) -> None:
    src = tmp_path / "src.docx"
    d = Document()
    d.add_paragraph("[[TEXT: Nom du client]]")
    d.add_paragraph("X [[TEXT: Description\t projet ]] Y")
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "[[LOGO]]"
    d.save(src)

    out = docx_general_placeholders(str(src))

    assert out["has_logo"] is True
    assert out["text_placeholders"] == [
        {"id": "text_1", "description": "Nom du client", "raw": "[[TEXT: Nom du client]]"},
        {"id": "text_2", "description": "Description\t projet", "raw": "[[TEXT: Description\t projet ]]"},
    ]


def test_apply_doc_pipeline_replaces_text_and_logo(tmp_path: Path) -> None:
    src = tmp_path / "tpl.docx"
    logo = tmp_path / "logo.png"
    logo.write_bytes(PNG_1X1)

    d = Document()
    d.add_paragraph("[[TEXT: Nom du client]]")
    d.add_paragraph("Avant [[TEXT: Adresse]] Après")
    d.add_paragraph("[[LOGO]]")
    d.save(src)

    out = tmp_path / "out.docx"
    apply_doc_pipeline(
        template_path=str(src),
        out_path=str(out),
        sections_text={},
        titles_to_remove=[],
        text_placeholder_values={"text_1": "ACME\nFrance", "text_2": "42 rue"},
        logo_path=str(logo),
    )

    result = Document(out)
    assert "ACME" in result.paragraphs[0].text
    assert "France" in result.paragraphs[0].text
    assert result.paragraphs[1].text == "Avant 42 rue Après"
    assert len(result.inline_shapes) == 1
