import os
import sys
import tempfile
import unittest

from docx import Document

ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
sys.path.insert(0, ROOT_DIR)

from app.schemas.ai_template_builder import FormQuestionSpec
from app.services.question_quality import enforce_question_quality
from app.services.section_context import (
    build_section_context,
    extract_sections_from_docx_by_headings,
)


class SectionContextTests(unittest.TestCase):
    def test_build_section_context_uses_heading_extraction(self) -> None:
        doc = Document()
        doc.add_heading("1.4 Périmètre", level=2)
        doc.add_paragraph(
            "Le scope couvre les équipements Catalyst 9404 et 9200L, "
            "les Firepower 2130, la gestion via FMC et ISE, "
            "ainsi que la supervision Prime."
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "sample.docx")
            doc.save(path)
            docx_sections = extract_sections_from_docx_by_headings(path)

        context_pack = build_section_context(
            title="1.4 Périmètre",
            sec_id="SEC_1_4",
            full_text="",
            doc_tables=[],
            max_chars=5000,
            docx_sections=docx_sections,
        )

        for token in ["Catalyst 9404", "9200L", "Firepower 2130", "FMC", "ISE", "Prime"]:
            self.assertIn(token, context_pack)

    def test_quality_gate_accepts_choice_only_contract(self) -> None:
        questions = [
            FormQuestionSpec(
                key="objectif_principal",
                label="Quel objectif principal doit être couvert ?",
                type="single_choice",
                choices=["Conception", "Migration", "Audit"],
                required=True,
            ),
            FormQuestionSpec(
                key="sites",
                label="Quels sites sont concernés ?",
                type="multi_choice",
                choices=["SITE_A", "SITE_B"],
                required=True,
            ),
        ]

        enforce_question_quality(questions, True, "Contexte")

    def test_quality_gate_rejects_generic_labels(self) -> None:
        questions = [
            FormQuestionSpec(
                key="q1",
                label="Détails pour la section",
                type="single_choice",
                choices=["A", "B"],
                required=True,
            )
        ]

        with self.assertRaises(RuntimeError):
            enforce_question_quality(questions, True, "Contexte")


if __name__ == "__main__":
    unittest.main()
