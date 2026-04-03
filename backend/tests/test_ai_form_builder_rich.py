import json
import os
import sys
import tempfile
import unittest

from docx import Document

ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
sys.path.insert(0, ROOT_DIR)

from app.services.ai_form_builder_rich import (  # noqa: E402
    build_section_context,
    extract_outline_from_docx,
    _extract_candidate_options,
    _is_option_grounded,
    validate_section_output,
)


class AIFormBuilderRichTests(unittest.TestCase):
    def test_extract_outline_from_docx(self):
        doc = Document()
        doc.add_heading("Titre principal", level=1)
        doc.add_heading("Contexte", level=2)
        doc.add_paragraph("Option A")
        doc.add_paragraph("Option B")
        doc.add_heading("Détail", level=3)
        doc.add_paragraph("Choix X")

        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "sample.docx")
            doc.save(path)
            sections = extract_outline_from_docx(path)

        self.assertEqual(len(sections), 2)
        self.assertEqual(sections[0].section_key, "SEC_1")
        self.assertEqual(sections[0].level, "H2")
        self.assertIn("Option A", sections[0].paragraphs)
        self.assertEqual(sections[1].level, "H3")

    def test_extract_outline_from_docx_supports_localized_heading_names(self):
        doc = Document()
        h2 = doc.add_paragraph("Contexte")
        h2.style = "Normal"
        h2.style.name = "Titre 2"
        doc.add_paragraph("Option A")
        h3 = doc.add_paragraph("Sous-section")
        h3.style = "Normal"
        h3.style.name = "Titre 3"
        doc.add_paragraph("Option B")

        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "localized.docx")
            doc.save(path)
            sections = extract_outline_from_docx(path)

        self.assertEqual(len(sections), 2)
        self.assertEqual(sections[0].section_key, "SEC_1")
        self.assertEqual(sections[1].section_key, "SEC_1_1")

    def test_validation_gate_rejects_unknown_options(self):
        section_data = {
            "section_key": "SEC_1",
            "title": "Contexte",
            "level": "H2",
            "intent": "Cadrer les choix disponibles.",
            "example": "Sélectionner une option existante.",
            "questions": [
                {
                    "order": 0,
                    "label": "Quel choix ?",
                    "qtype": "single_choice",
                    "is_required": True,
                    "options": [
                        {"order": 0, "value": "invented-option-1", "source_quote": "invented-option-1"},
                        {"order": 1, "value": "invented-option-2", "source_quote": "invented-option-2"},
                    ],
                }
            ],
        }

        with self.assertRaises(RuntimeError):
            validate_section_output(section_data, context="Texte: Option A et Option B")

    def test_parsing_json_strict_payload_shape(self):
        raw = json.dumps(
            {
                "section_key": "SEC_2",
                "title": "Objectif",
                "level": "H2",
                "intent": "Définir le périmètre.",
                "example": "Exemple: migration LAN.",
                "questions": [],
            }
        )
        parsed = json.loads(raw)
        validated = validate_section_output(parsed, context="", context_empty=True)
        self.assertEqual(validated["section_key"], "SEC_2")
        self.assertEqual(validated["questions"], [])

    def test_validation_accepts_exact_source_quote(self):
        section_data = {
            "section_key": "SEC_3",
            "title": "Contexte",
            "level": "H2",
            "intent": "Cadrer les choix disponibles.",
            "example": "Sélectionner une option exacte.",
            "questions": [
                {
                    "order": 0,
                    "label": "Quel site ?",
                    "qtype": "single_choice",
                    "is_required": True,
                    "options": [
                        {"order": 0, "value": "SITE_A", "source_quote": "SITE_A"},
                        {"order": 1, "value": "SITE_B", "source_quote": "SITE_B"},
                    ],
                }
            ],
        }

        validated = validate_section_output(section_data, context="Texte: SITE_A et SITE_B")
        self.assertEqual(validated["section_key"], "SEC_3")

    def test_extract_candidate_options(self):
        context = """Titre: Contexte\n\nTexte:\nSITE_A, SITE_B; NGFW/ISE\nMode: HA\n"""
        candidates = _extract_candidate_options(context)
        self.assertIn("SITE_A", candidates)
        self.assertIn("SITE_B", candidates)

    def test_is_option_grounded_accepts_containment_variants(self):
        normalized_context = "texte: proteger la zone serveurs par une solution pare feu de nouvelle generation"
        self.assertTrue(
            _is_option_grounded(
                normalized_context,
                option_value="pare-feu de nouvelle génération",
                source_quote="solution pare feu de nouvelle generation",
            )
        )

    def test_validation_keeps_single_option_question(self):
        section_data = {
            "section_key": "SEC_4",
            "title": "Risques",
            "level": "H2",
            "intent": "Qualifier les risques.",
            "example": "Risque d'infection.",
            "questions": [
                {
                    "order": 0,
                    "label": "Conséquence principale ?",
                    "qtype": "single_choice",
                    "is_required": True,
                    "options": [
                        {"order": 0, "value": "des attaques", "source_quote": "des attaques"},
                    ],
                }
            ],
        }

        validated = validate_section_output(
            section_data,
            context="Texte: des attaques sur le LAN non isolé",
        )
        self.assertEqual(len(validated["questions"]), 1)

    def test_validation_can_drop_invalid_question_when_enabled(self):
        section_data = {
            "section_key": "SEC_4",
            "title": "Risques",
            "level": "H2",
            "intent": "Qualifier les risques.",
            "example": "Risque d'infection.",
            "questions": [
                {
                    "order": 0,
                    "label": "Quelles conséquences sont observées ?",
                    "qtype": "single_choice",
                    "is_required": True,
                    "options": [],
                }
            ],
        }

        validated = validate_section_output(
            section_data,
            context="Texte: des attaques sur le LAN non isolé",
            drop_invalid_questions=True,
        )
        self.assertEqual(validated["questions"], [])


if __name__ == "__main__":
    unittest.main()
