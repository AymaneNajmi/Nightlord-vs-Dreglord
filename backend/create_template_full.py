import os
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

# ==========================================
# FONCTIONS UTILITAIRES OXML
# ==========================================

def add_dynamic_field(paragraph, field_code):
    """Ajoute un champ Word classique (Table des matières, Numéro de page)"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_code
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.extend([fldChar1, instrText, fldChar2, fldChar3])

def add_page_borders(doc, color_hex):
    """Ajoute un cadre (bordure) à toutes les pages du document via OXML"""
    for section in doc.sections:
        sectPr = section._sectPr
        existing_borders = sectPr.find(qn('w:pgBorders'))
        if existing_borders is not None:
            sectPr.remove(existing_borders)
            
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')     # Taille de la bordure
            border.set(qn('w:space'), '24')  # Marge
            border.set(qn('w:color'), color_hex)
            pgBorders.append(border)
            
        docGrid = sectPr.find(qn('w:docGrid'))
        if docGrid is not None:
            docGrid.addprevious(pgBorders)
        else:
            sectPr.append(pgBorders)

def add_page_numbers(doc):
    """Ajoute la numérotation au centre du pied de page"""
    for section in doc.sections:
        footer = section.footer
        for p in footer.paragraphs:
            p.text = ""
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run("Page ")
        add_dynamic_field(paragraph, ' PAGE \\* MERGEFORMAT ')
        paragraph.add_run(" / ")
        add_dynamic_field(paragraph, ' NUMPAGES \\* MERGEFORMAT ')

def force_style_color(style, rgb_color):
    """Force la couleur RGB en supprimant l'héritage du thème de Word"""
    style.font.color.rgb = rgb_color
    rPr = style.font.element
    color_elements = rPr.findall(qn('w:color'))
    if color_elements:
        for color_element in color_elements:
            for attr in ['themeColor', 'themeTint', 'themeShade']:
                q_attr = qn(f'w:{attr}')
                if q_attr in color_element.attrib:
                    del color_element.attrib[q_attr]

# ==========================================
# CONFIGURATION DES STYLES ET COULEURS
# ==========================================

def configure_styles(doc):
    """Paramètre les polices, couleurs, tailles, espacements et alinéas (jusqu'au niveau 5)"""
    
    COLOR_TITRE_DOC      = RGBColor(0, 51, 102)     # Bleu Nuit
    COLOR_CHAPITRE       = RGBColor(0, 51, 102)     # Chapitre (H1)
    COLOR_SECTION        = RGBColor(0, 102, 204)    # Section (H2)
    COLOR_SUBSECTION     = RGBColor(0, 153, 153)    # Sous-section (H3)
    COLOR_SUBSUBSECTION  = RGBColor(102, 102, 102)  # Sous-sous-section (H4)
    COLOR_SUBSUBSUBSEC   = RGBColor(153, 0, 0)      # Sous-sous-sous-section (H5)

    style_title = doc.styles['Title']
    style_title.font.name = 'Arial'
    style_title.font.size = Pt(26)
    style_title.font.bold = True
    force_style_color(style_title, COLOR_TITRE_DOC)
    style_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_title.paragraph_format.space_after = Pt(24)

    heading_configs = [
        (1, 16, COLOR_CHAPITRE,      0.0, 24, 12),
        (2, 14, COLOR_SECTION,       0.5, 18, 6),
        (3, 12, COLOR_SUBSECTION,    1.0, 12, 6),
        (4, 11, COLOR_SUBSUBSECTION, 1.5, 12, 6),
        (5, 10, COLOR_SUBSUBSUBSEC,  2.0, 10, 4)
    ]

    for level, size, color, indent, sp_before, sp_after in heading_configs:
        style = doc.styles[f'Heading {level}']
        style.font.name = 'Arial'
        style.font.size = Pt(size)
        style.font.bold = True
        
        force_style_color(style, color)
        
        style.paragraph_format.space_before = Pt(sp_before)
        style.paragraph_format.space_after = Pt(sp_after)
        style.paragraph_format.left_indent = Cm(indent)

def style_table(table):
    """Applique un style de grille basique aux tableaux"""
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ==========================================
# FONCTION PRINCIPALE
# ==========================================

def create_full_template():
    doc = Document()
    
    # Configuration globale
    configure_styles(doc)
    
    COLOR_CADRE_PAGE = "003366"  # Bleu Nuit
    add_page_borders(doc, color_hex=COLOR_CADRE_PAGE) 
    
    add_page_numbers(doc)
    
    # ==========================================
    # 2. PAGE DE GARDE
    # ==========================================
    for _ in range(3): doc.add_paragraph()
    
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.add_run("[[ INSÉRER LE LOGO ICI ]]").bold = True
    
    for _ in range(2): doc.add_paragraph()
    doc.add_heading("Document d’ingénierie", level=0)
    for _ in range(2): doc.add_paragraph()
    
    table_garde = doc.add_table(rows=4, cols=2)
    style_table(table_garde)
    items_garde = [("Nom du Client", "[[TEXT:Client]]"), 
                   ("Nom du Projet", "[[TEXT:Projet]]"), 
                   ("Date", "[[TEXT:Date]]"), 
                   ("Version", "[[TEXT:Version]]")]
    for i, (label, value) in enumerate(items_garde):
        table_garde.cell(i, 0).text = label
        table_garde.cell(i, 0).paragraphs[0].runs[0].bold = True
        table_garde.cell(i, 1).text = value

    doc.add_page_break()

    # ==========================================
    # 3. PROPRIÉTÉS DU DOCUMENT (Tableaux)
    # ==========================================
    doc.add_heading("Propriétés du Document", level=1)
    
    doc.add_heading("Rédacteurs", level=2)
    t_redacteurs = doc.add_table(rows=2, cols=3)
    style_table(t_redacteurs)
    for i, header in enumerate(["Nom", "Fonction", "Contact"]):
        t_redacteurs.cell(0, i).text = header
        t_redacteurs.cell(0, i).paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    doc.add_heading("Historique des modifications", level=2)
    t_historique = doc.add_table(rows=2, cols=3)
    style_table(t_historique)
    for i, header in enumerate(["Date", "Vérificateur", "Nature du changement"]):
        t_historique.cell(0, i).text = header
        t_historique.cell(0, i).paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    doc.add_heading("Diffusion", level=2)
    t_diffusion = doc.add_table(rows=2, cols=2)
    style_table(t_diffusion)
    for i, header in enumerate(["Entité", "Destinataires"]):
        t_diffusion.cell(0, i).text = header
        t_diffusion.cell(0, i).paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ==========================================
    # 4. SOMMAIRES DYNAMIQUES (Standard)
    # ==========================================
    doc.add_heading("SOMMAIRE", level=1)
    add_dynamic_field(doc.add_paragraph(), 'TOC \\o "1-5" \\h \\z \\u')
    p_note1 = doc.add_paragraph("(Note: Faites un clic droit et sélectionnez 'Mettre à jour les champs' dans Word)")
    p_note1.style.font.italic = True
    doc.add_page_break()

    doc.add_heading("TABLE DES FIGURES", level=1)
    add_dynamic_field(doc.add_paragraph(), 'TOC \\h \\z \\c "Figure"')
    doc.add_page_break()

    doc.add_heading("TABLE DES TABLEAUX", level=1)
    add_dynamic_field(doc.add_paragraph(), 'TOC \\h \\z \\c "Table"')
    doc.add_page_break()

    # ==========================================
    # 5. CONTENU PRINCIPAL (Parseur)
    # ==========================================
    raw_toc = """
1   APERÇU DU PROJET
1.1 OBJECTIF DU DOCUMENT
1.2 CONTEXTE
1.3 OBJECTIFS DE CONCEPTION
1.4 PERIMETRE
2   INTRODUCTION : ETUDE DE L’EXISTANT
2.1 APERÇU DU RESEAU EXISTANT
2.2 BLOCK CAMPUS
2.3 BLOCK WAN
2.4 BLOCK INTERNET
2.5 BLOCK CENTRE DE DONNEE
3   ARCHITECTURE RESEAU CIBLE
3.1 CONCEPTION DE L’ARCHITECTURE CIBLE
3.1.1   Conception hiérarchique deux tiers :
3.1.2   Conception modulaire :
3.1.3   Avantage de la Technologie Stack-Wise :
3.1.4   Conception de la couche CORE/Distribution :
3.2 TOPOLOGIE RESEAU CIBLE
3.3 LISTE DES EQUIPEMENTS PAR SITE
3.3.1   Aspect HARDWARE des équipements
3.3.1.1 Cisco Catalyst 9400 Séries châssis
3.3.1.2 Cisco Catalyst 9200 Séries châssis
3.3.2   Aspect fonctionnel des équipements
3.4 INFORMATION RACKAGE
3.4.1   Conditions environnementales
3.4.2   Site Centrale
3.4.3   Site DAFCI
3.5 CONVENTION DE DENOMINATION
3.5.1   HOSTNAMES
3.5.1.1 Site Centrale
3.5.1.2 Site Dafci
3.6 INFORMATION CABLAGE
3.6.1   Site centrale et Dafci
3.7 STACK-WISE-VIRTUAL SUR LES CATALYST 94004
3.7.1   Présentation de Stack-Wise Virtual
3.7.2   Prérequis pour Cisco Stack-Wise Virtual
3.7.3   Restrictions pour Cisco Stack-Wise Virtual
3.7.4   Architecture Stack-Wise virtuel
3.7.4.1 Adresses MAC Stack-Wise virtuel
3.7.4.2 Lien virtuel Stack-Wise
3.7.4.3 Conversion en mode virtuel Stack-Wise des deux c9404R
3.8 STACK-WISE-80 SUR LES CATALYST C9200L
3.8.1   Présentation de la pile de commutateurs
3.8.2   Restrictions pour les piles de commutateurs
3.8.3   Architecture empilable
3.8.3.1 Découverte de la pile
3.8.3.2 Élection ACTIVE
3.8.4   Conversion en mode Stack-Wise 80 des quatre C9200L-24T-4X-A
3.9 CONCEPTION DE LA COUCHE 1
3.9.1   Le paramètre global couche 1
3.9.2   Vitesse d'interface / duplex
3.9.3   Topologie réseau physique
3.10    CONCEPTION DE LA COUCHE 2
3.10.1  Les paramètres globaux niveau 2
3.10.1.1    Les vlan
3.10.1.2    Le protocole VTP
3.10.1.3    Conception du protocole Spanning Tree
3.10.1.3.1  Placement du Root Spanning-tree
3.10.1.3.2  Fonctionnalité spanning-tree Portfast :
3.10.1.3.3  Fonctionnalité spanning-tree BPDU Guard
3.10.1.3.4  Root Guard
3.10.1.3.5  Loop Guard
3.10.1.4    Ether-Channel
3.10.1.4.1  Ether-Channel Mode On
3.10.1.4.2  Le protocole PAgP
3.10.1.4.3  Le protocole LACP
3.10.1.5    Liens Ether-Channel multi châssis
3.10.2  Topologie réseau logique niveau 2
3.11    CONCEPTION DE LA COUCHE 3
3.11.1  Les paramètres globaux niveau 3
3.11.1.1    Les VRF
3.11.1.2    Liste VRF
3.11.1.3    Création de VRF et association avec une interface :
3.11.2  Schéma d’adressage IP
3.11.3  Serveur DHCP :
3.11.4  Les Passerelles
3.11.5  Le Routage
3.11.6  Route par défaut
3.11.7  Management de system
3.11.7.1    Adressage de management
3.11.7.1    NTP
3.11.7.1    Banner
3.11.7.1    Logging
3.11.7.1    SNMP
3.11.7.1    AAA
4   DESIGN DE LA SOLUTION ISE
OBJECTIF DU DESIGN DE LA SOLUTION ISE
4.1 APERÇU DE CISCO ISE (IDENTITY SERVICES ENGINE)
4.2 EXPRESSION DES BESOINS FONCTIONNELS DE PROSUMA
4.3 VERSION ISE CIBLE
4.4 LISTE LOGICIEL ET LICENSE (BOM)
4.5 PLATEFORME MATERIELLE ET LOGICIEL ISE
4.6 ARCHITECTURE CIBLE DE LA SOLUTION CISCO ISE
4.7 CONNECTIVITE PHYSIQUE & PLAN D’ADRESSAGE CIBLE
4.8 PARAMETRES GENERAUX DES DEUX APPLIANCES
4.9 LISTE DES NADS (NETWORK ACCESS DEVICE)
4.10    MATRICE DE FLUX (A AUTORISER SUR LES FIREWALLS)
4.11    SERVEURS DE SERVICE
4.12    STRATEGIES CISCO ISE CIBLE
4.12.1  Intégration avec Active Directory
4.13    SERVICE D’AUTHENTIFICATION
4.13.1  Strategie d’authentification PROSUMA :
4.13.2  Règle d’authentification
4.14    STRATEGIE D’AUTORISATION
4.14.1  Autorisation VLAND ID
4.15    STRATEGIE DE PROFILING
4.16    INVENTAIRE DES OS
4.17    STRATEGIE DE POSTURE
4.18    RECAP DES RÉGLE DE POSTURE
5   DESIGN DE LA SOLUTION DE SECURITE CENTRE DE DONNEE
OBJECTIF DU DESIGN SECURITE
SCOPE DE LA PARTIE SECURITE
DOCUMENTS DE REFERENCES
5.1 DESCRIPTION TECHNIQUE DE L’ARCHITECTURE ACTUELLE
5.1.1   Architecture actuelle
5.2 DESCRIPTION TECHNIQUE DE LA SOLUTION CIBLE
5.2.1   Architecture Physique cible
5.2.2   Descriptif de l’architecture cible
5.2.3   Digramme du flux Inspecté avec l’architecture cible
5.3 DEPLOIEMENT DU CISCO FIREPOWER MANAGEMENT CENTER
5.4 CONFIGURATION DU FIREWALL CISCO FIREPOWER
5.4.1   Conditions Environnements
5.4.2   Connectique physique
5.4.3   Rappel des licences
5.4.3.1 Base Licenses :
5.4.4   Haute disponibilité
A.  ACTIF /ACTIF
B.  ACTIF/PASSIF
5.4.5   Configuration NTP
5.5 CONFIGURATION DNS
5.6 CONFIGURATION SNMP
5.7 INTEGRATION SYSLOG
5.8 CONFIGURATION NETFLOW
5.9 CONTROL D’ACCESS AU FIREWALL
A.  LE PROFILE ADMINISTRATEUR
B.  METHODE D’AUTHENTIFICATION
5.10    SEGMENTATION PAR ZONES
5.11    ADRESSAGE DES INTERFACES
5.12    LE ROUTAGE
5.13    NETWORK DISCOVERY POLICIES
5.14    LE SYSTEME PREVENTION D’INTRUSION (IPS)
5.14.1  Variable Set
5.14.2  Définition des signatures IPS
5.15    CONFIGURER LES SERVICES PERSONNALISES
5.16    CONFIGURER LES OBJECTS D’ADRESSE
5.17    CONFIGURER LES POLITIQUES DE SECURITE
6   PRIME INFRASTRUCTURE
6.1 APERÇU
6.2 POINTS FORTS DE CISCO PRIME INFRASTRUCTURE
6.3 PRIME INFRASTRUCTURE SE PRESENTE SOUS DEUX FORMES PRINCIPALES :
6.4 OPTION D’INSTALLATION DE PRIME INFRASTRUCTURE :
6.5 LES LICENCES PRIME INFRASTRUCTURE
6.5.1   Les types de licence
6.5.2   Mode de licence ACTIF : licence traditionnelle
6.6 BOM
6.7 LES PORTS UTILISES PAR PRIME INFRASTRUCTURE ET ASSURANCE
6.8 CONFIGURER ET UTILISER LES TABLEAUX DE BORD
6.9 TACHES DE CONFIGURATION DU SERVEUR
6.9.1   Vérifiez les paramètres de sauvegarde
6.9.1.1 Planifier des sauvegardes automatiques des applications
6.9.1.2 Spécifier le référentiel de sauvegarde pour les sauvegardes automatiques
6.9.1.3 Utiliser un référentiel de sauvegarde à distance
6.9.2   Personnaliser les interruptions « traps » SNMP internes du serveur et transférer les interruptions
6.9.3   Activer le service FTP / TFTP / SFTP sur le serveur
6.9.4   Configurer le serveur de messagerie SMTP
6.9.5   Contrôle des tâches de collecte de données
6.9.5.1 Comment les paramètres de conservation des données affectent les données de l'interface graphique Web
6.9.5.2 À propos de la conservation des données historiques
6.9.5.3 Purge d'alarme, d'événement et de Syslog
6.9.5.4 Purge du journal
6.9.6   Configurer l'authentification locale
6.9.7   Configurer les stratégies de mot de passe globales pour l'authentification locale
6.9.8   Groupes d’utilisateurs d’interface WEB :
6.9.8.1 Tâches de configuration de la gestion des utilisateurs
6.9.8.2 Désactiver et activer l'utilisateur racine de l'interface graphique Web
6.10    FONCTIONNALITES.
6.10.1  Les domaines virtuels
6.10.2  Gestion de l'inventaire
6.10.2.1    Méthodes d'ajout d'appareils :
6.10.2.2    Comprendre le processus de découverte :
6.10.2.1    Découverte rapide :
6.10.2.2    Groupe d’équipement :
6.10.3  Gestion de la Configuration.
6.10.4  Surveillance et dépannage.
6.10.5  Gestion d’installation des images.
6.10.6  Les Rapports.
6.10.7  Les Notifications.
6.11    INSTALLATION DE PRIME INFRASTRUCTURE EN UTILISANT HYPER-V :
6.12    NAVIGATION POUR ACCEDER AUX FONCTIONNALITES :
7   CONCEPTION DE LA SECURITE CENTRE DE DONNEE
8   BOM
9   REFERENCES
REFERENCES DE CE DOCUMENT
10  ANNEXE A : TEMPLATE DE CONFIGURATION
VLAN
Access-List:  Management Access
AAA
11  ANNEXE A : HARDWARE
12  ANNEXE B : ACCEPTATION DOCUMENT
"""
    # Note: J'ai raccourci raw_toc ici pour l'exemple. 
    # Remettez votre liste complète à cette place.

    for line in raw_toc.strip().split('\n'):
        line = line.strip()
        if not line:
            continue
            
        match = re.match(r'^([\d\.]+)?(?:[\s\t]+)?(.*)$', line)
        if not match:
            continue
            
        number = match.group(1)
        title = match.group(2).strip()
        
        if number:
            level = number.count('.') + 1
            if level > 1 and not number.endswith('.'):
                pass
            sec_id_str = number.replace('.', '_').strip('_')
            full_title = f"{number} {title}"
        elif re.match(r'^[A-Z]\.', title): 
            level = 2
            sec_id_str = title[0]
            full_title = title
        else:
            level = 1
            sec_id_str = title.replace(' ', '_').replace('\'', '')[:10]
            full_title = title

        doc.add_heading(full_title, level=min(level, 5))
        
        # --- RETOUR AUX BALISES TEXTE SIMPLES ---
        txt_lower = title.lower()
        if any(word in txt_lower for word in ["bom", "liste des", "matrice", "inventaire"]):
            doc.add_paragraph(f"[[Excel: {title[:20].strip()}]]")
        elif any(word in txt_lower for word in ["architecture", "topologie", "schéma", "digramme", "diagramme"]):
            doc.add_paragraph(f"[[insérer Figure {title[:20].strip()}]]")
        
        doc.add_paragraph(f"[[SEC_{sec_id_str}]]")

    out_path = os.path.abspath("template_ingenierie.docx")
    doc.save(out_path)
    print(f"Document d'ingénierie créé avec succès : {out_path}")

if __name__ == "__main__":
    create_full_template()